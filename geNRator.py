import tkinter as tk
from tkinter import ttk, messagebox, font, filedialog
from docx import Document
from docx.shared import Inches
from datetime import datetime
import pandas as pd
import os
import json
import threading
import time
from openpyxl import load_workbook

# === Building Selection Function ===
def select_building():
    # Don't show root window yet - wait until UI is fully configured
    # root.deiconify()  # Removed - will show window after UI is configured
    building_window = tk.Toplevel()
    building_window.title("Select Building")
    building_window.configure(bg="black")
    # Center the window on the screen
    window_width = 400
    window_height = 350  # Increased height for Red Gym option
    screen_width = building_window.winfo_screenwidth()
    screen_height = building_window.winfo_screenheight()
    x = int((screen_width / 2) - (window_width / 2))
    y = int((screen_height / 2) - (window_height / 2))
    building_window.geometry(f"{window_width}x{window_height}+{x}+{y}")
    
    # Make sure building window is visible and focused
    building_window.lift()
    building_window.focus_force()
    
    # macOS specific: Force window to be topmost temporarily
    building_window.attributes('-topmost', True)
    building_window.update()
    building_window.attributes('-topmost', False)
    
    # Handle window close event - exit application if building selection is closed
    def on_building_close():
        stop_autosave()  # Stop any running autosave
        root.quit()  # Exit the mainloop
        root.destroy()  # Destroy the root window
    
    building_window.protocol("WM_DELETE_WINDOW", on_building_close)
    
    label = tk.Label(
        building_window, 
        text="SELECT BUILDING", 
        font=("Helvetica", 16, "bold"),
        fg="white", 
        bg="black"
    )
    label.pack(pady=20)
    
    selected_building = tk.StringVar()
    
    def confirm_selection():
        if not selected_building.get():
            messagebox.showwarning("Selection Required", "Please select a building before proceeding.")
            return
        
        # Set the global building variable
        global building
        building = selected_building.get()
        
        # Close the window
        building_window.destroy()
        
        # Now configure tabs based on selection
        configure_tabs_for_building()
    
    # Memorial Union Option
    mem_frame = tk.Frame(building_window, bg="black", bd=2, relief=tk.RIDGE, padx=10, pady=10)
    mem_frame.pack(fill=tk.X, padx=20, pady=5)
    
    mem_radio = tk.Radiobutton(
        mem_frame,
        text="Memorial Union",
        variable=selected_building,
        value="Memorial Union",
        bg="black",
        fg="white",
        selectcolor="black",
        font=("Helvetica", 12),
        activebackground="black",
        activeforeground="white"
    )
    mem_radio.pack(anchor=tk.W)
    
    # Union South Option
    south_frame = tk.Frame(building_window, bg="black", bd=2, relief=tk.RIDGE, padx=10, pady=10)
    south_frame.pack(fill=tk.X, padx=20, pady=5)
    
    south_radio = tk.Radiobutton(
        south_frame,
        text="Union South",
        variable=selected_building,
        value="Union South",
        bg="black",
        fg="white",
        selectcolor="black",
        font=("Helvetica", 12),
        activebackground="black",
        activeforeground="white"
    )
    south_radio.pack(anchor=tk.W)
    
    # Red Gym Option
    red_gym_frame = tk.Frame(building_window, bg="black", bd=2, relief=tk.RIDGE, padx=10, pady=10)
    red_gym_frame.pack(fill=tk.X, padx=20, pady=5)
    
    red_gym_radio = tk.Radiobutton(
        red_gym_frame,
        text="Red Gym",
        variable=selected_building,
        value="Red Gym",
        bg="black",
        fg="white",
        selectcolor="black",
        font=("Helvetica", 12),
        activebackground="black",
        activeforeground="white"
    )
    red_gym_radio.pack(anchor=tk.W)
    
    # Confirm Button
    confirm_btn = tk.Button(
        building_window,
        text="Confirm Selection",
        command=confirm_selection,
        bg="white",
        fg="black",
        font=("Helvetica", 11, "bold"),
        padx=10, 
        pady=5
    )
    confirm_btn.pack(pady=15)
    
    # Wait for this window to be destroyed before proceeding
    root.wait_window(building_window)

def configure_text_box(textbox, min_height=4):
    """Configure a text box to automatically resize based on content"""
    def update_height(event=None):
        # Get text content
        text = textbox.get("1.0", "end-1c")  # Exclude the automatic trailing newline
        
        # Get the width of the text box in pixels
        width = textbox.winfo_width()
        
        # If width is not properly initialized yet, use approximate calculation
        if width <= 1:
            if hasattr(textbox, 'text_font'):
                width = textbox.cget("width") * textbox.text_font.measure("0")
            else:
                # Default width approximation
                width = textbox.cget("width") * 7
        else:
            # Subtract some pixels for padding
            width = width - 10
        
        # Calculate the number of wrapped lines
        line_count = 0
        for line in text.split('\n'):
            if line.strip() == "":  # Empty line
                line_count += 1
                continue
                
            # Estimate line wrapping based on average character width
            if hasattr(textbox, 'text_font'):
                char_width = textbox.text_font.measure("0")  # Average character width
            else:
                char_width = 7  # Default approximation
            
            # Calculate how many lines this text would wrap to
            if char_width > 0:
                chars_per_line = max(1, width // char_width)
                wrapped_lines = (len(line) // chars_per_line) + 1
                line_count += wrapped_lines
            else:
                line_count += 1
        
        # Add 1 for the cursor line and any pending text
        line_count += 1
        
        # Set minimum height
        new_height = max(min_height, line_count)
        
        # Update height
        textbox.configure(height=new_height)
    
    # Create a font object to measure text
    try:
        font_name = textbox.cget("font")
        if isinstance(font_name, str):
            textbox.text_font = font.Font(font=font_name)
    except Exception:
        # If there's any error with font, we'll use approximations instead
        textbox.text_font = font.Font(family="Helvetica", size=11)
    
    # Bind events that should trigger height update
    textbox.bind("<KeyRelease>", update_height)
    textbox.bind("<FocusOut>", update_height)
    textbox.bind("<Configure>", update_height)  # Also update when the widget is resized
    
    # Initial height update (after a short delay to ensure the widget is rendered)
    textbox.after(100, update_height)
    
    return textbox

root = tk.Tk()
root.title("Night Report Generator")
root.configure(bg="black")
root.state('zoomed')  # Open in full screen
# root.geometry("750x700")  # Remove fixed geometry
root.withdraw()  # Hide the main window until building is selected

# Global building variable
building = ""

# Global autosave variables
autosave_thread = None
autosave_running = False

# Global draft loading variable
loaded_from_draft = False

label_font = ("Helvetica", 11)
entry_bg = "black"
entry_fg = "white"
entry_font = ("Helvetica", 11)
entries = {}
building_traffic_boxes = []

# Add these global lists for all note sections
mechanical_boxes = []
mechanical_note_tags = []
production_boxes = []
production_note_tags = []
decibel_entries = []
patron_boxes = []
patron_note_tags = []
patron_emergency_flags = []
access_note_boxes = []
access_note_tags = []
cash_boxes = []
cash_note_tags = []
dining_boxes = []
dining_note_tags = []
hotel_boxes = []
misc_boxes = []
misc_note_tags = []

# Global variables for Security/CSC section
csc_entries = {}
csc_shifts = ["Morning", "Evening", "Special Event", "Chair Watch"]

# Global frame variables that need to be available to widget creation functions
traffic_notes_frame = None
mechanical_notes_frame = None
production_notes_frame = None
patron_notes_frame = None
cash_frame = None
dining_frame = None
hotel_frame = None
misc_frame = None
# Memorial Union specific frames
carding_frame = None
terrace_frame = None
enforcement_frame = None
alumni_frame = None
pier_frame = None
# Red Gym specific frames
mail_frame = None
# Access frame containers
access_notes_container = None
decibel_rows_container = None
enforcement_components = []  # Track enforcement components for proper ordering

# Add these global variables for Red Gym after the existing global lists
red_gym_building_tours_box = None
red_gym_deviations_entry = None
red_gym_deviation_boxes = []
red_gym_door_check_time = None
red_gym_door_check_day_type = None
red_gym_mail_boxes = []
red_gym_misc_boxes = []

# Memorial Union-specific
carding_boxes = []
terrace_boxes = []
enforcement_boxes = []
enforcement_note_tags = []
enforcement_images = []  # Add this for storing image paths
enforcement_items = []  # Special list for enforcement image dict items (Memorial Union only)
enforcement_components = []  # Track enforcement components for proper ordering
alumni_boxes = []
alumni_note_tags = []
pier_boxes = []
pier_note_tags = []

# Memorial Union specific function for enforcement component ordering
def reorder_enforcement_components():
    """Reorder enforcement components in the correct order"""
    # Forget all components
    for component in enforcement_components:
        component.pack_forget()
    # Re-pack in correct order
    for component in enforcement_components:
        component.pack(fill="x", pady=5)

# Tag options by tab (global)
MECHANICAL_TAG_OPTIONS = ["None", "FAMIS report", "Door Lock Failure", "Reset Elevator", "Custodial Support"]
PRODUCTION_TAG_OPTIONS = ["None", "Production Support", "AV Troubleshooting for Guest/Patron"]
PATRON_TAG_OPTIONS = ["None", "Lost & Found support", "First Aid Kit", "Incident report", "No dog policy", "Confirmed service animal", "Ban patron encounter", "Repeat Offender Encounter", "Severe Disruption / Verbal Escalation", "Patron Subject to 24-Hour Ban", "UWPD/911 Called", "Non-Patron Escort / Trespass", "Ensure outside alcohol policy", "Enforce smoking policy", "Removed Bikes & electric transport", "Wellness Check", "Disability Assistance/Patron Support", "CESO Support/Wedding walkthrough", "Patron Services/inquires/General Assistance"]
ACCESS_TAG_OPTIONS = ["None", "Loading dock access", "Employee Locker unlock", "Health Room unlock", "Door/cooler unlock"]
CASH_TAG_OPTIONS = ["None", "Cash Equipment Jam", "Cash Machine Problem"]
TERRACE_TAG_OPTIONS = ["None", "No dog policy", "Ensure outside alcohol policy & remove outside alcohol", "Enforce smoking policy", "Removed Bikes & electric transport", "Enforce no fishing policy", "Ban patron encounter", "Repeat Offender Encounter", "Confirmed service animal", "Wellness Check"]
DINING_TAG_OPTIONS = ["None", "Catering Order Pickup/inquiry"]
RED_GYM_MISC_TAG_OPTIONS = ["None", "Physical Plant"]

# Helper for adding tag dropdowns to a note (must be defined before use)
def add_tagging_to_note(tag_frame, tag_options, tag_vars, tag_dropdowns):
    def add_tag_dropdown():
        var = tk.StringVar(value="None")
        dropdown = ttk.Combobox(tag_frame, textvariable=var, values=tag_options, state="readonly", width=30)
        dropdown.pack(side="left", padx=(0, 5))
        tag_vars.append(var)
        tag_dropdowns.append(dropdown)
        def on_tag_change(event=None):
            if var.get() != "None" and not hasattr(dropdown, 'add_tag_btn'):
                add_btn = tk.Button(tag_frame, text="+ Add Tag", bg="white", fg="black", font=("Helvetica", 9, "bold"),
                                   command=lambda: [add_tag_dropdown(), add_btn.destroy()])
                add_btn.pack(side="left", padx=(0, 5))
                dropdown.add_tag_btn = add_btn
            elif var.get() == "None" and hasattr(dropdown, 'add_tag_btn'):
                dropdown.add_tag_btn.destroy()
                delattr(dropdown, 'add_tag_btn')
        dropdown.bind("<<ComboboxSelected>>", on_tag_change)
        return add_tag_dropdown  # Return the function for programmatic use
    
    def update_add_tag_buttons():
        """Helper function to show/hide add tag buttons based on current tag values"""
        # First, remove all existing add tag buttons
        for dropdown in tag_dropdowns:
            if hasattr(dropdown, 'add_tag_btn'):
                dropdown.add_tag_btn.destroy()
                delattr(dropdown, 'add_tag_btn')
        
        # Find the last dropdown with a non-"None" value
        last_valid_index = -1
        for i, var in enumerate(tag_vars):
            if var and var.get() != "None":
                last_valid_index = i
        
        # Only add the "Add Tag" button to the last dropdown with a valid tag
        if last_valid_index >= 0 and last_valid_index < len(tag_dropdowns):
            dropdown = tag_dropdowns[last_valid_index]
            add_btn = tk.Button(tag_frame, text="+ Add Tag", bg="white", fg="black", font=("Helvetica", 9, "bold"),
                               command=lambda: [add_tag_dropdown(), add_btn.destroy()])
            add_btn.pack(side="left", padx=(0, 5))
            dropdown.add_tag_btn = add_btn
    
    # Store the functions for later use
    tag_frame.add_tag_dropdown = add_tag_dropdown
    tag_frame.update_add_tag_buttons = update_add_tag_buttons
    add_tag_dropdown()

# === Tabs ===
notebook = ttk.Notebook(root)
notebook.pack(expand=1, fill="both")

style = ttk.Style()
style.theme_use('default')
style.configure('TNotebook', background='black', borderwidth=0)
style.configure('TNotebook.Tab', background='black', foreground='white', padding=10)
style.map('TNotebook.Tab', background=[('selected', '#333')], foreground=[('selected', 'white')])

def create_tab(title):
    frame = tk.Frame(notebook, bg="black")
    notebook.add(frame, text=title)
    return frame

def add_labeled_entry(parent, label_text, key, default=""):
    frame = tk.Frame(parent, bg="black")
    label = tk.Label(frame, text=label_text, fg="white", bg="black", font=label_font)
    entry = tk.Entry(frame, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, width=60)
    entry.insert(0, default)
    label.pack(anchor="w")
    entry.pack(fill="x")
    frame.pack(pady=5, padx=10, fill="x")
    entries[key] = entry

# === Utility Functions ===

def set_note_tags(note_tags_list, note_index, tags_to_set):
    """
    Utility function to set tag values for a specific note's tag comboboxes.
    
    Args:
        note_tags_list: The global list containing tag variables for all notes (e.g., mechanical_note_tags)
        note_index: Index of the note whose tags should be set (0-based)
        tags_to_set: List of tag values to set
    
    Example:
        set_note_tags(mechanical_note_tags, 0, ["FAMIS report", "Door Lock Failure"])
    """
    if note_index < len(note_tags_list) and tags_to_set:
        # Filter out empty/None tags
        valid_tags = [tag for tag in tags_to_set if tag and tag != "None"]
        note_tag_vars = note_tags_list[note_index]
        
        # Create additional tag dropdowns if needed
        while len(note_tag_vars) < len(valid_tags):
            # We need to find and use the add_tag_dropdown function for this note
            # This requires finding the tag frame associated with this note
            # We'll handle this in the restore function instead
            break
        
        for tag_index, tag_value in enumerate(valid_tags):
            if tag_index < len(note_tag_vars):
                note_tag_vars[tag_index].set(tag_value)

def ensure_note_boxes(note_boxes_list, note_tags_list, add_function, target_count):
    """
    Utility function to ensure a specific number of note boxes exist.
    Adds or removes boxes as needed to match the target count.
    
    Args:
        note_boxes_list: The global list containing note textboxes (e.g., mechanical_boxes)
        note_tags_list: The global list containing tag variables for notes (e.g., mechanical_note_tags)
        add_function: The function to call to add new note boxes (e.g., add_mechanical_box)
        target_count: Desired number of note boxes
    
    Example:
        ensure_note_boxes(mechanical_boxes, mechanical_note_tags, add_mechanical_box, 3)
    """
    current_count = len(note_boxes_list)
    
    if current_count < target_count:
        # Add more boxes
        for _ in range(target_count - current_count):
            add_function("")
    elif current_count > target_count:
        # Remove excess boxes
        while len(note_boxes_list) > target_count:
            # Get the widget to destroy - it might be the textbox or its parent frame
            widget_to_destroy = note_boxes_list[-1]
            parent = widget_to_destroy.master
            
            # Try to destroy the parent frame if it exists and looks like a note frame
            try:
                if hasattr(parent, 'winfo_children') and len(parent.winfo_children()) <= 3:
                    parent.destroy()
                else:
                    widget_to_destroy.destroy()
            except:
                # Fallback to just destroying the widget
                try:
                    widget_to_destroy.destroy()
                except:
                    pass
            
            note_boxes_list.pop()
        
        # Also clean up tag lists to match
        while len(note_tags_list) > target_count:
            note_tags_list.pop()

# === Promoted Widget Creation Helper Functions ===

def add_building_traffic_box(default_text=""):
    """Add a building traffic note box (applies to all buildings)"""
    frame = tk.Frame(traffic_notes_frame, bg="black")
    label = tk.Label(frame, text=f"Building Traffic Note #{len(building_traffic_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    # Configure text box to auto-resize
    configure_text_box(textbox)
    building_traffic_boxes.append(textbox)

def add_mechanical_box(default_text=""):
    """Add a mechanical note box with dynamic tagging (Memorial Union & Union South)"""
    frame = tk.Frame(mechanical_notes_frame, bg="black")
    label = tk.Label(frame, text=f"Mechanical Note #{len(mechanical_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    configure_text_box(textbox)
    mechanical_boxes.append(textbox)

    # Tagging logic
    tag_vars = []  # List of tk.StringVar for this note
    tag_dropdowns = []  # List of dropdown widgets for this note

    tag_frame = tk.Frame(frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))

    def add_tag_dropdown():
        var = tk.StringVar(value="None")
        dropdown = ttk.Combobox(tag_frame, textvariable=var, values=MECHANICAL_TAG_OPTIONS, state="readonly", width=22)
        dropdown.pack(side="left", padx=(0, 5))
        tag_vars.append(var)
        tag_dropdowns.append(dropdown)

        def on_tag_change(event=None):
            # Show +Add Tag button if a valid tag is selected and no button exists
            if var.get() != "None" and not hasattr(dropdown, 'add_tag_btn'):
                add_btn = tk.Button(tag_frame, text="+ Add Tag", bg="white", fg="black", font=("Helvetica", 9, "bold"),
                                   command=lambda: [add_tag_dropdown(), add_btn.destroy()])
                add_btn.pack(side="left", padx=(0, 5))
                dropdown.add_tag_btn = add_btn
            elif var.get() == "None" and hasattr(dropdown, 'add_tag_btn'):
                dropdown.add_tag_btn.destroy()
                delattr(dropdown, 'add_tag_btn')
        dropdown.bind("<<ComboboxSelected>>", on_tag_change)
        return add_tag_dropdown  # Return function for programmatic access
    
    def update_add_tag_buttons():
        """Helper function to show/hide add tag buttons based on current tag values"""
        # First, remove all existing add tag buttons
        for dropdown in tag_dropdowns:
            if hasattr(dropdown, 'add_tag_btn'):
                dropdown.add_tag_btn.destroy()
                delattr(dropdown, 'add_tag_btn')
        
        # Find the last dropdown with a non-"None" value
        last_valid_index = -1
        for i, var in enumerate(tag_vars):
            if var and var.get() != "None":
                last_valid_index = i
        
        # Only add the "Add Tag" button to the last dropdown with a valid tag
        if last_valid_index >= 0 and last_valid_index < len(tag_dropdowns):
            dropdown = tag_dropdowns[last_valid_index]
            add_btn = tk.Button(tag_frame, text="+ Add Tag", bg="white", fg="black", font=("Helvetica", 9, "bold"),
                               command=lambda: [add_tag_dropdown(), add_btn.destroy()])
            add_btn.pack(side="left", padx=(0, 5))
            dropdown.add_tag_btn = add_btn
    
    # Store the functions for later use during restoration
    tag_frame.add_tag_dropdown = add_tag_dropdown
    tag_frame.update_add_tag_buttons = update_add_tag_buttons
    add_tag_dropdown()  # Add the first dropdown
    mechanical_note_tags.append(tag_vars)
    frame.pack(pady=5, fill="x")

def add_production_note_box(default_text=""):
    """Add a production note box with tagging (Memorial Union & Union South)"""
    frame = tk.Frame(production_notes_frame, bg="black")
    label = tk.Label(frame, text=f"Production Note #{len(production_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    production_boxes.append(textbox)
    # Tagging
    tag_vars = []
    tag_dropdowns = []
    tag_frame = tk.Frame(frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))
    add_tagging_to_note(tag_frame, PRODUCTION_TAG_OPTIONS, tag_vars, tag_dropdowns)
    production_note_tags.append(tag_vars)
    frame.pack(pady=5, fill="x")

def add_decibel_row():
    """Add a decibel reading row (Memorial Union & Union South)"""
    row_frame = tk.Frame(decibel_rows_container, bg="black")
    time_entry = tk.Entry(row_frame, width=15, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
    reading_entry = tk.Entry(row_frame, width=15, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
    location_entry = tk.Entry(row_frame, width=40, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)

    time_entry.insert(0, "Time")
    reading_entry.insert(0, "Reading (db)")
    location_entry.insert(0, "Location")

    # Add focus event handlers to select all text when clicked
    def on_focus_in(event):
        event.widget.select_range(0, 'end')
        event.widget.icursor('end')
    
    time_entry.bind("<FocusIn>", on_focus_in)
    reading_entry.bind("<FocusIn>", on_focus_in)
    location_entry.bind("<FocusIn>", on_focus_in)

    time_entry.pack(side="left", padx=5)
    reading_entry.pack(side="left", padx=5)
    location_entry.pack(side="left", padx=5)
    row_frame.pack(pady=3, anchor="w", fill="x")
    
    decibel_entries.append((time_entry, reading_entry, location_entry))

def add_patron_note_box(default_text=""):
    """Add a patron services note box with tagging (Memorial Union & Union South)"""
    frame = tk.Frame(patron_notes_frame, bg="black")
    label = tk.Label(frame, text=f"Patron Note #{len(patron_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=6, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    configure_text_box(textbox, min_height=6)
    patron_boxes.append(textbox)
    # Tagging
    tag_vars = []
    tag_dropdowns = []
    tag_frame = tk.Frame(frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))
    add_tagging_to_note(tag_frame, PATRON_TAG_OPTIONS, tag_vars, tag_dropdowns)
    patron_note_tags.append(tag_vars)
    frame.pack(pady=5, fill="x")

def add_access_note():
    """Add an access note box with tagging (Memorial Union & Union South)"""
    frame = tk.Frame(access_notes_container, bg="black")
    label = tk.Label(frame, text=f"Access Note #{len(access_note_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    label.pack(anchor="w")
    textbox.pack(fill="x", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    access_note_boxes.append(textbox)
    # Tagging
    tag_vars = []
    tag_dropdowns = []
    tag_frame = tk.Frame(frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))
    add_tagging_to_note(tag_frame, ACCESS_TAG_OPTIONS, tag_vars, tag_dropdowns)
    access_note_tags.append(tag_vars)
    frame.pack(pady=5, fill="x")

def add_cash_note_box(default_text=""):
    """Add a cash office note box with tagging (Memorial Union & Union South)"""
    frame = tk.Frame(cash_frame, bg="black")
    label = tk.Label(frame, text=f"Cash Office Note #{len(cash_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    cash_boxes.append(textbox)
    # Tagging
    tag_vars = []
    tag_dropdowns = []
    tag_frame = tk.Frame(frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))
    add_tagging_to_note(tag_frame, CASH_TAG_OPTIONS, tag_vars, tag_dropdowns)
    cash_note_tags.append(tag_vars)
    frame.pack(pady=5, fill="x")

def add_dining_note_box(default_text=""):
    """Add a dining note box with tagging (Memorial Union & Union South)"""
    frame = tk.Frame(dining_frame, bg="black")
    label = tk.Label(frame, text=f"Dining Note #{len(dining_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    dining_boxes.append(textbox)
    # Tagging
    tag_vars = []
    tag_dropdowns = []
    tag_frame = tk.Frame(frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))
    add_tagging_to_note(tag_frame, DINING_TAG_OPTIONS, tag_vars, tag_dropdowns)
    dining_note_tags.append(tag_vars)
    frame.pack(pady=5, fill="x")

def add_hotel_note_box(default_text=""):
    """Add a hotel note box (Memorial Union & Union South)"""
    frame = tk.Frame(hotel_frame, bg="black")
    label = tk.Label(frame, text=f"Hotel Note #{len(hotel_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    configure_text_box(textbox, min_height=3)
    hotel_boxes.append(textbox)
    frame.pack(pady=5, fill="x")

def add_misc_note_box(default_text=""):
    """Add a miscellaneous note box (Memorial Union & Union South)"""
    frame = tk.Frame(misc_frame, bg="black")
    label = tk.Label(frame, text=f"Misc Note #{len(misc_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    misc_boxes.append(textbox)
    frame.pack(pady=5, fill="x")

# === Memorial Union Specific Functions ===

def add_carding_note_box(default_text=""):
    """Add a carding note box (Memorial Union only)"""
    frame = tk.Frame(carding_frame, bg="black")
    label = tk.Label(frame, text=f"Carding Run Note #{len(carding_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    carding_boxes.append(textbox)

def add_terrace_note_box(default_text=""):
    """Add a terrace traffic note box (Memorial Union only)"""
    frame = tk.Frame(terrace_frame, bg="black")
    label = tk.Label(frame, text=f"Terrace Traffic Note #{len(terrace_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    terrace_boxes.append(textbox)

def add_enforcement_note_box(default_text=""):
    """Add an enforcement text note box with tagging (Memorial Union only)"""
    frame = tk.Frame(enforcement_frame, bg="black")
    label = tk.Label(frame, text=f"Enforcement Note #{len(enforcement_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    configure_text_box(textbox)
    enforcement_boxes.append(textbox)
    # Tagging
    tag_vars = []
    tag_dropdowns = []
    tag_frame = tk.Frame(frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))
    add_tagging_to_note(tag_frame, TERRACE_TAG_OPTIONS, tag_vars, tag_dropdowns)
    enforcement_note_tags.append(tag_vars)
    # Add to components and reorder
    enforcement_components.append(frame)
    reorder_enforcement_components()

def add_enforcement_image():
    """Add an enforcement image with description and tagging (Memorial Union only).
    Returns dict with image info and pushes to enforcement_items global list."""
    image_frame = tk.Frame(enforcement_frame, bg="black")
    
    # Label for the image section
    image_label = tk.Label(image_frame, text=f"Enforcement Image #{len(enforcement_images)+1}:", fg="white", bg="black", font=label_font)
    image_label.pack(anchor="w")
    
    # Frame for image upload button and status
    upload_frame = tk.Frame(image_frame, bg="black")
    upload_frame.pack(fill="x", pady=2)
    
    # Variable to store image path
    image_path_var = tk.StringVar()
    enforcement_images.insert(0, image_path_var)  # Insert at beginning
    
    # Status label
    status_label = tk.Label(upload_frame, text="No image selected", fg="gray", bg="black", font=entry_font)
    status_label.pack(side="left", padx=(0, 10))
    
    def select_image():
        file_path = filedialog.askopenfilename(
            title="Select Enforcement Image",
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.gif *.bmp"), ("All files", "*.*")]
        )
        if file_path:
            image_path_var.set(file_path)
            # RULE: Missing image files - still show filename; allow reselection
            import os
            if os.path.exists(file_path):
                status_label.config(text=f"Selected: {os.path.basename(file_path)}", fg="green")
            else:
                status_label.config(text=f"Missing: {os.path.basename(file_path)} (click to reselect)", fg="orange")
    
    # Update status when image path is set programmatically (for draft loading)
    def update_status_from_var(*args):
        current_path = image_path_var.get()
        if current_path:
            import os
            if os.path.exists(current_path):
                status_label.config(text=f"Selected: {os.path.basename(current_path)}", fg="green")
            else:
                # RULE: Missing image files - still show filename; allow reselection
                status_label.config(text=f"Missing: {os.path.basename(current_path)} (click to reselect)", fg="orange")
        else:
            status_label.config(text="No image selected", fg="gray")
    
    # Trace the variable to update status when set programmatically
    image_path_var.trace_add("write", update_status_from_var)
    
    upload_btn = tk.Button(upload_frame, text="Select Image", command=select_image,
                         bg="white", fg="black", font=("Helvetica", 9, "bold"))
    upload_btn.pack(side="left")
    
    # Text box for image description
    desc_label = tk.Label(image_frame, text="Description of enforcement action:", fg="white", bg="black", font=label_font)
    desc_label.pack(anchor="w", pady=(5, 0))
    
    textbox = tk.Text(image_frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    textbox.pack(fill="x", padx=5)
    configure_text_box(textbox, min_height=3)
    enforcement_boxes.insert(0, textbox)  # Insert at beginning
    
    # Tagging
    tag_vars = []
    tag_dropdowns = []
    tag_frame = tk.Frame(image_frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))
    add_tagging_to_note(tag_frame, TERRACE_TAG_OPTIONS, tag_vars, tag_dropdowns)
    enforcement_note_tags.insert(0, tag_vars)  # Insert at beginning
    
    # Pack the image frame and reorder all components
    enforcement_components.insert(0, image_frame)
    reorder_enforcement_components()
    
    # Create and store the enforcement item dict
    enforcement_item = {
        "type": "image",
        "image_path": image_path_var,
        "description_textbox": textbox,
        "tags": tag_vars,
        "frame": image_frame
    }
    enforcement_items.append(enforcement_item)
    
    return enforcement_item

def add_alumni_note_box(default_text=""):
    """Add an alumni park note box with tagging (Memorial Union only)"""
    frame = tk.Frame(alumni_frame, bg="black")
    label = tk.Label(frame, text=f"Alumni Park Note #{len(alumni_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    alumni_boxes.append(textbox)
    # Tagging
    tag_vars = []
    tag_dropdowns = []
    tag_frame = tk.Frame(frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))
    add_tagging_to_note(tag_frame, TERRACE_TAG_OPTIONS, tag_vars, tag_dropdowns)
    alumni_note_tags.append(tag_vars)
    frame.pack(pady=5, fill="x")

def add_pier_note_box(default_text=""):
    """Add a pier note box with tagging (Memorial Union only)"""
    frame = tk.Frame(pier_frame, bg="black")
    label = tk.Label(frame, text=f"Pier Note #{len(pier_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    pier_boxes.append(textbox)
    # Tagging
    tag_vars = []
    tag_dropdowns = []
    tag_frame = tk.Frame(frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))
    add_tagging_to_note(tag_frame, TERRACE_TAG_OPTIONS, tag_vars, tag_dropdowns)
    pier_note_tags.append(tag_vars)
    frame.pack(pady=5, fill="x")

# === Red Gym Specific Functions ===

def add_red_gym_mail_box(default_text=""):
    """Add a mail note box (Red Gym only)"""
    frame = tk.Frame(mail_frame, bg="black")
    label = tk.Label(frame, text=f"Mail Note #{len(red_gym_mail_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    red_gym_mail_boxes.append(textbox)

def add_red_gym_misc_box(default_text=""):
    """Add a misc note box with tagging (Red Gym only)"""
    frame = tk.Frame(misc_frame, bg="black")
    label = tk.Label(frame, text=f"Misc Note #{len(red_gym_misc_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    red_gym_misc_boxes.append(textbox)
    
    # Tagging for Red Gym misc (with ability to add multiple tags)
    tag_vars = []
    tag_dropdowns = []
    tag_frame = tk.Frame(frame, bg="black")
    tag_frame.pack(anchor="w", pady=(2, 0))
    
    # Use the standard tagging function for consistency
    add_tagging_to_note(tag_frame, RED_GYM_MISC_TAG_OPTIONS, tag_vars, tag_dropdowns)
    misc_note_tags.append(tag_vars)
    frame.pack(pady=5, fill="x")

# Function to configure tabs based on selected building
def configure_tabs_for_building():
    global tabs
    
    # Update window title with selected building
    root.title(f"{building} - Night Report Generator")
    
    # Define all possible tabs
    all_tab_keys = [
        "Supervisor Info", "Building Traffic", "Mechanical", "Production", 
        "Patron Services", "Access", "Cash Office", "Carding Runs", 
        "Terrace Traffic", "Terrace Enforcement", "Alumni Park", "Goodspeed Pier",
        "Dining & Markets", "Hotel", "Misc", "Security"
    ]
    
    # Define the tabs to exclude for different buildings
    exclude_for_union_south = [
        "Carding Runs", "Terrace Traffic", "Terrace Enforcement", 
        "Alumni Park", "Goodspeed Pier", "Mail"
    ]
    
    exclude_for_memorial_union = [
        "Mail"
    ]
    
    red_gym_tabs_only = [
        "Supervisor Info", "Building Traffic", "Security", "Mail", "Misc"
    ]
    
    # Determine which tabs to create
    if building == "Red Gym":
        tab_keys_to_create = red_gym_tabs_only
    elif building == "Union South":
        tab_keys_to_create = [key for key in all_tab_keys if key not in exclude_for_union_south]
    else:  # Memorial Union
        tab_keys_to_create = [key for key in all_tab_keys if key not in exclude_for_memorial_union]
    
    # Create tabs
    tabs = {}
    for key in tab_keys_to_create:
        tab_title = key
        if key == "Access":
            tab_title = "Access/Lock/Unlock"
        elif key == "Dining & Markets":
            tab_title = "Dining Service & Markets"
        elif key == "Goodspeed Pier":
            tab_title = "Goodspeed Family Pier"
            
        tabs[key] = create_tab(tab_title)
    
    # Continue with the rest of the UI setup
    setup_ui_components()
    
    # Start autosave after UI is configured
    start_autosave(interval_min=3)
    
    # Handle main window close event - exit application properly
    def on_main_close():
        # Show confirmation dialog
        response = messagebox.askyesno(
            "Confirm Exit", 
            "Are you sure you want to close the Night Report Generator?\n\nAny unsaved changes will be lost.",
            icon='warning'
        )
        if response:  # User clicked "Yes"
            stop_autosave()  # Stop any running autosave
            root.quit()  # Exit the mainloop
            root.destroy()  # Destroy the root window
    
    root.protocol("WM_DELETE_WINDOW", on_main_close)
    
    # Now show the main window
    root.deiconify()

def setup_ui_components():
    global traffic_notes_frame, mechanical_notes_frame, production_notes_frame, patron_notes_frame
    global cash_frame, dining_frame, hotel_frame, misc_frame
    global carding_frame, terrace_frame, enforcement_frame, alumni_frame, pier_frame
    global mail_frame, access_notes_container, decibel_rows_container
    
    # === Supervisor Info tab ===
    today_str = datetime.now().strftime("%A, %B %d, %Y")
    add_labeled_entry(tabs["Supervisor Info"], "Date", "date", default=today_str)
    add_labeled_entry(tabs["Supervisor Info"], "Shift Hours", "shift_hours")
    add_labeled_entry(tabs["Supervisor Info"], "Building Manager(s)", "bms")
    
    # Only add additional fields for Memorial Union and Union South
    if building != "Red Gym":
        # Terrace Manager(s) only for Memorial Union
        if building == "Memorial Union":
            add_labeled_entry(tabs["Supervisor Info"], "Terrace Manager(s)", "terrace_managers")
        add_labeled_entry(tabs["Supervisor Info"], "Guest Service Specialist", "gss")
        add_labeled_entry(tabs["Supervisor Info"], "Operation Manager(s)", "operation_managers")
        add_labeled_entry(tabs["Supervisor Info"], "Custodial Supervisor(s)", "custodial")
        add_labeled_entry(tabs["Supervisor Info"], "Production Supervisor(s)", "production")
        add_labeled_entry(tabs["Supervisor Info"], "Retail & Dining Supervisor(s)", "retail")
        add_labeled_entry(tabs["Supervisor Info"], "Catering Supervisor(s)", "catering")
        add_labeled_entry(tabs["Supervisor Info"], "Event Manager(s)", "eventmanagers")
        add_labeled_entry(tabs["Supervisor Info"], "CAVR Desk Staff", "cavr")
    
    # === Building Traffic Tab ===
    # Create a container frame inside the tab to hold text boxes
    traffic_tab = tabs["Building Traffic"]
    traffic_notes_frame = tk.Frame(traffic_tab, bg="black")
    traffic_notes_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    # Add first required box (using promoted function)
    add_building_traffic_box()

    # Add the + Add Note button, stays at bottom
    add_box_btn = tk.Button(
        traffic_tab, text="+ Add Note", command=add_building_traffic_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_box_btn.pack(pady=10)

    # Red Gym specific tabs
    if building == "Red Gym":
        # === Red Gym Security Tab ===
        global red_gym_building_tours_box, red_gym_deviations_entry, red_gym_deviation_boxes
        global red_gym_door_check_time, red_gym_door_check_day_type
        
        security_tab = tabs["Security"]
        security_frame = tk.Frame(security_tab, bg="black")
        security_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Building Tours
        tours_label = tk.Label(security_frame, text="Building Tours:", fg="white", bg="black", font=label_font)
        tours_label.pack(anchor="w")
        red_gym_building_tours_box = tk.Text(security_frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        red_gym_building_tours_box.pack(fill="x", padx=5, pady=(0, 10))
        configure_text_box(red_gym_building_tours_box, min_height=3)
        
        # Deviations section
        deviations_label = tk.Label(security_frame, text="Deviations from standard building locking protocol:", fg="white", bg="black", font=label_font)
        deviations_label.pack(anchor="w", pady=(10, 0))
        
        deviations_frame = tk.Frame(security_frame, bg="black")
        deviations_frame.pack(fill="x", pady=5)
        
        deviations_text_label = tk.Label(deviations_frame, text="There were", fg="white", bg="black", font=label_font)
        deviations_text_label.pack(side="left")
        
        red_gym_deviations_entry = tk.Entry(deviations_frame, width=5, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        red_gym_deviations_entry.insert(0, "0")
        red_gym_deviations_entry.pack(side="left", padx=(5, 5))
        
        deviations_text_label2 = tk.Label(deviations_frame, text="deviations from the standard building locking protocol today.", fg="white", bg="black", font=label_font)
        deviations_text_label2.pack(side="left")
        
        # Container for deviation notes
        deviations_notes_frame = tk.Frame(security_frame, bg="black")
        deviations_notes_frame.pack(fill="x", pady=5)
        
        def update_deviation_notes():
            # Clear existing notes
            for widget in deviations_notes_frame.winfo_children():
                widget.destroy()
            red_gym_deviation_boxes.clear()
            
            try:
                num_deviations = int(red_gym_deviations_entry.get() or "0")
                if num_deviations > 0:
                    for i in range(num_deviations):
                        frame = tk.Frame(deviations_notes_frame, bg="black")
                        label = tk.Label(frame, text=f"Deviation {chr(97+i)}:", fg="white", bg="black", font=label_font)
                        textbox = tk.Text(frame, height=2, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
                        label.pack(anchor="w")
                        textbox.pack(fill="x", padx=5)
                        frame.pack(pady=2, fill="x")
                        configure_text_box(textbox, min_height=2)
                        red_gym_deviation_boxes.append(textbox)
            except ValueError:
                pass
        
        red_gym_deviations_entry.bind('<KeyRelease>', lambda e: update_deviation_notes())
        
        # Door check section
        door_check_frame = tk.Frame(security_frame, bg="black")
        door_check_frame.pack(fill="x", pady=(20, 5))
        
        door_check_label1 = tk.Label(door_check_frame, text="I checked to confirm that the Building Doors were locked and the door swipe scanner was red at", fg="white", bg="black", font=label_font)
        door_check_label1.pack(side="left")
        
        red_gym_door_check_time = tk.Entry(door_check_frame, width=10, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        red_gym_door_check_time.pack(side="left", padx=(5, 5))
        
        door_check_label2 = tk.Label(door_check_frame, text="on a", fg="white", bg="black", font=label_font)
        door_check_label2.pack(side="left")
        
        red_gym_door_check_day_type = ttk.Combobox(door_check_frame, values=["weekday", "weekend"], state="readonly", width=10)
        red_gym_door_check_day_type.pack(side="left", padx=(5, 0))
        
        door_check_label3 = tk.Label(door_check_frame, text=".", fg="white", bg="black", font=label_font)
        door_check_label3.pack(side="left")
        
        # === Red Gym Mail Tab ===
        global red_gym_mail_boxes
        mail_tab = tabs["Mail"]
        mail_frame = tk.Frame(mail_tab, bg="black")
        mail_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))
        
        # Use promoted function
        add_red_gym_mail_box()
        
        tk.Button(
            mail_tab, text="+ Add Note", command=add_red_gym_mail_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        ).pack(pady=10)
        
        # === Red Gym Misc Tab ===
        global red_gym_misc_boxes
        misc_tab = tabs["Misc"]
        misc_frame = tk.Frame(misc_tab, bg="black")
        misc_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))
        
        # Use promoted function
        add_red_gym_misc_box()
        
        tk.Button(
            misc_tab, text="+ Add Note", command=add_red_gym_misc_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        ).pack(pady=10)
        
        # === Generate Report Buttons ===
        button_frame = tk.Frame(root, bg="black")
        button_frame.pack(pady=10)
        
        save_btn = tk.Button(
            button_frame, text="Save Report", command=save_report_draft,
            bg="white", fg="black", font=("Helvetica", 12, "bold"), padx=10, pady=6,
            relief="raised", activebackground="white", activeforeground="black"
        )
        save_btn.pack(side="left", padx=8)
        
        submit_btn = tk.Button(
            button_frame, text="End Shift", command=end_shift_and_generate,
            bg="white", fg="black", font=("Helvetica", 12, "bold"), padx=10, pady=6,
            relief="raised", activebackground="white", activeforeground="black"
        )
        submit_btn.pack(side="left")
        
        return  # Exit early for Red Gym
    
    # Continue with existing code for Memorial Union and Union South
    # === Mechanical/Repairs/Custodial Tab ===
    mechanical_notes_frame = tk.Frame(tabs["Mechanical"], bg="black")
    mechanical_notes_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    def add_mechanical_box(default_text=""):
        frame = tk.Frame(mechanical_notes_frame, bg="black")
        label = tk.Label(frame, text=f"Mechanical Note #{len(mechanical_boxes)+1}:", fg="white", bg="black", font=label_font)
        textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        if default_text:
            textbox.insert("1.0", default_text)
        label.pack(anchor="w")
        textbox.pack(fill="both", expand=True, padx=5)
        configure_text_box(textbox)
        mechanical_boxes.append(textbox)

        # Tagging logic
        tag_vars = []  # List of tk.StringVar for this note
        tag_dropdowns = []  # List of dropdown widgets for this note

        tag_frame = tk.Frame(frame, bg="black")
        tag_frame.pack(anchor="w", pady=(2, 0))

        def add_tag_dropdown():
            var = tk.StringVar(value="None")
            dropdown = ttk.Combobox(tag_frame, textvariable=var, values=MECHANICAL_TAG_OPTIONS, state="readonly", width=22)
            dropdown.pack(side="left", padx=(0, 5))
            tag_vars.append(var)
            tag_dropdowns.append(dropdown)

            def on_tag_change(event=None):
                # Show +Add Tag button if a valid tag is selected and no button exists
                if var.get() != "None" and not hasattr(dropdown, 'add_tag_btn'):
                    add_btn = tk.Button(tag_frame, text="+ Add Tag", bg="white", fg="black", font=("Helvetica", 9, "bold"),
                                       command=lambda: [add_tag_dropdown(), add_btn.destroy()])
                    add_btn.pack(side="left", padx=(0, 5))
                    dropdown.add_tag_btn = add_btn
                elif var.get() == "None" and hasattr(dropdown, 'add_tag_btn'):
                    dropdown.add_tag_btn.destroy()
                    delattr(dropdown, 'add_tag_btn')
            dropdown.bind("<<ComboboxSelected>>", on_tag_change)
            return add_tag_dropdown  # Return function for programmatic access
        
        # Store the add_tag_dropdown function for later use during restoration
        tag_frame.add_tag_dropdown = add_tag_dropdown
        add_tag_dropdown()  # Add the first dropdown
        mechanical_note_tags.append(tag_vars)
        frame.pack(pady=5, fill="x")

    # Add first required box
    add_mechanical_box()

    # Add Note button
    add_mechanical_btn = tk.Button(
        tabs["Mechanical"], text="+ Add Note", command=add_mechanical_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_mechanical_btn.pack(pady=10)

    # === Production Services Tab ===
    production_notes_frame = tk.Frame(tabs["Production"], bg="black")
    production_notes_frame.pack(fill="x", padx=10, pady=(10, 5))

    # Use promoted function
    add_production_note_box()

    # Button to add more production notes
    add_note_btn = tk.Button(
        tabs["Production"], text="+ Add Note", command=add_production_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_note_btn.pack(pady=(0, 10))

    # === Decibel Reading Table ===
    decibel_frame = tk.Frame(tabs["Production"], bg="black")
    decibel_frame.pack(padx=10, pady=(10, 5), fill="x")

    header = tk.Label(decibel_frame, text="Decibel Readings", fg="white", bg="black", font=("Helvetica", 12, "bold"))
    header.pack(anchor="w")

    # Container for decibel rows - this will help with proper placement
    decibel_rows_container = tk.Frame(decibel_frame, bg="black")
    decibel_rows_container.pack(fill="x", expand=True)

    # Use promoted function
    add_decibel_row()

    # Button to add more decibel rows - kept outside the container
    add_decibel_btn = tk.Button(
        decibel_frame, text="+ Add Decibel Reading", command=add_decibel_row,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_decibel_btn.pack(pady=10)

    # === Patron Services Tab ===
    patron_notes_frame = tk.Frame(tabs["Patron Services"], bg="black")
    patron_notes_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    # Use promoted function
    add_patron_note_box()

    # Add button to add more notes
    add_patron_btn = tk.Button(
        tabs["Patron Services"], text="+ Add Note", command=add_patron_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_patron_btn.pack(pady=10)

    # === Access/Lock/Unlock Tab ===
    global access_inputs, csc_shifts, csc_entries  # <-- This line ensures both are global
    access_tab = tabs["Access"]
    access_frame = tk.Frame(access_tab, bg="black")
    access_frame.pack(fill="both", expand=True, padx=10, pady=10)
    
    access_inputs = {}
    access_note_boxes = []

    def add_dropdown(label_text, key, options):
        frame = tk.Frame(access_frame, bg="black")
        label = tk.Label(frame, text=label_text, fg="white", bg="black", font=label_font)
        var = tk.StringVar(value=options[0])
        dropdown = ttk.Combobox(frame, textvariable=var, values=options, state="readonly", width=20)
        label.pack(anchor="w")
        dropdown.pack(fill="x")
        frame.pack(pady=5, fill="x")
        access_inputs[key] = var

    def add_entry(label_text, key):
        frame = tk.Frame(access_frame, bg="black")
        label = tk.Label(frame, text=label_text, fg="white", bg="black", font=label_font)
        entry = tk.Entry(frame, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, width=30)
        label.pack(anchor="w")
        entry.pack(fill="x")
        frame.pack(pady=5, fill="x")
        access_inputs[key] = entry

    add_dropdown("Loading Dock Arm Gate at Early Check:", "early_gate", ["Open", "Closed"])
    add_entry("Time of Early Check:", "early_time")

    add_dropdown("Loading Dock Arm Gate at Closing Check:", "close_gate", ["Open", "Closed"])
    add_entry("Time of Closing Check:", "close_time")

    add_dropdown("HID Scanners Status at Close:", "hid_status", ["Locked", "Unlocked"])
    
    # Set default based on building type
    if building == "Memorial Union":
        add_dropdown("Overhead Door Secured:", "door_status", ["Unsuccessfully", "Successfully"])
    else:  # Union South and other buildings
        add_dropdown("Overhead Door Secured:", "door_status", ["Successfully", "Unsuccessfully"])

    # === Optional Access Notes ===
    access_notes_label = tk.Label(access_frame, text="Additional Access Notes (Optional):", fg="white", bg="black", font=label_font)
    access_notes_label.pack(anchor="w", pady=(10, 0))

    access_notes_container = tk.Frame(access_frame, bg="black")
    access_notes_container.pack(fill="both", expand=True)

    # Add Note button (using promoted function)
    add_note_btn = tk.Button(
        access_frame, text="+ Add Note", command=add_access_note,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_note_btn.pack(pady=10)

    # === Cash Office Tab ===
    cash_frame = tk.Frame(tabs["Cash Office"], bg="black")
    cash_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    # Use promoted function
    add_cash_note_box()

    add_cash_btn = tk.Button(
        tabs["Cash Office"], text="+ Add Note", command=add_cash_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_cash_btn.pack(pady=10)

    # Create the Memorial Union-specific tabs only if the selected building is Memorial Union
    if building == "Memorial Union":
        # === Carding Runs Tab ===
        carding_frame = tk.Frame(tabs["Carding Runs"], bg="black")
        carding_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        # Use promoted function
        add_carding_note_box()

        # Add note button
        add_carding_btn = tk.Button(
            tabs["Carding Runs"], text="+ Add Note", command=add_carding_note_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        )
        add_carding_btn.pack(pady=10)

        # === Terrace Traffic Tab ===
        terrace_frame = tk.Frame(tabs["Terrace Traffic"], bg="black")
        terrace_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        # Use promoted function
        add_terrace_note_box()

        add_terrace_btn = tk.Button(
            tabs["Terrace Traffic"], text="+ Add Note", command=add_terrace_note_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        )
        add_terrace_btn.pack(pady=10)

        # === Terrace Enforcement Tab ===
        enforcement_frame = tk.Frame(tabs["Terrace Enforcement"], bg="black")
        enforcement_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        # Use promoted functions
        # Button to add enforcement image with description
        add_image_btn = tk.Button(
            enforcement_frame, text="+ Add Enforcement Image", command=add_enforcement_image,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        )
        add_image_btn.pack(pady=5)

        add_enforcement_note_box()

        add_enforcement_btn = tk.Button(
            tabs["Terrace Enforcement"], text="+ Add Note", command=add_enforcement_note_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        )
        add_enforcement_btn.pack(pady=10)

        # === Alumni Park Tab ===
        alumni_frame = tk.Frame(tabs["Alumni Park"], bg="black")
        alumni_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        # Use promoted function
        add_alumni_note_box()

        tk.Button(
            tabs["Alumni Park"], text="+ Add Note", command=add_alumni_note_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        ).pack(pady=10)

        # === Goodspeed Pier Tab ===
        pier_frame = tk.Frame(tabs["Goodspeed Pier"], bg="black")
        pier_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        # Use promoted function
        add_pier_note_box()

        tk.Button(
            tabs["Goodspeed Pier"], text="+ Add Note", command=add_pier_note_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        ).pack(pady=10)

    # === Dining & Markets Tab ===
    dining_frame = tk.Frame(tabs["Dining & Markets"], bg="black")
    dining_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    # Use promoted function
    add_dining_note_box()

    tk.Button(
        tabs["Dining & Markets"], text="+ Add Note", command=add_dining_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    ).pack(pady=10)

    # === Hotel Tab ===
    hotel_frame = tk.Frame(tabs["Hotel"], bg="black")
    hotel_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    # Use promoted function
    add_hotel_note_box()
    tk.Button(
        tabs["Hotel"], text="+ Add Note", command=add_hotel_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    ).pack(pady=10)

    # === Miscellaneous Tab ===
    misc_frame = tk.Frame(tabs["Misc"], bg="black")
    misc_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    # Use promoted function
    add_misc_note_box()

    tk.Button(
        tabs["Misc"], text="+ Add Note", command=add_misc_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    ).pack(pady=10)

    # === Security Tab === (Changed from CSC Log)
    csc_tab = tabs["Security"]  # Changed from "CSC Log"
    csc_frame = tk.Frame(csc_tab, bg="black")
    csc_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    # Remove the local csc_entries and csc_shifts declarations since they're now global
    for shift in csc_shifts:
        section = tk.LabelFrame(csc_frame, text=shift, fg="white", bg="black", font=label_font, labelanchor="n", padx=10, pady=10)
        section.pack(fill="x", pady=5)

        # Requested
        req_label = tk.Label(section, text="Staff Requested:", fg="white", bg="black", font=label_font)
        req_entry = tk.Entry(section, width=10, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        req_label.grid(row=0, column=0, sticky="w")
        req_entry.grid(row=0, column=1, padx=5)

        # Present
        pres_label = tk.Label(section, text="Staff Present:", fg="white", bg="black", font=label_font)
        pres_entry = tk.Entry(section, width=10, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        pres_label.grid(row=0, column=2, sticky="w")
        pres_entry.grid(row=0, column=3, padx=5)

        # Names (optional)
        names_label = tk.Label(section, text="Names (comma-separated):", fg="white", bg="black", font=label_font)
        names_entry = tk.Entry(section, width=40, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        names_label.grid(row=1, column=0, columnspan=2, sticky="w")
        names_entry.grid(row=1, column=2, columnspan=2, sticky="we", pady=5)

        csc_entries[shift] = {
            "requested": req_entry,
            "present": pres_entry,
            "names": names_entry
        }

    # === Generate Report Buttons ===
    button_frame = tk.Frame(root, bg="black")
    button_frame.pack(pady=10)
    
    save_btn = tk.Button(
        button_frame, text="Save Report", command=save_report_draft,
        bg="white", fg="black", font=("Helvetica", 12, "bold"), padx=10, pady=6,
        relief="raised", activebackground="white", activeforeground="black"
    )
    save_btn.pack(side="left", padx=8)
    
    submit_btn = tk.Button(
        button_frame, text="End Shift", command=end_shift_and_generate,
        bg="white", fg="black", font=("Helvetica", 12, "bold"), padx=10, pady=6,
        relief="raised", activebackground="white", activeforeground="black"
    )
    submit_btn.pack(side="left")

# === Load Draft Functions ===
def load_draft_report():
    """Load a draft report from JSON file and restore the UI state"""
    # Declare global variables at the start
    global building, loaded_from_draft
    
    try:
        # Determine initial directory based on existing entries
        initial_dir = None
        if "date" in entries and entries["date"].get():
            try:
                user_date = entries["date"].get()
                parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
                current_year = parsed_date.strftime("%Y")
                current_month = parsed_date.strftime("%B")
                initial_dir = f"M:\\Sh_BM\\{building}\\Night Reports\\{current_year}\\{current_month}\\drafts"
            except:
                # Fall back to desktop if date parsing fails
                initial_dir = os.path.expanduser("~/Desktop")
        else:
            # Default to desktop if no date entry
            initial_dir = os.path.expanduser("~/Desktop")
        
        # Get the file path from user
        file_path = filedialog.askopenfilename(
            title="Select Draft Report to Load",
            filetypes=[
                ("JSON files", "*.json"),
                ("All files", "*.*")
            ],
            initialdir=initial_dir
        )
        
        if not file_path:
            return  # User cancelled
        
        # Load and validate the JSON data
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Validate the data structure with backward compatibility
        if not isinstance(data, dict):
            raise ValueError("Invalid file format: not a JSON object")
        
        if "building" not in data:
            raise ValueError("Invalid draft file: missing 'building' field")
        
        # ENHANCED: Create missing 'entries' section with sensible defaults
        if "entries" not in data:
            data["entries"] = {}
        
        # ENHANCED: Ensure 'notes' section exists with sensible defaults
        if "notes" not in data:
            data["notes"] = {}
        
        # Store current state in case we need to restore on error
        current_building = building
        current_loaded_state = loaded_from_draft
        
        # Set global variables
        building = data["building"]
        loaded_from_draft = True
        
        # Clear existing tabs if they exist
        if notebook.tabs():
            for tab_id in notebook.tabs():
                notebook.forget(tab_id)
        
        # Configure tabs for the loaded building
        configure_tabs_for_building()
        
        # Populate the form with the loaded data
        populate_form_from_data(data)
        
        # Show success message
        messagebox.showinfo("Success", "Report loaded successfully!")
        
    except FileNotFoundError:
        messagebox.showerror("Load Error", "The selected draft file could not be found.")
    except json.JSONDecodeError as e:
        messagebox.showerror("Load Error", f"Invalid JSON file: {str(e)}")
    except ValueError as e:
        messagebox.showerror("Load Error", str(e))
    except Exception as e:
        # Restore original state on any error
        try:
            building = current_building
            loaded_from_draft = current_loaded_state
        except:
            pass
        messagebox.showerror("Load Error", str(e))

def populate_form_from_data(data):
    """Populate the form fields with data from a loaded draft"""
    try:
        # Clear existing data first
        clear_all_form_data()
        
        # Restore entries
        if "entries" in data:
            entries_data = data.get("entries", {})
            for key, value in entries_data.items():
                if key in entries and hasattr(entries[key], 'delete') and hasattr(entries[key], 'insert'):
                    entries[key].delete(0, tk.END)
                    entries[key].insert(0, str(value))
                    # RULE: Leave everything editable
                    try:
                        entries[key].config(state="normal")
                    except:
                        pass
        
        # Restore building traffic notes
        if "notes" in data and "building_traffic" in data["notes"]:
            restore_note_section(
                data["notes"]["building_traffic"], 
                building_traffic_boxes,
                add_building_traffic_box
            )
        
        # Building-specific data restoration
        if building == "Red Gym":
            restore_red_gym_data(data)
        else:
            restore_union_data(data)
            
    except Exception as e:
        messagebox.showerror("Error", f"Failed to populate form: {str(e)}")

def clear_all_form_data():
    """Clear all form data to prepare for loading a draft"""
    # Clear entries
    for entry in entries.values():
        if hasattr(entry, 'delete'):
            entry.delete(0, tk.END)
    
    # Clear building traffic boxes
    clear_note_boxes(building_traffic_boxes)
    
    if building == "Red Gym":
        clear_red_gym_data()
    else:
        clear_union_data()

def clear_note_boxes(note_boxes):
    """Helper to clear a list of text boxes"""
    for box in note_boxes:
        if hasattr(box, 'delete'):
            box.delete("1.0", tk.END)
            # RULE: Leave everything editable
            try:
                box.config(state="normal")
            except:
                pass

def restore_note_section(note_data, note_boxes, add_function):
    """Helper to restore a section of notes with their text content"""
    if not note_data:
        return
    
    # ENHANCED: Handle type mismatches - convert string to array if needed
    if isinstance(note_data, str):
        note_data = [note_data]  # Convert single string to array
    elif not isinstance(note_data, list):
        note_data = [str(note_data)]  # Convert other types to string array
    
    # RULE: Never assume counts - dynamically add rows/boxes to match JSON
    # Ensure we have enough boxes to match the data
    while len(note_boxes) < len(note_data):
        add_function("")
    
    # RULE: Leave everything editable (state="normal")
    # Set content for each note box
    for i, note_text in enumerate(note_data):
        if i < len(note_boxes):
            # Clear and populate the text box
            note_boxes[i].delete("1.0", tk.END)
            note_boxes[i].insert("1.0", str(note_text))
            
            # Ensure text box is editable
            try:
                note_boxes[i].config(state="normal")
            except:
                pass

def restore_note_section_with_tags(note_data, note_boxes, note_tags, add_function):
    """Helper to restore a section of notes with text and tags"""
    if not note_data:
        return
    
    # ENHANCED: Handle type mismatches - convert string to array if needed
    if isinstance(note_data, str):
        note_data = [note_data]  # Convert single string to array
    elif not isinstance(note_data, list):
        note_data = [str(note_data)]  # Convert other types to string array
    
    # RULE: Never assume counts - dynamically add rows/boxes to match JSON
    # Ensure we have enough boxes to match the data
    while len(note_boxes) < len(note_data):
        add_function("")
    
    # Set content for each note box
    for i, note_entry in enumerate(note_data):
        if i < len(note_boxes):
            # ENHANCED: Handle mixed types within the array
            if isinstance(note_entry, dict):
                text = note_entry.get("text", "")
                tags = note_entry.get("tags", [])
            elif isinstance(note_entry, str):
                text = note_entry
                tags = []
            else:
                text = str(note_entry)
                tags = []
            
            # Clear and populate the text box
            note_boxes[i].delete("1.0", tk.END)
            note_boxes[i].insert("1.0", text)
            
            # RULE: Leave everything editable (state="normal")
            try:
                note_boxes[i].config(state="normal")
            except:
                pass
            
            # RULE: Tag arrays vs note counts - clamp safely
            # RULE: Set comboboxes via .set(value)
            if i < len(note_tags) and tags:
                # Filter out empty/None tags
                valid_tags = [tag for tag in tags if tag and tag != "None"]
                
                # Ensure we have enough tag dropdowns for all valid tags
                while len(note_tags[i]) < len(valid_tags):
                    # Find the tag frame for this note and add more dropdowns
                    # We need to access the parent frame's add_tag_dropdown function
                    # Get the textbox and find its parent frame, then the tag frame
                    textbox = note_boxes[i]
                    note_frame = textbox.master
                    # Look for tag frame (it should be the frame with tag dropdowns)
                    tag_frame = None
                    for child in note_frame.winfo_children():
                        if hasattr(child, 'add_tag_dropdown'):
                            tag_frame = child
                            break
                    
                    if tag_frame and hasattr(tag_frame, 'add_tag_dropdown'):
                        tag_frame.add_tag_dropdown()
                    else:
                        break  # Can't add more dropdowns
                
                # Set the tag values
                for j, tag in enumerate(valid_tags):
                    if j < len(note_tags[i]) and tag:
                        try:
                            note_tags[i][j].set(str(tag))
                        except:
                            pass
                
                # Update add tag buttons after setting values
                # Find the tag frame and call update_add_tag_buttons if available
                textbox = note_boxes[i]
                note_frame = textbox.master
                tag_frame = None
                for child in note_frame.winfo_children():
                    if hasattr(child, 'update_add_tag_buttons'):
                        tag_frame = child
                        break
                
                if tag_frame and hasattr(tag_frame, 'update_add_tag_buttons'):
                    tag_frame.update_add_tag_buttons()

def restore_red_gym_data(data):
    """Restore Red Gym specific data"""
    # Clear Red Gym specific data
    clear_red_gym_data()
    
    # Building tours
    if "red_gym_building_tours" in data and red_gym_building_tours_box:
        red_gym_building_tours_box.delete("1.0", tk.END)
        red_gym_building_tours_box.insert("1.0", str(data.get("red_gym_building_tours", "")))
        # RULE: Leave everything editable
        try:
            red_gym_building_tours_box.config(state="normal")
        except:
            pass
    
    # Deviations count
    if "red_gym_deviations_count" in data and red_gym_deviations_entry:
        red_gym_deviations_entry.delete(0, tk.END)
        red_gym_deviations_entry.insert(0, str(data.get("red_gym_deviations_count", "0")))
        # RULE: Leave everything editable
        try:
            red_gym_deviations_entry.config(state="normal")
        except:
            pass
        # Trigger update to create deviation boxes
        red_gym_deviations_entry.event_generate('<KeyRelease>')
    
    # Deviation notes (populated after the boxes are created)
    if "red_gym_deviation_notes" in data:
        # RULE: Never assume counts - dynamically add to match JSON
        def populate_deviations():
            deviation_notes = data.get("red_gym_deviation_notes", [])
            for i, note_text in enumerate(deviation_notes):
                if i < len(red_gym_deviation_boxes):
                    red_gym_deviation_boxes[i].delete("1.0", tk.END)
                    red_gym_deviation_boxes[i].insert("1.0", str(note_text))
                    # RULE: Leave everything editable
                    try:
                        red_gym_deviation_boxes[i].config(state="normal")
                    except:
                        pass
        root.after(100, populate_deviations)
    
    # Door check
    if "red_gym_door_check_time" in data and red_gym_door_check_time:
        red_gym_door_check_time.delete(0, tk.END)
        red_gym_door_check_time.insert(0, str(data.get("red_gym_door_check_time", "")))
        # RULE: Leave everything editable
        try:
            red_gym_door_check_time.config(state="normal")
        except:
            pass
    
    # RULE: Set comboboxes via .set(value)
    if "red_gym_door_check_day_type" in data and red_gym_door_check_day_type:
        try:
            red_gym_door_check_day_type.set(str(data.get("red_gym_door_check_day_type", "")))
            red_gym_door_check_day_type.config(state="readonly")  # Comboboxes should be readonly
        except:
            pass
    
    # Mail notes
    if "notes" in data and "red_gym_mail" in data["notes"]:
        restore_note_section(
            data["notes"]["red_gym_mail"], 
            red_gym_mail_boxes,
            add_red_gym_mail_box
        )
    
    # Misc notes with tags
    if "notes" in data and "red_gym_misc" in data["notes"]:
        restore_note_section_with_tags(
            data["notes"]["red_gym_misc"], 
            red_gym_misc_boxes,
            misc_note_tags,
            add_red_gym_misc_box
        )

def restore_union_data(data):
    """Restore Memorial Union and Union South specific data"""
    # Clear union specific data
    clear_union_data()
    
    # Mechanical notes with tags
    if "notes" in data and "mechanical" in data["notes"]:
        restore_note_section_with_tags(
            data["notes"]["mechanical"], 
            mechanical_boxes,
            mechanical_note_tags,
            add_mechanical_box
        )
    
    # Production notes with tags
    if "notes" in data and "production" in data["notes"]:
        restore_note_section_with_tags(
            data["notes"]["production"], 
            production_boxes,
            production_note_tags,
            add_production_note_box
        )
    
    # Decibel readings
    if "decibel_readings" in data:
        # RULE: Never assume counts - dynamically add rows to match JSON
        decibel_data = data.get("decibel_readings", [])
        
        # Ensure we have enough rows to match the data
        while len(decibel_entries) < len(decibel_data):
            add_decibel_row()
        
        # Populate decibel readings
        for i, reading_data in enumerate(decibel_data):
            if i < len(decibel_entries):
                time_entry, reading_entry, location_entry = decibel_entries[i]
                
                # RULE: Be backward-compatible - use .get(...) with defaults
                time_entry.delete(0, tk.END)
                time_entry.insert(0, str(reading_data.get("time", "")))
                reading_entry.delete(0, tk.END)  
                reading_entry.insert(0, str(reading_data.get("reading", "")))
                location_entry.delete(0, tk.END)
                location_entry.insert(0, str(reading_data.get("location", "")))
                
                # RULE: Leave everything editable
                try:
                    time_entry.config(state="normal")
                    reading_entry.config(state="normal")
                    location_entry.config(state="normal")
                except:
                    pass
    
    # Patron notes with tags
    if "notes" in data and "patron" in data["notes"]:
        restore_note_section_with_tags(
            data["notes"]["patron"], 
            patron_boxes,
            patron_note_tags,
            add_patron_note_box
        )
    
    # Access inputs
    if "access_inputs" in data:
        access_data = data.get("access_inputs", {})
        for key, value in access_data.items():
            if key in access_inputs:
                # RULE: Set comboboxes via .set(value) and entries via insert
                if hasattr(access_inputs[key], 'delete') and hasattr(access_inputs[key], 'insert'):
                    access_inputs[key].delete(0, tk.END)
                    access_inputs[key].insert(0, str(value))
                    # RULE: Leave everything editable
                    try:
                        access_inputs[key].config(state="normal")
                    except:
                        pass
                elif hasattr(access_inputs[key], 'set'):
                    # RULE: Set comboboxes via .set(value)
                    try:
                        access_inputs[key].set(str(value))
                        access_inputs[key].config(state="readonly")  # Comboboxes should be readonly
                    except:
                        pass
    
    # Access notes with tags
    if "notes" in data and "access" in data["notes"]:
        restore_note_section_with_tags(
            data["notes"]["access"], 
            access_note_boxes,
            access_note_tags,
            add_access_note
        )
    
    # Cash notes with tags
    if "notes" in data and "cash" in data["notes"]:
        restore_note_section_with_tags(
            data["notes"]["cash"], 
            cash_boxes,
            cash_note_tags,
            add_cash_note_box
        )
    
    # Dining notes with tags
    if "notes" in data and "dining" in data["notes"]:
        restore_note_section_with_tags(
            data["notes"]["dining"], 
            dining_boxes,
            dining_note_tags,
            add_dining_note_box
        )
    
    # Hotel notes
    if "notes" in data and "hotel" in data["notes"]:
        restore_note_section(
            data["notes"]["hotel"], 
            hotel_boxes,
            add_hotel_note_box
        )
    
    # Misc notes
    if "notes" in data and "misc" in data["notes"]:
        restore_note_section(
            data["notes"]["misc"], 
            misc_boxes,
            add_misc_note_box
        )
    
    # Memorial Union specific sections
    if building == "Memorial Union":
        # Carding notes
        if "notes" in data and "carding" in data["notes"]:
            restore_note_section(
                data["notes"]["carding"], 
                carding_boxes,
                add_carding_note_box
            )
        
        # Terrace notes without tags
        if "notes" in data and "terrace" in data["notes"]:
            restore_note_section(
                data["notes"]["terrace"], 
                terrace_boxes,
                add_terrace_note_box
            )
        
        # Enforcement notes with tags and images
        if "notes" in data and "enforcement" in data["notes"]:
            restore_enforcement_notes(data["notes"]["enforcement"])
        
        # Alumni notes with tags
        if "notes" in data and "alumni" in data["notes"]:
            restore_note_section_with_tags(
                data["notes"]["alumni"], 
                alumni_boxes,
                alumni_note_tags,
                add_alumni_note_box
            )
        
        # Pier notes with tags
        if "notes" in data and "pier" in data["notes"]:
            restore_note_section_with_tags(
                data["notes"]["pier"], 
                pier_boxes,
                pier_note_tags,
                add_pier_note_box
            )
    
    # CSC/Security data
    if "csc" in data:
        csc_data = data.get("csc", {})
        for shift, shift_data in csc_data.items():
            if shift in csc_entries:
                for field, value in shift_data.items():
                    if field in csc_entries[shift]:
                        csc_entries[shift][field].delete(0, tk.END)
                        csc_entries[shift][field].insert(0, str(value))
                        # RULE: Leave everything editable
                        try:
                            csc_entries[shift][field].config(state="normal")
                        except:
                            pass

def restore_enforcement_notes(enforcement_data):
    """Special handler for enforcement notes with images"""
    if not enforcement_data:
        return
    
    # RULE: Never assume counts - dynamically add to match JSON
    # Clear existing enforcement data by clearing the lists
    enforcement_boxes.clear()
    enforcement_note_tags.clear()
    enforcement_images.clear()
    
    # RULE: Never assume counts - dynamically add entries to match JSON
    # Restore enforcement entries
    for i, note_entry in enumerate(enforcement_data):
        # RULE: Be backward-compatible - use .get(...) with defaults
        if isinstance(note_entry, dict):
            text = note_entry.get("text", "")
            tags = note_entry.get("tags", [])
            image_path = note_entry.get("image_path", "")
        else:
            text = str(note_entry)
            tags = []
            image_path = ""
        
        if image_path:
            # Add as image entry
            add_enforcement_image()
            if len(enforcement_images) > 0:
                # RULE: Missing image files - still show filename in status label; allow reselection
                enforcement_images[-1].set(str(image_path))
                # Find the status label and update it to show the filename
                # (The add_enforcement_image function will handle missing files gracefully)
        else:
            # Add as text note
            add_enforcement_note_box("")
        
        # Set text content
        if len(enforcement_boxes) > 0:
            enforcement_boxes[-1].delete("1.0", tk.END)
            enforcement_boxes[-1].insert("1.0", str(text))
            # RULE: Leave everything editable
            try:
                enforcement_boxes[-1].config(state="normal")
            except:
                pass
        
        # RULE: Set comboboxes via .set(value)
        # RULE: Tag arrays vs note counts - clamp safely
        if len(enforcement_note_tags) > 0 and tags:
            for j, tag in enumerate(tags):
                if j < len(enforcement_note_tags[-1]) and tag:
                    try:
                        enforcement_note_tags[-1][j].set(str(tag))
                    except:
                        pass
            
            # Update add tag buttons after setting values for enforcement
            if len(enforcement_boxes) > 0:
                textbox = enforcement_boxes[-1]
                note_frame = textbox.master
                tag_frame = None
                for child in note_frame.winfo_children():
                    if hasattr(child, 'update_add_tag_buttons'):
                        tag_frame = child
                        break
                
                if tag_frame and hasattr(tag_frame, 'update_add_tag_buttons'):
                    tag_frame.update_add_tag_buttons()

def clear_red_gym_data():
    """Clear Red Gym specific form data"""
    if red_gym_building_tours_box:
        red_gym_building_tours_box.delete("1.0", tk.END)
        # RULE: Leave everything editable
        try:
            red_gym_building_tours_box.config(state="normal")
        except:
            pass
    
    if red_gym_deviations_entry:
        red_gym_deviations_entry.delete(0, tk.END)
        red_gym_deviations_entry.insert(0, "0")
        # RULE: Leave everything editable
        try:
            red_gym_deviations_entry.config(state="normal")
        except:
            pass
    
    if red_gym_door_check_time:
        red_gym_door_check_time.delete(0, tk.END)
        # RULE: Leave everything editable
        try:
            red_gym_door_check_time.config(state="normal")
        except:
            pass
    
    if red_gym_door_check_day_type:
        red_gym_door_check_day_type.set("")
        # RULE: Comboboxes should remain readonly but functional
        try:
            red_gym_door_check_day_type.config(state="readonly")
        except:
            pass
    
    clear_note_boxes(red_gym_mail_boxes)
    clear_note_boxes(red_gym_misc_boxes)

def clear_union_data():
    """Clear Memorial Union and Union South specific form data"""
    clear_note_boxes(mechanical_boxes)
    clear_note_boxes(production_boxes)
    clear_note_boxes(patron_boxes)
    clear_note_boxes(access_note_boxes)
    clear_note_boxes(cash_boxes)
    clear_note_boxes(dining_boxes)
    clear_note_boxes(hotel_boxes)
    clear_note_boxes(misc_boxes)
    
    # Clear access inputs
    for widget in access_inputs.values():
        if hasattr(widget, 'delete') and hasattr(widget, 'insert'):
            widget.delete(0, tk.END)
            # RULE: Leave everything editable
            try:
                widget.config(state="normal")
            except:
                pass
        elif hasattr(widget, 'set'):
            widget.set("")
            # RULE: Comboboxes should remain readonly but functional
            try:
                widget.config(state="readonly")
            except:
                pass
    
    # Clear CSC data
    for shift_data in csc_entries.values():
        for field_widget in shift_data.values():
            field_widget.delete(0, tk.END)
            # RULE: Leave everything editable
            try:
                field_widget.config(state="normal")
            except:
                pass
    
    # Clear decibel entries
    for time_entry, reading_entry, location_entry in decibel_entries:
        time_entry.delete(0, tk.END)
        reading_entry.delete(0, tk.END)
        location_entry.delete(0, tk.END)
        time_entry.insert(0, "Time")
        reading_entry.insert(0, "Reading (db)")
        location_entry.insert(0, "Location")
        # RULE: Leave everything editable
        try:
            time_entry.config(state="normal")
            reading_entry.config(state="normal")
            location_entry.config(state="normal")
        except:
            pass
    
    # Memorial Union specific
    if building == "Memorial Union":
        clear_note_boxes(carding_boxes)
        clear_note_boxes(terrace_boxes)
        clear_note_boxes(enforcement_boxes)
        clear_note_boxes(alumni_boxes)
        clear_note_boxes(pier_boxes)

# === Draft Save Function ===
def save_report_draft():
    try:
        # Create timestamp
        now = datetime.now()
        timestamp = now.isoformat()
        
        # Build the draft data structure
        draft_data = {
            "timestamp": timestamp,
            "building": building,
            "entries": {},
            "notes": {},
            "decibel_readings": [],
            "csc": {},
            "access_inputs": {}
        }
        
        # Save all entries
        for key, entry in entries.items():
            if hasattr(entry, 'get'):
                draft_data["entries"][key] = entry.get()
        
        # Save building traffic notes
        traffic_notes = []
        for box in building_traffic_boxes:
            traffic_notes.append(box.get("1.0", "end-1c"))
        draft_data["notes"]["building_traffic"] = traffic_notes
        
        # Red Gym specific data
        if building == "Red Gym":
            # Building tours
            if red_gym_building_tours_box:
                draft_data["red_gym_building_tours"] = red_gym_building_tours_box.get("1.0", "end-1c")
            
            # Deviations
            if red_gym_deviations_entry:
                draft_data["red_gym_deviations_count"] = red_gym_deviations_entry.get()
            
            deviation_notes = []
            for box in red_gym_deviation_boxes:
                deviation_notes.append(box.get("1.0", "end-1c"))
            draft_data["red_gym_deviation_notes"] = deviation_notes
            
            # Door check
            if red_gym_door_check_time:
                draft_data["red_gym_door_check_time"] = red_gym_door_check_time.get()
            if red_gym_door_check_day_type:
                draft_data["red_gym_door_check_day_type"] = red_gym_door_check_day_type.get()
            
            # Mail notes
            mail_notes = []
            for box in red_gym_mail_boxes:
                mail_notes.append(box.get("1.0", "end-1c"))
            draft_data["notes"]["red_gym_mail"] = mail_notes
            
            # Misc notes with tags
            misc_notes = []
            for i, box in enumerate(red_gym_misc_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(misc_note_tags):
                    note_data["tags"] = [var.get() for var in misc_note_tags[i]]
                misc_notes.append(note_data)
            draft_data["notes"]["red_gym_misc"] = misc_notes
        else:
            # Memorial Union and Union South specific data
            
            # Save mechanical notes with tags
            mechanical_notes = []
            for i, box in enumerate(mechanical_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(mechanical_note_tags):
                    note_data["tags"] = [var.get() for var in mechanical_note_tags[i]]
                mechanical_notes.append(note_data)
            draft_data["notes"]["mechanical"] = mechanical_notes
            
            # Save production notes with tags
            production_notes = []
            for i, box in enumerate(production_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(production_note_tags):
                    note_data["tags"] = [var.get() for var in production_note_tags[i]]
                production_notes.append(note_data)
            draft_data["notes"]["production"] = production_notes
            
            # Save decibel readings
            for time_entry, reading_entry, location_entry in decibel_entries:
                reading_data = {
                    "time": time_entry.get(),
                    "reading": reading_entry.get(),
                    "location": location_entry.get()
                }
                draft_data["decibel_readings"].append(reading_data)
            
            # Save patron notes with tags
            patron_notes = []
            for i, box in enumerate(patron_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(patron_note_tags):
                    note_data["tags"] = [var.get() for var in patron_note_tags[i]]
                patron_notes.append(note_data)
            draft_data["notes"]["patron"] = patron_notes
            
            # Save access inputs
            for key, widget in access_inputs.items():
                if hasattr(widget, 'get'):
                    draft_data["access_inputs"][key] = widget.get()
            
            # Save access notes with tags
            access_notes = []
            for i, box in enumerate(access_note_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(access_note_tags):
                    note_data["tags"] = [var.get() for var in access_note_tags[i]]
                access_notes.append(note_data)
            draft_data["notes"]["access"] = access_notes
            
            # Save cash notes with tags
            cash_notes = []
            for i, box in enumerate(cash_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(cash_note_tags):
                    note_data["tags"] = [var.get() for var in cash_note_tags[i]]
                cash_notes.append(note_data)
            draft_data["notes"]["cash"] = cash_notes
            
            # Save dining notes with tags
            dining_notes = []
            for i, box in enumerate(dining_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(dining_note_tags):
                    note_data["tags"] = [var.get() for var in dining_note_tags[i]]
                dining_notes.append(note_data)
            draft_data["notes"]["dining"] = dining_notes
            
            # Save hotel notes
            hotel_notes = []
            for box in hotel_boxes:
                hotel_notes.append(box.get("1.0", "end-1c"))
            draft_data["notes"]["hotel"] = hotel_notes
            
            # Save misc notes
            misc_notes = []
            for box in misc_boxes:
                misc_notes.append(box.get("1.0", "end-1c"))
            draft_data["notes"]["misc"] = misc_notes
            
            # Memorial Union specific sections
            if building == "Memorial Union":
                # Save carding notes
                carding_notes = []
                for box in carding_boxes:
                    carding_notes.append(box.get("1.0", "end-1c"))
                draft_data["notes"]["carding"] = carding_notes
                
                # Save terrace notes without tags
                terrace_notes = []
                for box in terrace_boxes:
                    terrace_notes.append(box.get("1.0", "end-1c"))
                draft_data["notes"]["terrace"] = terrace_notes
                
                # Save enforcement notes with tags and images
                enforcement_notes = []
                for i, box in enumerate(enforcement_boxes):
                    note_data = {
                        "text": box.get("1.0", "end-1c"),
                        "tags": [],
                        "image_path": ""
                    }
                    if i < len(enforcement_note_tags):
                        note_data["tags"] = [var.get() for var in enforcement_note_tags[i]]
                    if i < len(enforcement_images):
                        note_data["image_path"] = enforcement_images[i].get()
                    enforcement_notes.append(note_data)
                draft_data["notes"]["enforcement"] = enforcement_notes
                
                # Save alumni notes with tags
                alumni_notes = []
                for i, box in enumerate(alumni_boxes):
                    note_data = {
                        "text": box.get("1.0", "end-1c"),
                        "tags": []
                    }
                    if i < len(alumni_note_tags):
                        note_data["tags"] = [var.get() for var in alumni_note_tags[i]]
                    alumni_notes.append(note_data)
                draft_data["notes"]["alumni"] = alumni_notes
                
                # Save pier notes with tags
                pier_notes = []
                for i, box in enumerate(pier_boxes):
                    note_data = {
                        "text": box.get("1.0", "end-1c"),
                        "tags": []
                    }
                    if i < len(pier_note_tags):
                        note_data["tags"] = [var.get() for var in pier_note_tags[i]]
                    pier_notes.append(note_data)
                draft_data["notes"]["pier"] = pier_notes
            
            # Save CSC/Security data
            for shift in csc_shifts:
                if shift in csc_entries:
                    draft_data["csc"][shift] = {
                        "requested": csc_entries[shift]["requested"].get(),
                        "present": csc_entries[shift]["present"].get(),
                        "names": csc_entries[shift]["names"].get()
                    }
        
        # Create draft folder path
        user_date = entries["date"].get()
        parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
        current_year = parsed_date.strftime("%Y")
        current_month = parsed_date.strftime("%B")
        
        base_dir = f"M:\\Sh_BM\\{building}\\Night Reports"
        year_dir = os.path.join(base_dir, current_year)
        month_dir = os.path.join(year_dir, current_month)
        drafts_dir = os.path.join(month_dir, "drafts")
        os.makedirs(drafts_dir, exist_ok=True)
        
        # Create draft filename with timestamp
        draft_filename = f"draft_{now.strftime('%Y-%m-%d_%H-%M')}.json"
        draft_path = os.path.join(drafts_dir, draft_filename)
        
        # Save the JSON file (overwrite if exists)
        with open(draft_path, 'w', encoding='utf-8') as f:
            json.dump(draft_data, f, indent=2, ensure_ascii=False)
        
        messagebox.showinfo("Draft Saved", f"Report draft saved to:\n{draft_path}")
        
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save draft: {str(e)}")

# === Autosave Functions ===
def start_autosave(interval_min=5):
    """Start autosave mechanism that saves drafts every interval_min minutes"""
    global autosave_thread, autosave_running
    
    # Don't start multiple autosave threads
    if autosave_thread is not None and autosave_thread.is_alive():
        return
    
    autosave_running = True
    
    def autosave_worker():
        """Background worker that triggers autosave periodically"""
        while autosave_running:
            # Sleep for the specified interval
            time.sleep(interval_min * 60)
            
            # Check if we should still be running
            if not autosave_running:
                break
                
            # Schedule save_report_draft on the main thread
            try:
                root.after(0, lambda: save_report_draft_silent())
            except Exception as e:
                print(f"Autosave error: {e}")
    
    # Create and start daemon thread
    autosave_thread = threading.Thread(target=autosave_worker, daemon=True)
    autosave_thread.start()

def stop_autosave():
    """Stop the autosave mechanism"""
    global autosave_running
    autosave_running = False

def save_report_draft_silent():
    """Silent version of save_report_draft for autosave (no popup messages)"""
    try:
        # Create timestamp
        now = datetime.now()
        timestamp = now.isoformat()
        
        # Build the draft data structure (same as save_report_draft)
        draft_data = {
            "timestamp": timestamp,
            "building": building,
            "entries": {},
            "notes": {},
            "decibel_readings": [],
            "csc": {},
            "access_inputs": {}
        }
        
        # Save all entries
        for key, entry in entries.items():
            if hasattr(entry, 'get'):
                draft_data["entries"][key] = entry.get()
        
        # Save building traffic notes
        traffic_notes = []
        for box in building_traffic_boxes:
            traffic_notes.append(box.get("1.0", "end-1c"))
        draft_data["notes"]["building_traffic"] = traffic_notes
        
        # Red Gym specific data
        if building == "Red Gym":
            # Building tours
            if red_gym_building_tours_box:
                draft_data["red_gym_building_tours"] = red_gym_building_tours_box.get("1.0", "end-1c")
            
            # Deviations
            if red_gym_deviations_entry:
                draft_data["red_gym_deviations_count"] = red_gym_deviations_entry.get()
            
            deviation_notes = []
            for box in red_gym_deviation_boxes:
                deviation_notes.append(box.get("1.0", "end-1c"))
            draft_data["red_gym_deviation_notes"] = deviation_notes
            
            # Door check
            if red_gym_door_check_time:
                draft_data["red_gym_door_check_time"] = red_gym_door_check_time.get()
            if red_gym_door_check_day_type:
                draft_data["red_gym_door_check_day_type"] = red_gym_door_check_day_type.get()
            
            # Mail notes
            mail_notes = []
            for box in red_gym_mail_boxes:
                mail_notes.append(box.get("1.0", "end-1c"))
            draft_data["notes"]["red_gym_mail"] = mail_notes
            
            # Misc notes with tags
            misc_notes = []
            for i, box in enumerate(red_gym_misc_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(misc_note_tags):
                    note_data["tags"] = [var.get() for var in misc_note_tags[i]]
                misc_notes.append(note_data)
            draft_data["notes"]["red_gym_misc"] = misc_notes
        else:
            # Memorial Union and Union South specific data
            
            # Save mechanical notes with tags
            mechanical_notes = []
            for i, box in enumerate(mechanical_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(mechanical_note_tags):
                    note_data["tags"] = [var.get() for var in mechanical_note_tags[i]]
                mechanical_notes.append(note_data)
            draft_data["notes"]["mechanical"] = mechanical_notes
            
            # Save production notes with tags
            production_notes = []
            for i, box in enumerate(production_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(production_note_tags):
                    note_data["tags"] = [var.get() for var in production_note_tags[i]]
                production_notes.append(note_data)
            draft_data["notes"]["production"] = production_notes
            
            # Save decibel readings
            for time_entry, reading_entry, location_entry in decibel_entries:
                reading_data = {
                    "time": time_entry.get(),
                    "reading": reading_entry.get(),
                    "location": location_entry.get()
                }
                draft_data["decibel_readings"].append(reading_data)
            
            # Save patron notes with tags
            patron_notes = []
            for i, box in enumerate(patron_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(patron_note_tags):
                    note_data["tags"] = [var.get() for var in patron_note_tags[i]]
                patron_notes.append(note_data)
            draft_data["notes"]["patron"] = patron_notes
            
            # Save access inputs
            for key, widget in access_inputs.items():
                if hasattr(widget, 'get'):
                    draft_data["access_inputs"][key] = widget.get()
            
            # Save access notes with tags
            access_notes = []
            for i, box in enumerate(access_note_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(access_note_tags):
                    note_data["tags"] = [var.get() for var in access_note_tags[i]]
                access_notes.append(note_data)
            draft_data["notes"]["access"] = access_notes
            
            # Save cash notes with tags
            cash_notes = []
            for i, box in enumerate(cash_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(cash_note_tags):
                    note_data["tags"] = [var.get() for var in cash_note_tags[i]]
                cash_notes.append(note_data)
            draft_data["notes"]["cash"] = cash_notes
            
            # Save dining notes with tags
            dining_notes = []
            for i, box in enumerate(dining_boxes):
                note_data = {
                    "text": box.get("1.0", "end-1c"),
                    "tags": []
                }
                if i < len(dining_note_tags):
                    note_data["tags"] = [var.get() for var in dining_note_tags[i]]
                dining_notes.append(note_data)
            draft_data["notes"]["dining"] = dining_notes
            
            # Save hotel notes
            hotel_notes = []
            for box in hotel_boxes:
                hotel_notes.append(box.get("1.0", "end-1c"))
            draft_data["notes"]["hotel"] = hotel_notes
            
            # Save misc notes
            misc_notes = []
            for box in misc_boxes:
                misc_notes.append(box.get("1.0", "end-1c"))
            draft_data["notes"]["misc"] = misc_notes
            
            # Memorial Union specific sections
            if building == "Memorial Union":
                # Save carding notes
                carding_notes = []
                for box in carding_boxes:
                    carding_notes.append(box.get("1.0", "end-1c"))
                draft_data["notes"]["carding"] = carding_notes
                
                # Save terrace notes without tags
                terrace_notes = []
                for box in terrace_boxes:
                    terrace_notes.append(box.get("1.0", "end-1c"))
                draft_data["notes"]["terrace"] = terrace_notes
                
                # Save enforcement notes with tags and images
                enforcement_notes = []
                for i, box in enumerate(enforcement_boxes):
                    note_data = {
                        "text": box.get("1.0", "end-1c"),
                        "tags": [],
                        "image_path": ""
                    }
                    if i < len(enforcement_note_tags):
                        note_data["tags"] = [var.get() for var in enforcement_note_tags[i]]
                    if i < len(enforcement_images):
                        note_data["image_path"] = enforcement_images[i].get()
                    enforcement_notes.append(note_data)
                draft_data["notes"]["enforcement"] = enforcement_notes
                
                # Save alumni notes with tags
                alumni_notes = []
                for i, box in enumerate(alumni_boxes):
                    note_data = {
                        "text": box.get("1.0", "end-1c"),
                        "tags": []
                    }
                    if i < len(alumni_note_tags):
                        note_data["tags"] = [var.get() for var in alumni_note_tags[i]]
                    alumni_notes.append(note_data)
                draft_data["notes"]["alumni"] = alumni_notes
                
                # Save pier notes with tags
                pier_notes = []
                for i, box in enumerate(pier_boxes):
                    note_data = {
                        "text": box.get("1.0", "end-1c"),
                        "tags": []
                    }
                    if i < len(pier_note_tags):
                        note_data["tags"] = [var.get() for var in pier_note_tags[i]]
                    pier_notes.append(note_data)
                draft_data["notes"]["pier"] = pier_notes
            
            # Save CSC/Security data
            for shift in csc_shifts:
                if shift in csc_entries:
                    draft_data["csc"][shift] = {
                        "requested": csc_entries[shift]["requested"].get(),
                        "present": csc_entries[shift]["present"].get(),
                        "names": csc_entries[shift]["names"].get()
                    }
        
        # Create draft folder path
        user_date = entries["date"].get()
        parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
        current_year = parsed_date.strftime("%Y")
        current_month = parsed_date.strftime("%B")
        
        base_dir = f"M:\\Sh_BM\\{building}\\Night Reports"
        year_dir = os.path.join(base_dir, current_year)
        month_dir = os.path.join(year_dir, current_month)
        drafts_dir = os.path.join(month_dir, "drafts")
        os.makedirs(drafts_dir, exist_ok=True)
        
        # Delete previous autosave files to keep only the latest one
        try:
            for filename in os.listdir(drafts_dir):
                if filename.startswith("autosave_") and filename.endswith(".json"):
                    old_autosave_path = os.path.join(drafts_dir, filename)
                    os.remove(old_autosave_path)
                    print(f"Deleted old autosave: {filename}")
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up old autosave files: {cleanup_error}")
        
        # Create autosave filename (different from manual saves)
        draft_filename = f"autosave_{now.strftime('%Y-%m-%d_%H-%M')}.json"
        draft_path = os.path.join(drafts_dir, draft_filename)
        
        # Save the JSON file (overwrite if exists)
        with open(draft_path, 'w', encoding='utf-8') as f:
            json.dump(draft_data, f, indent=2, ensure_ascii=False)
        
        # Optional: Print to console for debugging (no popup)
        print(f"Autosave completed: {draft_path}")
        
    except Exception as e:
        # Silent error handling for autosave
        print(f"Autosave failed: {str(e)}")

# === Rename existing function ===
def end_shift_and_generate():
    try:
        # Stop autosave before generating final report
        stop_autosave()
        
        # Generate the report using existing logic
        generate_report()
        
        # After successful generation, delete the drafts folder
        try:
            user_date = entries["date"].get()
            parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
            current_year = parsed_date.strftime("%Y")
            current_month = parsed_date.strftime("%B")
            drafts_dir = os.path.join(f"M:\\Sh_BM\\{building}\\Night Reports", current_year, current_month, "drafts")
            
            if os.path.exists(drafts_dir):
                import shutil
                shutil.rmtree(drafts_dir)
        except Exception as e:
            # Don't fail the entire operation if draft cleanup fails
            print(f"Warning: Could not clean up drafts folder: {e}")
            
    except Exception as e:
        messagebox.showerror("Error", str(e))

# === Generate Report Logic === (renamed from generate_report)
def generate_report():
    try:
        doc = Document()
        # Include building name in the heading
        doc.add_heading(f'{building.upper()}\nBUILDING MANAGER\'S NIGHT REPORT', level=1)
        
        # Add bold paragraphs with regular user input
        def add_bold_para_with_input(bold_text, user_input):
            p = doc.add_paragraph()
            p.add_run(f"{bold_text}: ").bold = True
            p.add_run(user_input)
        
        add_bold_para_with_input("Date", entries["date"].get())
        add_bold_para_with_input("Shift Hours", entries["shift_hours"].get())
        add_bold_para_with_input("Building Manager(s)", entries["bms"].get())
        
        # Red Gym specific report format
        if building == "Red Gym":
            # Bold section headers
            p = doc.add_paragraph()
            p.add_run("\nNotes:").bold = True
            
            def add_bold_section_header(text):
                p = doc.add_paragraph()
                p.add_run(text).bold = True
                return p
                
            # Function to add indented paragraphs using Inches for better control
            def add_indented_paragraph(number, content):
                p = doc.add_paragraph("")
                p.paragraph_format.left_indent = Inches(0.25)
                p.add_run(f"{number}. ").bold = True
                p.add_run(content)
                return p
            
            def add_sub_indented_paragraph(letter, content):
                p = doc.add_paragraph("")
                p.paragraph_format.left_indent = Inches(0.75)
                p.add_run(f"{letter}. ").bold = True
                p.add_run(content)
                return p

            note_counter = 1

            # Building Traffic
            add_bold_section_header("Building Traffic")
            has_traffic_notes = False
            for box in building_traffic_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_traffic_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            if not has_traffic_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # Security
            add_bold_section_header("Security")
            
            # Building Tours
            tours_content = red_gym_building_tours_box.get("1.0", "end").strip()
            tours_text = f"Building Tours: {tours_content}" if tours_content else "Building Tours:"
            add_indented_paragraph(note_counter, tours_text)
            note_counter += 1
            
            # Deviations
            num_deviations = int(red_gym_deviations_entry.get() or "0")
            deviation_text = f"There were {num_deviations} deviations from the standard building locking protocol today."
            
            # Create paragraph for deviations with bold formatting
            p = doc.add_paragraph("")
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(f"{note_counter}. ").bold = True
            p.add_run("There were ")
            p.add_run(str(num_deviations)).bold = True
            p.add_run(" deviations from the standard building locking protocol today.")
            
            if num_deviations > 0:
                for i, box in enumerate(red_gym_deviation_boxes):
                    content = box.get("1.0", "end").strip()
                    letter = chr(97 + i)  # a, b, c, etc.
                    add_sub_indented_paragraph(letter, content)
            
            note_counter += 1
            
            # Door check
            door_time = red_gym_door_check_time.get().strip()   
            door_day_type = red_gym_door_check_day_type.get()
            
            # Create paragraph for door check with bold formatting
            p = doc.add_paragraph("")
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(f"{note_counter}. ").bold = True
            p.add_run("I checked to confirm that the Building Doors were locked and the door swipe scanner was red at ")
            p.add_run(door_time).bold = True
            p.add_run(" on a ")
            p.add_run(door_day_type).bold = True
            p.add_run(".")
            
            note_counter += 1

            # Mail
            add_bold_section_header("Mail")
            has_mail_notes = False
            for box in red_gym_mail_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_mail_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            if not has_mail_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # Miscellaneous
            add_bold_section_header("Miscellaneous")
            has_misc_notes = False
            for box in red_gym_misc_boxes:
                content = box.get("1.0", "end").strip()
               
                if content:
                    has_misc_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            if not has_misc_notes:
                add_indented_paragraph(note_counter, "None")
                note_counter += 1

        else:
            # Existing code for Memorial Union and Union South
            # Terrace Manager(s) only for Memorial Union
            if building == "Memorial Union":
                add_bold_para_with_input("Terrace Manager(s)", entries["terrace_managers"].get())
            add_bold_para_with_input("Event Manager(s)", entries["eventmanagers"].get())
            add_bold_para_with_input("Guest Service Specialist", entries["gss"].get())
            add_bold_para_with_input("Operation Manager(s)", entries["operation_managers"].get())
            add_bold_para_with_input("Custodial Supervisor(s)", entries["custodial"].get())
            add_bold_para_with_input("Production Supervisor(s)", entries["production"].get())
            add_bold_para_with_input("Retail & Dining Supervisor(s)", entries["retail"].get())
            add_bold_para_with_input("Catering Supervisor(s)", entries["catering"].get())
            add_bold_para_with_input("CAVR Desk Staff", entries["cavr"].get())

            # Bold section headers
            p = doc.add_paragraph()
            p.add_run("\nNotes:").bold = True
            
            def add_bold_section_header(text):
                p = doc.add_paragraph()
                p.add_run(text).bold = True
                return p
                
            # Function to add indented paragraphs using Inches for better control
            def add_indented_paragraph(number, content):
                # Create a paragraph with the content and number
                p = doc.add_paragraph("")
                p.paragraph_format.left_indent = Inches(0.25)  # Change from 0.5 to 0.25 inch indent
                
                # Add the number with bold formatting
                p.add_run(f"{number}. ").bold = True
                
                # Add the content directly
                p.add_run(content)
                
                return p

            # Start a global counter for note numbering across all sections
            note_counter = 1

            add_bold_section_header("Building Traffic")
            
            # Add building traffic notes with global counter
            has_traffic_notes = False
            for box in building_traffic_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_traffic_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_traffic_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            add_bold_section_header("Mechanical/Repairs/Custodial")
            
            # Add mechanical notes with continuing counter
            has_mechanical_notes = False
            for box in mechanical_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_mechanical_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_mechanical_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            add_bold_section_header("Production Services (Meetings, Events, Set-ups, AV)")
            
            # Add production notes with continuing counter
            has_production_notes = False
            for box in production_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_production_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_production_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1
                
            # === Decibel reading and Security table modifications ===
            
            if decibel_entries:
                add_bold_section_header("Decibel Readings")
                # Create a table for decibel readings
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Add header row - bold the headers
                header_cells = table.rows[0].cells
                header_cells[0].paragraphs[0].add_run("Time").bold = True
                header_cells[1].paragraphs[0].add_run("Reading (dB)").bold = True
                header_cells[2].paragraphs[0].add_run("Location").bold = True
                
                # Center-align header cells
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 1  # 1 = CENTER

                # Add data rows
                for time_entry, reading_entry, location_entry in decibel_entries:
                    time = time_entry.get().strip()
                    reading = reading_entry.get().strip()
                    location = location_entry.get().strip()
                    if time and reading and location and time != "Time" and reading != "Reading (db)" and location != "Location":
                        row_cells = table.add_row().cells
                        row_cells[0].text = time
                        row_cells[1].text = reading
                        row_cells[2].text = location
                        
                        # Center-align all cells in this row
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = 1  # 1 = CENTER

            add_bold_section_header("Patron Services (Membership, Patron Assistance, Problem Patrons)")
            
            # Add patron service notes with continuing counter
            has_patron_notes = False
            for box in patron_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_patron_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_patron_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            add_bold_section_header("Access/Lock/Unlock")

            # Special case for access notes - bold the user input instead
            access_notes = [
                {
                    "text": f"At the early check, the loading dock arm gate was {access_inputs['early_gate'].get().lower()} at {access_inputs['early_time'].get()}.",
                    "bold_parts": [access_inputs['early_gate'].get().lower(), access_inputs['early_time'].get()]
                },
                {
                    "text": f"At the closing check, the loading dock arm gate was {access_inputs['close_gate'].get().lower()} at {access_inputs['close_time'].get()}.",
                    "bold_parts": [access_inputs['close_gate'].get().lower(), access_inputs['close_time'].get()]
                },
                {
                    "text": f"At the closing check, the HID scanners were {access_inputs['hid_status'].get().lower()}.",
                    "bold_parts": [access_inputs['hid_status'].get().lower()]
                },
                {
                    "text": f"I {access_inputs['door_status'].get().lower()} secured the loading dock overhead door for the night.",
                    "bold_parts": [access_inputs['door_status'].get().lower()]
                }
            ]

            # Auto-generated sentences with bold user inputs and continuing counter with indentation
            for note_data in access_notes:
                p = doc.add_paragraph("")
                # Indent the entire paragraph using Inches
                p.paragraph_format.left_indent = Inches(0.25)  # Change from 0.5 to 0.25
                
                # Add the number with bold formatting
                p.add_run(f"{note_counter}. ").bold = True
                
                sentence = note_data["text"]
                bold_parts = note_data["bold_parts"]
                
                current_pos = 0
                for bold_part in bold_parts:
                    if bold_part in sentence[current_pos:]:
                        # Find position of bold part
                        part_pos = sentence.find(bold_part, current_pos)
                        
                        # Add text before the bold part
                        if part_pos > current_pos:
                            p.add_run(sentence[current_pos:part_pos])
                        
                        # Add the bold part
                        p.add_run(bold_part).bold = True
                        
                        # Update position
                        current_pos = part_pos + len(bold_part)
                
                # Add any remaining text after the last bold part
                if current_pos < len(sentence):
                    p.add_run(sentence[current_pos:])
                    
                note_counter += 1
                
            # User-entered access notes
            for box in access_note_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1

            add_bold_section_header("Cash Office")
            
            # Add cash office notes with continuing counter
            has_cash_notes = False
            for box in cash_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_cash_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_cash_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # Only include Memorial Union specific sections if the building is Memorial Union
            if building == "Memorial Union":
                add_bold_section_header("Carding Runs")
                
                # Add carding notes with continuing counter
                has_carding_notes = False
                for box in carding_boxes:
                    content = box.get("1.0", "end").strip()
                    if content:
                        has_carding_notes = True
                        add_indented_paragraph(note_counter, content)
                        note_counter += 1
                
                # Add an empty numbered note if no content
                if not has_carding_notes:
                    add_indented_paragraph(note_counter, "")
                    note_counter += 1

                add_bold_section_header("Terrace Traffic")
                
                # Add terrace traffic notes with continuing counter
                has_terrace_notes = False
                for box in terrace_boxes:
                    content = box.get("1.0", "end").strip()
                    if content:
                        has_terrace_notes = True
                        add_indented_paragraph(note_counter, content)
                        note_counter += 1
                
                # Add an empty numbered note if no content
                if not has_terrace_notes:
                    add_indented_paragraph(note_counter, "")
                    note_counter += 1

                add_bold_section_header("Terrace Enforcement")
                
                # Add enforcement notes with continuing counter
                has_enforcement_notes = False
                enforcement_index = 0  # Track index for both images and text notes
                
                for i, box in enumerate(enforcement_boxes):
                    content = box.get("1.0", "end").strip()
                    
                    # Check if this is an image entry (has corresponding image path)
                    if i < len(enforcement_images) and enforcement_images[i].get():
                        # This is an image with description
                        image_path = enforcement_images[i].get()
                        
                        # Add the image to the document
                        try:
                            # Insert image first with width fitting document and height of 400
                            doc.add_picture(image_path, width=Inches(6.5), height=Inches(250/72))  # 6.5 inches width (standard doc width), 400px height converted to inches
                            
                            # Add numbered description after the image
                            if content:
                                p = doc.add_paragraph("")
                                p.paragraph_format.left_indent = Inches(0.25)
                                p.add_run(f"{note_counter}. ").bold = True
                                p.add_run(content)
                            else:
                                # If no description, still add the number
                                p = doc.add_paragraph("")
                                p.paragraph_format.left_indent = Inches(0.25)
                                p.add_run(f"{note_counter}. ").bold = True
                            
                            has_enforcement_notes = True
                            note_counter += 1
                        except Exception as e:
                            # If image fails to load, just add the text
                            if content:
                                add_indented_paragraph(note_counter, content)
                                has_enforcement_notes = True
                                note_counter += 1
                    else:
                        # Regular text note
                        if content:
                            has_enforcement_notes = True
                            add_indented_paragraph(note_counter, content)
                            note_counter +=  1
                

                # Add an empty numbered note if no content
                if not has_enforcement_notes:
                    add_indented_paragraph(note_counter, "")
                    note_counter += 1

                add_bold_section_header("Alumni Park")
                
                # Add alumni park notes with continuing counter
                has_alumni_notes = False
                for box in alumni_boxes:
                    content = box.get("1.0", "end").strip()
                    if content:
                        has_alumni_notes = True
                        add_indented_paragraph(note_counter, content)
                        note_counter += 1
                
                # Add an empty numbered note if no content
                if not has_alumni_notes:
                    add_indented_paragraph(note_counter, "")
                    note_counter += 1

                add_bold_section_header("Goodspeed Family Pier")
                
                # Add pier notes with continuing counter
                has_pier_notes = False
                for box in pier_boxes:
                    content = box.get("1.0", "end").strip()
                    if content:
                        has_pier_notes = True
                        add_indented_paragraph(note_counter, content)
                        note_counter += 1
                
                # Add an empty numbered note if no content
                if not has_pier_notes:
                    add_indented_paragraph(note_counter, "")
                    note_counter += 1

            # === Dining Service & Markets ===
            add_bold_section_header("Dining Service & Markets")
            
            # Add dining notes with continuing counter
            has_dining_notes = False
            for box in dining_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_dining_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_dining_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # === Hotel ===
            add_bold_section_header("Hotel")
            
            # Add hotel notes with continuing counter
            has_hotel_notes = False
            for box in hotel_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_hotel_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_hotel_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # === Miscellaneous ===
            add_bold_section_header("Miscellaneous")
            
            # Add miscellaneous notes with continuing counter
            has_misc_notes = False
            for box in misc_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_misc_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_misc_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # === Security Section === Only for Memorial Union and Union South
            add_bold_section_header("Security")  # Changed from "CSC Log"

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            # Bold the headers
            hdr_cells[0].paragraphs[0].add_run("Shift").bold = True
            hdr_cells[1].paragraphs[0].add_run("Staff Requested").bold = True
            hdr_cells[2].paragraphs[0].add_run("Staff Present").bold = True
            
            # Center-align header cells
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1  # 1 = CENTER

            for shift in csc_shifts:
                req_val = csc_entries[shift]["requested"].get().strip()
                pres_val = csc_entries[shift]["present"].get().strip()
                names_val = csc_entries[shift]["names"].get().strip()

                # Format names in parentheses only if present is a number
                pres_display = pres_val
                if pres_val.isdigit() and names_val:
                    pres_display = f"{pres_val} ({names_val})"

                row = table.add_row().cells
                
                # Bold the shift name - no tab character needed
                shift_run = row[0].paragraphs[0].add_run(shift)
                shift_run.bold = True
                
                # Set text directly
                row[1].text = req_val if req_val else "-"
                row[2].text = pres_display if pres_display else "-"
                
                # Center-align all cells in this row
                for cell in row:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 1  # 1 = CENTER

        # Excel Tally Update - Only for non-Red Gym buildings
        if building != "Red Gym":
            try:
                # Extract year and date for folder and filename
                user_date = entries["date"].get()
                parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
                current_year = parsed_date.strftime("%Y")
                current_month = parsed_date.strftime("%B")
                # Build master tag set from all tab tag lists
                master_tag_set = set()
                for taglist in [MECHANICAL_TAG_OPTIONS, PRODUCTION_TAG_OPTIONS, PATRON_TAG_OPTIONS, ACCESS_TAG_OPTIONS, CASH_TAG_OPTIONS, DINING_TAG_OPTIONS, TERRACE_TAG_OPTIONS]:
                    master_tag_set.update(tag for tag in taglist if tag != "None")
                # Add special-case tags
                master_tag_set.add("Hotel Request")
                master_tag_set.add("Carding Support/Lead Carding")
                master_tag_set.add("Decibel checked")
                master_tag_set.add("Patron Services/inquires/General Assistance")
                categories = sorted(master_tag_set)
                months = [
                    "January", "February", "March", "April", "May", "June",
                    "July", "August", "September", "October", "November", "December"
                ]
                # Create year folder in building-specific directory
                base_dir = f"M:\\Sh_BM\\{building}\\Night Reports"
                year_dir = os.path.join(base_dir, current_year)
                os.makedirs(year_dir, exist_ok=True)
                # Save tally as building_Tally_YYYY.xlsx in year folder
                building_short = "MU" if building == "Memorial Union" else "US"
                tally_filename = f"{building_short}_Tally_{current_year}.xlsx"
                tally_path = os.path.join(year_dir, tally_filename)
                # Load or create the Excel file for the current year/building
                if os.path.exists(tally_path):
                    df = pd.read_excel(tally_path, index_col=0)
                    for category in categories:
                        if category not in df.index:
                            df.loc[category] = [0] * len(df.columns)
                    # Remove any rows not in master list
                    for row in list(df.index):
                        if row not in categories:
                            df = df.drop(row)
                else:
                    df = pd.DataFrame(0, index=categories, columns=months)
                # Tally tags for all notes
                tag_counts = {cat: 0 for cat in categories}
                # Mechanical
                for tag_var_list, box in zip(mechanical_note_tags, mechanical_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Production
                for tag_var_list, box in zip(production_note_tags, production_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Patron
                for tag_var_list, box in zip(patron_note_tags, patron_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    if tags:
                        for tag in tags:
                            tag_counts[tag] += 1
                    else:
                        # If no tag but text, count as General Assistance
                        tag_counts["Patron Services/inquires/General Assistance"] += 1
                # Access
                for tag_var_list, box in zip(access_note_tags, access_note_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Cash
                for tag_var_list, box in zip(cash_note_tags, cash_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Dining
                for tag_var_list, box in zip(dining_note_tags, dining_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Memorial Union specific sections
                if building == "Memorial Union":
                    # Enforcement
                    for tag_var_list, box in zip(enforcement_note_tags, enforcement_boxes):
                        content = box.get("1.0", "end").strip()
                        if not content:
                            continue
                        tags = set(var.get() for var in tag_var_list if var.get() != "None")
                        for tag in tags:
                            tag_counts[tag] += 1
                    # Alumni
                    for tag_var_list, box in zip(alumni_note_tags, alumni_boxes):
                        content = box.get("1.0", "end").strip()
                        if not content:
                            continue
                        tags = set(var.get() for var in tag_var_list if var.get() != "None")
                        for tag in tags:
                            tag_counts[tag] += 1
                    # Pier
                    for tag_var_list, box in zip(pier_note_tags, pier_boxes):
                        content = box.get("1.0", "end").strip()
                        if not content:
                            continue
                        tags = set(var.get() for var in tag_var_list if var.get() != "None")
                        for tag in tags:
                            tag_counts[tag] += 1
                # Hotel (special case)
                for box in hotel_boxes:
                    content = box.get("1.0", "end").strip()
                    if content:
                        tag_counts["Hotel Request"] +=   1
                # Decibel (special case)
                if any(time.get().strip() and time.get().strip() != "Time" for time, _, _ in decibel_entries):
                    tag_counts["Decibel checked"] += 1
                # Write tallies to Excel
                for category, count in tag_counts.items():
                    if category in df.index and current_month in df.columns:
                        df.loc[category, current_month] += count
                df = df.loc[categories]  # Ensure order
                df.to_excel(tally_path)
            except Exception as e:
                messagebox.showerror("Excel Error", f"Failed to update Excel tally: {e}")
        else:
            # Red Gym Excel Tally Update
            try:
                # Extract year and date for folder and filename
                user_date = entries["date"].get()
                parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
                current_year = parsed_date.strftime("%Y")
                current_month = parsed_date.strftime("%B")
                # Build Red Gym tag set
                red_gym_tag_set = set()
                for taglist in [RED_GYM_MISC_TAG_OPTIONS]:
                    red_gym_tag_set.update(tag for tag in taglist if tag != "None")
                categories = sorted(red_gym_tag_set)
                months = [
                    "January", "February", "March", "April", "May", "June",
                    "July", "August", "September", "October", "November", "December"
                ]
                # Create year folder in Red Gym directory
                base_dir = f"M:\\Sh_BM\\{building}\\Night Reports"
                year_dir = os.path.join(base_dir, current_year)
                os.makedirs(year_dir, exist_ok=True)
                # Save tally as RG_Tally_YYYY.xlsx in year folder
                tally_filename = f"RG_Tally_{current_year}.xlsx"
                tally_path = os.path.join(year_dir, tally_filename)
                # Load or create the Excel file for the current year/building
                if os.path.exists(tally_path):
                    df = pd.read_excel(tally_path, index_col=0)
                    for category in categories:
                        if category not in df.index:
                            df.loc[category] = [0] * len(df.columns)
                    # Remove any rows not in master list
                    for row in list(df.index):
                        if row not in categories:
                            df = df.drop(row)
                else:
                    df = pd.DataFrame(0, index=categories, columns=months)
                # Tally tags for Red Gym misc notes
                tag_counts = {cat: 0 for cat in categories}
                # Red Gym Misc
                for tag_var_list, box in zip(misc_note_tags, red_gym_misc_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Write tallies to Excel
                for category, count in tag_counts.items():
                    if category in df.index and current_month in df.columns:
                        df.loc[category, current_month] += count
                df = df.loc[categories]  # Ensure order
                df.to_excel(tally_path)
            except Exception as e:
                messagebox.showerror("Excel Error", f"Failed to update Red Gym Excel tally: {e}")

        # Save report as MM-DD-YY.docx in the correct folder (for all building types)
        user_date = entries["date"].get()
        parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
        current_year = parsed_date.strftime("%Y")
        current_month = parsed_date.strftime("%B")
        
        # Create year and month folders in building-specific directory
        base_dir = f"M:\\Sh_BM\\{building}\\Night Reports"
        year_dir = os.path.join(base_dir, current_year)
        month_dir = os.path.join(year_dir, current_month)
        os.makedirs(month_dir, exist_ok=True)
        
        report_filename = f"{parsed_date.month}-{parsed_date.day}-{str(parsed_date.year)[2:]}.docx"
        report_path = os.path.join(month_dir, report_filename)
        doc.save(report_path)
        
        messagebox.showinfo("Success", f"Report saved as {report_path}")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# Wrap the initial UI setup in a function to be called after mainloop starts
def show_startup_modal():
    """Show initial modal to choose between creating new report or loading saved report"""
    startup_window = tk.Toplevel()
    startup_window.title("Night Report Generator")
    startup_window.configure(bg="black")
    startup_window.resizable(False, False)
    
    # Center the window on the screen
    window_width = 450
    window_height = 300
    screen_width = startup_window.winfo_screenwidth()
    screen_height = startup_window.winfo_screenheight()
    x = int((screen_width / 2) - (window_width / 2))
    y = int((screen_height / 2) - (window_height / 2))
    startup_window.geometry(f"{window_width}x{window_height}+{x}+{y}")
    
    # Make it a modal dialog but don't make it transient since root is withdrawn
    # startup_window.transient(root)  # Commented out - causes issues when root is withdrawn on macOS
    startup_window.grab_set()
    startup_window.lift()  # Bring window to front
    startup_window.focus_force()  # Force focus to the window
    
    # macOS specific: Force window to be topmost temporarily
    startup_window.attributes('-topmost', True)
    startup_window.update()
    startup_window.attributes('-topmost', False)
    
    # Handle window close event - exit application if startup modal is closed
    def on_startup_close():
        stop_autosave()  # Stop any running autosave
        root.quit()  # Exit the mainloop
        root.destroy()  # Destroy the root window
    
    startup_window.protocol("WM_DELETE_WINDOW", on_startup_close)
    
    # Title
    title_label = tk.Label(
        startup_window,
        text="NIGHT REPORT GENERATOR",
        font=("Helvetica", 18, "bold"),
        fg="white",
        bg="black"
    )
    title_label.pack(pady=(30, 20))
    
    # Subtitle
    subtitle_label = tk.Label(
        startup_window,
        text="Choose an option to get started:",
        font=("Helvetica", 12),
        fg="white",
        bg="black"
    )
    subtitle_label.pack(pady=(0, 30))
    
    # Button frame
    button_frame = tk.Frame(startup_window, bg="black")
    button_frame.pack(pady=20)
    
    def create_new_report():
        startup_window.destroy()
        select_building()
    
    def load_saved_report():
        startup_window.destroy()
        load_draft_report_startup()
    
    # Create New Report button
    new_report_btn = tk.Button(
        button_frame,
        text="Create New Report",
        command=create_new_report,
        bg="white",
        fg="black",
        font=("Helvetica", 14, "bold"),
        padx=20,
        pady=10,
        width=18
    )
    new_report_btn.pack(pady=10)
    
    # Load Saved Report button
    load_report_btn = tk.Button(
        button_frame,
        text="Load Saved Report",
        command=load_saved_report,
        bg="white",
        fg="black",
        font=("Helvetica", 14, "bold"),
        padx=20,
        pady=10,
        width=18
    )
    load_report_btn.pack(pady=10)
    
    # Instructions
    instructions_label = tk.Label(
        startup_window,
        text="• Create New Report: Start a fresh report for your shift\n• Load Saved Report: Continue working on a previously saved draft",
        font=("Helvetica", 10),
        fg="gray",
        bg="black",
        justify=tk.LEFT
    )
    instructions_label.pack(pady=(20, 10))
    
    # Wait for this window to be destroyed before proceeding
    root.wait_window(startup_window)

def load_draft_report_startup():
    """Special version of load_draft_report for startup flow"""
    try:
        # Get the file path from user
        file_path = filedialog.askopenfilename(
            title="Select Draft Report to Load",
            filetypes=[
                ("JSON files", "*.json"),
                ("All files", "*.*")
            ],
            initialdir=os.path.expanduser("~/Desktop")  # Start at Desktop for startup
        )
        
        if not file_path:
            # User cancelled - show the startup modal again
            show_startup_modal()
            return
        
        # Load and validate the JSON data
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Validate the data structure with backward compatibility
        if not isinstance(data, dict):
            raise ValueError("Invalid file format: not a JSON object")
        
        if "building" not in data:
            raise ValueError("Invalid draft file: missing 'building' field")
        
        # ENHANCED: Create missing 'entries' section with sensible defaults
        if "entries" not in data:
            data["entries"] = {}
        
        # ENHANCED: Ensure 'notes' section exists with sensible defaults
        if "notes" not in data:
            data["notes"] = {}
        
        # Set the global building variable from the loaded data
        global building, loaded_from_draft
        building = data["building"]
        loaded_from_draft = True
        
        # Update window title with loaded building
        root.title(f"{building} - Night Report Generator")
        
        # Configure tabs for the loaded building
        configure_tabs_for_building()
        
        # Populate the form with the loaded data
        populate_form_from_data(data)
        
        # Show success message
        messagebox.showinfo("Success", "Report loaded successfully!")
        
    except FileNotFoundError:
        messagebox.showerror("Load Error", "The selected draft file could not be found.")
        show_startup_modal()
    except json.JSONDecodeError as e:
        messagebox.showerror("Load Error", f"Invalid JSON file: {str(e)}")
        show_startup_modal()
    except ValueError as e:
        messagebox.showerror("Load Error", str(e))
        show_startup_modal()
    except Exception as e:
        messagebox.showerror("Load Error", str(e))
        show_startup_modal()

def start_app():
    show_startup_modal()

# Schedule the app start after the mainloop starts
root.after(0, start_app)
root.mainloop()
