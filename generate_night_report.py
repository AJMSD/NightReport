import tkinter as tk
from tkinter import ttk, messagebox, font
from docx import Document
from datetime import datetime
import pandas as pd
import os
from openpyxl import load_workbook

def configure_text_box(textbox, min_height=4):
    """Configure a text box to automatically resize based on content"""
    def update_height(event=None):
        # Get text content
        text = textbox.get("1.0", "end-1c")  # Exclude the automatic trailing newline
        
        # Get the width of the text box in pixels
        width = textbox.winfo_width()
        
        # If width is not properly initialized yet, use approximate calculation
        if width <= 1:
            if hasattr(textbox, 'text_font'):
                width = textbox.cget("width") * textbox.text_font.measure("0")
            else:
                # Default width approximation
                width = textbox.cget("width") * 7
        else:
            # Subtract some pixels for padding
            width = width - 10
        
        # Calculate the number of wrapped lines
        line_count = 0
        for line in text.split('\n'):
            if line.strip() == "":  # Empty line
                line_count += 1
                continue
                
            # Estimate line wrapping based on average character width
            if hasattr(textbox, 'text_font'):
                char_width = textbox.text_font.measure("0")  # Average character width
            else:
                char_width = 7  # Default approximation
            
            # Calculate how many lines this text would wrap to
            if char_width > 0:
                chars_per_line = max(1, width // char_width)
                wrapped_lines = (len(line) // chars_per_line) + 1
                line_count += wrapped_lines
            else:
                line_count += 1
        
        # Add 1 for the cursor line and any pending text
        line_count += 1
        
        # Set minimum height
        new_height = max(min_height, line_count)
        
        # Update height
        textbox.configure(height=new_height)
    
    # Create a font object to measure text
    try:
        font_name = textbox.cget("font")
        if isinstance(font_name, str):
            textbox.text_font = font.Font(font=font_name)
    except Exception:
        # If there's any error with font, we'll use approximations instead
        textbox.text_font = font.Font(family="Helvetica", size=11)
    
    # Bind events that should trigger height update
    textbox.bind("<KeyRelease>", update_height)
    textbox.bind("<FocusOut>", update_height)
    textbox.bind("<Configure>", update_height)  # Also update when the widget is resized
    
    # Initial height update (after a short delay to ensure the widget is rendered)
    textbox.after(100, update_height)
    
    return textbox

root = tk.Tk()
root.title("Night Report Generator")
root.configure(bg="black")
root.geometry("750x700")

label_font = ("Helvetica", 11)
entry_bg = "black"
entry_fg = "white"
entry_font = ("Helvetica", 11)
entries = {}
building_traffic_boxes = []

# === Tabs ===
notebook = ttk.Notebook(root)
notebook.pack(expand=1, fill="both")

style = ttk.Style()
style.theme_use('default')
style.configure('TNotebook', background='black', borderwidth=0)
style.configure('TNotebook.Tab', background='black', foreground='white', padding=10)
style.map('TNotebook.Tab', background=[('selected', '#333')], foreground=[('selected', 'white')])

def create_tab(title):
    frame = tk.Frame(notebook, bg="black")
    notebook.add(frame, text=title)
    return frame

def add_labeled_entry(parent, label_text, key, default=""):
    frame = tk.Frame(parent, bg="black")
    label = tk.Label(frame, text=label_text, fg="white", bg="black", font=label_font)
    entry = tk.Entry(frame, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, width=60)
    entry.insert(0, default)
    label.pack(anchor="w")
    entry.pack(fill="x")
    frame.pack(pady=5, padx=10, fill="x")
    entries[key] = entry

# === Create tabs ===
tabs = {
    "Supervisor Info": create_tab("Supervisor Info"),
    "Building Traffic": create_tab("Building Traffic"),
    "Mechanical": create_tab("Mechanical"),
    "Production": create_tab("Production"),
    "Patron Services": create_tab("Patron Services"),
    "Access": create_tab("Access/Lock/Unlock"),
    "Cash Office": create_tab("Cash Office"),
    "Carding Runs": create_tab("Carding Runs"),
    "Terrace Traffic": create_tab("Terrace Traffic"),
    "Terrace Enforcement": create_tab("Terrace Enforcement"),
    "Alumni Park": create_tab("Alumni Park"),
    "Goodspeed Pier": create_tab("Goodspeed Family Pier"),
    "Dining & Markets": create_tab("Dining Service & Markets"),
    "Hotel": create_tab("Hotel"),
    "Misc": create_tab("Miscellaneous"),
    "CSC Log": create_tab("CSC Log")
}

# === Supervisor Info tab ===
today_str = datetime.now().strftime("%A, %B %d, %Y")
add_labeled_entry(tabs["Supervisor Info"], "Date", "date", default=today_str)
add_labeled_entry(tabs["Supervisor Info"], "Shift Hours", "shift_hours")
add_labeled_entry(tabs["Supervisor Info"], "Building Manager(s)", "bms")
add_labeled_entry(tabs["Supervisor Info"], "Guest Service Specialist", "gss")
add_labeled_entry(tabs["Supervisor Info"], "Custodial Supervisor(s)", "custodial")
add_labeled_entry(tabs["Supervisor Info"], "Production Supervisor(s)", "production")
add_labeled_entry(tabs["Supervisor Info"], "Retail & Dining Supervisor(s)", "retail")
add_labeled_entry(tabs["Supervisor Info"], "Catering Supervisor(s)", "catering")
add_labeled_entry(tabs["Supervisor Info"], "Event Manager(s)", "eventmanagers")
add_labeled_entry(tabs["Supervisor Info"], "CAVR Desk Staff", "cavr")

# === Building Traffic Tab ===
# Create a container frame inside the tab to hold text boxes
traffic_tab = tabs["Building Traffic"]
traffic_notes_frame = tk.Frame(traffic_tab, bg="black")
traffic_notes_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_building_traffic_box(default_text=""):
    frame = tk.Frame(traffic_notes_frame, bg="black")
    label = tk.Label(frame, text=f"Building Traffic Note #{len(building_traffic_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    # Configure text box to auto-resize
    configure_text_box(textbox)
    building_traffic_boxes.append(textbox)

# Add first required box
add_building_traffic_box()

# Add the + Add Note button, stays at bottom
add_box_btn = tk.Button(
    traffic_tab, text="+ Add Note", command=add_building_traffic_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
)
add_box_btn.pack(pady=10)

# === Mechanical/Repairs/Custodial Tab ===
mechanical_boxes = []

mechanical_tab = tabs["Mechanical"]
mechanical_notes_frame = tk.Frame(mechanical_tab, bg="black")
mechanical_notes_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_mechanical_box(default_text=""):
    frame = tk.Frame(mechanical_notes_frame, bg="black")
    label = tk.Label(frame, text=f"Mechanical Note #{len(mechanical_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    mechanical_boxes.append(textbox)

# Add first required box
add_mechanical_box()

# Add Note button
add_mechanical_btn = tk.Button(
    mechanical_tab, text="+ Add Note", command=add_mechanical_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
)
add_mechanical_btn.pack(pady=10)

# === Production Services Tab ===
production_tab = tabs["Production"]
production_notes = []
decibel_entries = []

# Frame for general notes
production_notes_frame = tk.Frame(production_tab, bg="black")
production_notes_frame.pack(fill="x", padx=10, pady=(10, 5))

def add_production_note_box(default_text=""):
    frame = tk.Frame(production_notes_frame, bg="black")
    label = tk.Label(frame, text=f"Production Note #{len(production_notes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    production_notes.append(textbox)

# Add first required note
add_production_note_box()

# Button to add more production notes
add_note_btn = tk.Button(
    production_tab, text="+ Add Note", command=add_production_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
)
add_note_btn.pack(pady=(0, 10))

# === Decibel Reading Table ===
decibel_frame = tk.Frame(production_tab, bg="black")
decibel_frame.pack(padx=10, pady=(10, 5), fill="x")

header = tk.Label(decibel_frame, text="Decibel Readings", fg="white", bg="black", font=("Helvetica", 12, "bold"))
header.pack(anchor="w")

# Container for decibel rows - this will help with proper placement
decibel_rows_container = tk.Frame(decibel_frame, bg="black")
decibel_rows_container.pack(fill="x", expand=True)

def add_decibel_row():
    row_frame = tk.Frame(decibel_rows_container, bg="black")
    time_entry = tk.Entry(row_frame, width=15, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
    reading_entry = tk.Entry(row_frame, width=15, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
    location_entry = tk.Entry(row_frame, width=40, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)

    time_entry.insert(0, "Time")
    reading_entry.insert(0, "Reading (db)")
    location_entry.insert(0, "Location")

    time_entry.pack(side="left", padx=5)
    reading_entry.pack(side="left", padx=5)
    location_entry.pack(side="left", padx=5)
    row_frame.pack(pady=3, anchor="w", fill="x")
    
    decibel_entries.append((time_entry, reading_entry, location_entry))

# Add first decibel row
add_decibel_row()

# Button to add more decibel rows - kept outside the container
add_decibel_btn = tk.Button(
    decibel_frame, text="+ Add Decibel Reading", command=add_decibel_row,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
)
add_decibel_btn.pack(pady=10)

# === Patron Services Tab ===
patron_tab = tabs["Patron Services"]
patron_boxes = []

patron_notes_frame = tk.Frame(patron_tab, bg="black")
patron_notes_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_patron_note_box(default_text=""):
    frame = tk.Frame(patron_notes_frame, bg="black")
    label = tk.Label(frame, text=f"Patron Note #{len(patron_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=6, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=6)
    patron_boxes.append(textbox)

# Add first required box
add_patron_note_box()

# Add button to add more notes
add_patron_btn = tk.Button(
    patron_tab, text="+ Add Note", command=add_patron_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
)
add_patron_btn.pack(pady=10)

# === Access/Lock/Unlock Tab ===
access_tab = tabs["Access"]
access_inputs = {}
access_note_boxes = []

access_frame = tk.Frame(access_tab, bg="black")
access_frame.pack(fill="both", expand=True, padx=10, pady=10)

def add_dropdown(label_text, key, options):
    frame = tk.Frame(access_frame, bg="black")
    label = tk.Label(frame, text=label_text, fg="white", bg="black", font=label_font)
    var = tk.StringVar(value=options[0])
    dropdown = ttk.Combobox(frame, textvariable=var, values=options, state="readonly", width=20)
    label.pack(anchor="w")
    dropdown.pack(fill="x")
    frame.pack(pady=5, fill="x")
    access_inputs[key] = var

def add_entry(label_text, key):
    frame = tk.Frame(access_frame, bg="black")
    label = tk.Label(frame, text=label_text, fg="white", bg="black", font=label_font)
    entry = tk.Entry(frame, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, width=30)
    label.pack(anchor="w")
    entry.pack(fill="x")
    frame.pack(pady=5, fill="x")
    access_inputs[key] = entry

add_dropdown("Loading Dock Arm Gate at Early Check:", "early_gate", ["Open", "Closed"])
add_entry("Time of Early Check:", "early_time")

add_dropdown("Loading Dock Arm Gate at Closing Check:", "close_gate", ["Open", "Closed"])
add_entry("Time of Closing Check:", "close_time")

add_dropdown("HID Scanners Status at Close:", "hid_status", ["Locked", "Unlocked"])
add_dropdown("Overhead Door Secured:", "door_status", ["Successfully", "Unsuccessfully"])

# === Optional Access Notes ===
access_notes_label = tk.Label(access_frame, text="Additional Access Notes (Optional):", fg="white", bg="black", font=label_font)
access_notes_label.pack(anchor="w", pady=(10, 0))

access_notes_container = tk.Frame(access_frame, bg="black")
access_notes_container.pack(fill="both", expand=True)

def add_access_note():
    frame = tk.Frame(access_notes_container, bg="black")
    label = tk.Label(frame, text=f"Access Note #{len(access_note_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    label.pack(anchor="w")
    textbox.pack(fill="x", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    access_note_boxes.append(textbox)

# Add Note button
add_note_btn = tk.Button(
    access_frame, text="+ Add Note", command=add_access_note,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
)
add_note_btn.pack(pady=10)

# === Cash Office Tab ===
cash_tab = tabs["Cash Office"]
cash_boxes = []

cash_frame = tk.Frame(cash_tab, bg="black")
cash_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_cash_note_box(default_text=""):
    frame = tk.Frame(cash_frame, bg="black")
    label = tk.Label(frame, text=f"Cash Office Note #{len(cash_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    cash_boxes.append(textbox)

add_cash_note_box()

add_cash_btn = tk.Button(
    cash_tab, text="+ Add Note", command=add_cash_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
)
add_cash_btn.pack(pady=10)

# === Carding Runs Tab ===
carding_tab = tabs["Carding Runs"]
carding_boxes = []

carding_frame = tk.Frame(carding_tab, bg="black")
carding_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_carding_note_box(default_text=""):
    frame = tk.Frame(carding_frame, bg="black")
    label = tk.Label(frame, text=f"Carding Run Note #{len(carding_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    carding_boxes.append(textbox)

# Add first required note
add_carding_note_box()

# Add note button
add_carding_btn = tk.Button(
    carding_tab, text="+ Add Note", command=add_carding_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
)
add_carding_btn.pack(pady=10)

# === Terrace Traffic Tab ===
terrace_tab = tabs["Terrace Traffic"]
terrace_boxes = []

terrace_frame = tk.Frame(terrace_tab, bg="black")
terrace_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_terrace_note_box(default_text=""):
    frame = tk.Frame(terrace_frame, bg="black")
    label = tk.Label(frame, text=f"Terrace Traffic Note #{len(terrace_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    terrace_boxes.append(textbox)

add_terrace_note_box()

add_terrace_btn = tk.Button(
    terrace_tab, text="+ Add Note", command=add_terrace_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
)
add_terrace_btn.pack(pady=10)

# === Terrace Enforcement Tab ===
enforcement_tab = tabs["Terrace Enforcement"]
enforcement_boxes = []

enforcement_frame = tk.Frame(enforcement_tab, bg="black")
enforcement_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_enforcement_note_box(default_text=""):
    frame = tk.Frame(enforcement_frame, bg="black")
    label = tk.Label(frame, text=f"Enforcement Note #{len(enforcement_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    enforcement_boxes.append(textbox)

add_enforcement_note_box()

add_enforcement_btn = tk.Button(
    enforcement_tab, text="+ Add Note", command=add_enforcement_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
)
add_enforcement_btn.pack(pady=10)

# === Alumni Park Tab ===
alumni_tab = tabs["Alumni Park"]
alumni_boxes = []

alumni_frame = tk.Frame(alumni_tab, bg="black")
alumni_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_alumni_note_box(default_text=""):
    frame = tk.Frame(alumni_frame, bg="black")
    label = tk.Label(frame, text=f"Alumni Park Note #{len(alumni_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    alumni_boxes.append(textbox)

add_alumni_note_box()

tk.Button(
    alumni_tab, text="+ Add Note", command=add_alumni_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
).pack(pady=10)

# === Goodspeed Pier Tab ===
pier_tab = tabs["Goodspeed Pier"]
pier_boxes = []

pier_frame = tk.Frame(pier_tab, bg="black")
pier_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_pier_note_box(default_text=""):
    frame = tk.Frame(pier_frame, bg="black")
    label = tk.Label(frame, text=f"Pier Note #{len(pier_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    pier_boxes.append(textbox)

add_pier_note_box()

tk.Button(
    pier_tab, text="+ Add Note", command=add_pier_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
).pack(pady=10)

# === Dining & Markets Tab ===
dining_tab = tabs["Dining & Markets"]
dining_boxes = []

dining_frame = tk.Frame(dining_tab, bg="black")
dining_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_dining_note_box(default_text=""):
    frame = tk.Frame(dining_frame, bg="black")
    label = tk.Label(frame, text=f"Dining Note #{len(dining_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    dining_boxes.append(textbox)

add_dining_note_box()

tk.Button(
    dining_tab, text="+ Add Note", command=add_dining_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
).pack(pady=10)

# === Hotel Tab ===
hotel_tab = tabs["Hotel"]
hotel_boxes = []

hotel_frame = tk.Frame(hotel_tab, bg="black")
hotel_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_hotel_note_box(default_text=""):
    frame = tk.Frame(hotel_frame, bg="black")
    label = tk.Label(frame, text=f"Hotel Note #{len(hotel_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox, min_height=3)
    hotel_boxes.append(textbox)

add_hotel_note_box()

tk.Button(
    hotel_tab, text="+ Add Note", command=add_hotel_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
).pack(pady=10)

# === Miscellaneous Tab ===
misc_tab = tabs["Misc"]
misc_boxes = []

misc_frame = tk.Frame(misc_tab, bg="black")
misc_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

def add_misc_note_box(default_text=""):
    frame = tk.Frame(misc_frame, bg="black")
    label = tk.Label(frame, text=f"Misc Note #{len(misc_boxes)+1}:", fg="white", bg="black", font=label_font)
    textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
    if default_text:
        textbox.insert("1.0", default_text)
    label.pack(anchor="w")
    textbox.pack(fill="both", expand=True, padx=5)
    frame.pack(pady=5, fill="x")
    configure_text_box(textbox)
    misc_boxes.append(textbox)

add_misc_note_box()

tk.Button(
    misc_tab, text="+ Add Note", command=add_misc_note_box,
    bg="white", fg="black", font=("Helvetica", 10, "bold")
).pack(pady=10)

# === CSC Log Tab ===
csc_tab = tabs["CSC Log"]
csc_entries = {}

csc_shifts = ["Morning", "Evening", "Special Event", "Chair Watch"]

csc_frame = tk.Frame(csc_tab, bg="black")
csc_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

for shift in csc_shifts:
    section = tk.LabelFrame(csc_frame, text=shift, fg="white", bg="black", font=label_font, labelanchor="n", padx=10, pady=10)
    section.pack(fill="x", pady=5)

    # Requested
    req_label = tk.Label(section, text="Staff Requested:", fg="white", bg="black", font=label_font)
    req_entry = tk.Entry(section, width=10, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
    req_label.grid(row=0, column=0, sticky="w")
    req_entry.grid(row=0, column=1, padx=5)

    # Present
    pres_label = tk.Label(section, text="Staff Present:", fg="white", bg="black", font=label_font)
    pres_entry = tk.Entry(section, width=10, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
    pres_label.grid(row=0, column=2, sticky="w")
    pres_entry.grid(row=0, column=3, padx=5)

    # Names (optional)
    names_label = tk.Label(section, text="Names (comma-separated):", fg="white", bg="black", font=label_font)
    names_entry = tk.Entry(section, width=40, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
    names_label.grid(row=1, column=0, columnspan=2, sticky="w")
    names_entry.grid(row=1, column=2, columnspan=2, sticky="we", pady=5)

    csc_entries[shift] = {
        "requested": req_entry,
        "present": pres_entry,
        "names": names_entry
    }

# === Generate Report Logic ===
def generate_report():
    try:
        doc = Document()
        doc.add_heading('MEMORIAL UNION\nBUILDING MANAGER\'S NIGHT REPORT', level=1)
        doc.add_paragraph(f'Date: {entries["date"].get()}')
        doc.add_paragraph(f'Shift Hours: {entries["shift_hours"].get()}')
        doc.add_paragraph(f'Building Manager(s): {entries["bms"].get()}')
        doc.add_paragraph(f'Event Manager(s): {entries["eventmanagers"].get()}')
        doc.add_paragraph(f'Guest Service Specialist: {entries["gss"].get()}')
        doc.add_paragraph(f'Custodial Supervisor(s): {entries["custodial"].get()}')
        doc.add_paragraph(f'Production Supervisor(s): {entries["production"].get()}')
        doc.add_paragraph(f'Retail & Dining Supervisor(s): {entries["retail"].get()}')
        doc.add_paragraph(f'Catering Supervisor(s): {entries["catering"].get()}')
        doc.add_paragraph(f'CAVR Desk Staff: {entries["cavr"].get()}')

        doc.add_paragraph("\nNotes:")
        doc.add_paragraph("Building Traffic")

        for i, box in enumerate(building_traffic_boxes, start=1):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")

        doc.add_paragraph("Mechanical/Repairs/Custodial")

        start_index = len(building_traffic_boxes) + 1
        for i, box in enumerate(mechanical_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")

        doc.add_paragraph("Production Services (Meetings, Events, Set-ups, AV)")

        start_index = len(building_traffic_boxes) + len(mechanical_boxes) + 1
        for i, box in enumerate(production_notes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")

        if decibel_entries:
            doc.add_paragraph("Decibel Readings:")
            # Create a table for decibel readings
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Time"
            header_cells[1].text = "Reading (dB)"
            header_cells[2].text = "Location"
            
            # Add data rows
            for time_entry, reading_entry, location_entry in decibel_entries:
                time = time_entry.get().strip()
                reading = reading_entry.get().strip()
                location = location_entry.get().strip()
                if time and reading and location and time != "Time" and reading != "Reading (db)" and location != "Location":
                    row_cells = table.add_row().cells
                    row_cells[0].text = time
                    row_cells[1].text = reading
                    row_cells[2].text = location

        doc.add_paragraph("Patron Services (Membership, Patron Assistance, Problem Patrons)")

        start_index = len(building_traffic_boxes) + len(mechanical_boxes) + len(production_notes) + 1
        for i, box in enumerate(patron_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")

        doc.add_paragraph("Access/Lock/Unlock")

        start_index = (
            len(building_traffic_boxes)
            + len(mechanical_boxes)
            + len(production_notes)
            + len(patron_boxes)
            + 1
        )

        access_notes = [
            f"At the early check, the loading dock arm gate was {access_inputs['early_gate'].get().lower()} at {access_inputs['early_time'].get()}.",
            f"At the closing check, the loading dock arm gate was {access_inputs['close_gate'].get().lower()} at {access_inputs['close_time'].get()}.",
            f"At the closing check, the HID scanners were {access_inputs['hid_status'].get().lower()}.",
            f"I {access_inputs['door_status'].get().lower()} secured the loading dock overhead door for the night."
        ]

        # Auto-generated sentences
        for i, sentence in enumerate(access_notes, start=start_index):
            doc.add_paragraph(f"{i}. {sentence}")

        # User-entered access notes
        for j, box in enumerate(access_note_boxes):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{start_index + len(access_notes) + j}. {content}")

        doc.add_paragraph("Cash Office")

        start_index += len(access_notes) + len(access_note_boxes)
        for i, box in enumerate(cash_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")
        start_index += len(cash_boxes)

        doc.add_paragraph("Carding Runs")

        for i, box in enumerate(carding_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")
        start_index += len(carding_boxes)

        doc.add_paragraph("Terrace Traffic")
        for i, box in enumerate(terrace_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")
        start_index += len(terrace_boxes)

        doc.add_paragraph("Terrace Enforcement")
        for i, box in enumerate(enforcement_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")
        start_index += len(enforcement_boxes)

        # === Alumni Park ===
        doc.add_paragraph("Alumni Park")
        for i, box in enumerate(alumni_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")
        start_index += len(alumni_boxes)

        # === Goodspeed Family Pier ===
        doc.add_paragraph("Goodspeed Family Pier")
        for i, box in enumerate(pier_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")
        start_index += len(pier_boxes)

        # === Dining Service & Markets ===
        doc.add_paragraph("Dining Service & Markets")
        for i, box in enumerate(dining_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")
        start_index += len(dining_boxes)

        # === Hotel ===
        doc.add_paragraph("Hotel")
        for i, box in enumerate(hotel_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")
        start_index += len(hotel_boxes)

        # === Miscellaneous ===
        doc.add_paragraph("Miscellaneous")
        for i, box in enumerate(misc_boxes, start=start_index):
            content = box.get("1.0", "end").strip()
            if content:
                doc.add_paragraph(f"{i}. {content}")
        start_index += len(misc_boxes)

        # === CSC Log Section ===
        doc.add_paragraph("CSC Log")

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Shift"
        hdr_cells[1].text = "Staff Requested"
        hdr_cells[2].text = "Staff Present"

        for shift in csc_shifts:
            req_val = csc_entries[shift]["requested"].get().strip()
            pres_val = csc_entries[shift]["present"].get().strip()
            names_val = csc_entries[shift]["names"].get().strip()

            # Format names in parentheses only if present is a number
            pres_display = pres_val
            if pres_val.isdigit() and names_val:
                pres_display = f"{pres_val} ({names_val})"

            row = table.add_row().cells
            row[0].text = shift
            row[1].text = req_val if req_val else "-"
            row[2].text = pres_display if pres_display else "-"

        filename = f"NightReport_{datetime.now().strftime('%Y-%m-%d_%H%M')}.docx"
        doc.save(filename)
        # === Excel Tally Update ===
        try:
            tally_file = "NightReport_Tally.xlsx"
            categories = [
                "Building Traffic", "Mechanical/Repairs/Custodial", "Production Services",
                "Patron Services", "Access/Lock/Unlock", "Cash Office", "Carding Runs",
                "Terrace Traffic", "Terrace Enforcement", "Alumni Park", "Goodspeed Family Pier",
                "Dining Service & Markets", "Hotel", "Miscellaneous"
            ]
            months = [
                "January", "February", "March", "April", "May", "June",
                "July", "August", "September", "October", "November", "December"
            ]

            # Load or create the Excel file
            if os.path.exists(tally_file):
                df = pd.read_excel(tally_file, index_col=0)
            else:
                df = pd.DataFrame(0, index=categories, columns=months)

            # Extract month from the date field
            user_date = entries["date"].get()
            parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
            current_month = parsed_date.strftime("%B")

            # Count notes
            counts = {
                "Building Traffic": sum(1 for box in building_traffic_boxes if box.get("1.0", "end").strip()),
                "Mechanical/Repairs/Custodial": sum(1 for box in mechanical_boxes if box.get("1.0", "end").strip()),
                "Production Services": sum(1 for box in production_notes if box.get("1.0", "end").strip()),
                "Patron Services": sum(1 for box in patron_boxes if box.get("1.0", "end").strip()),
                "Access/Lock/Unlock": 1 + sum(1 for box in access_note_boxes if box.get("1.0", "end").strip()),  # Always at least 1 dropdown result
                "Cash Office": sum(1 for box in cash_boxes if box.get("1.0", "end").strip()),
                "Carding Runs": sum(1 for box in carding_boxes if box.get("1.0", "end").strip()),
                "Terrace Traffic": sum(1 for box in terrace_boxes if box.get("1.0", "end").strip()),
                "Terrace Enforcement": sum(1 for box in enforcement_boxes if box.get("1.0", "end").strip()),
                "Alumni Park": sum(1 for box in alumni_boxes if box.get("1.0", "end").strip()),
                "Goodspeed Family Pier": sum(1 for box in pier_boxes if box.get("1.0", "end").strip()),
                "Dining Service & Markets": sum(1 for box in dining_boxes if box.get("1.0", "end").strip()),
                "Hotel": sum(1 for box in hotel_boxes if box.get("1.0", "end").strip()),
                "Miscellaneous": sum(1 for box in misc_boxes if box.get("1.0", "end").strip())
            }

            for category, count in counts.items():
                if category in df.index and current_month in df.columns:
                    df.loc[category, current_month] += count
            print("Debug: Building Traffic Box Content Count")
            for box in building_traffic_boxes:
                content = box.get("1.0", "end").strip()
                print(f"Content: '{content}'")

            df.to_excel(tally_file)

        except Exception as e:
            messagebox.showerror("Excel Error", f"Failed to update Excel tally: {e}")
        messagebox.showinfo("Success", f"Report saved as {filename}")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# === Submit Button ===
submit_btn = tk.Button(
    root, text="Generate Report", command=generate_report,
    bg="white", fg="black", font=("Helvetica", 12, "bold"), padx=10, pady=6
)
submit_btn.pack(pady=10)

root.mainloop()
