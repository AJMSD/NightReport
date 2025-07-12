import tkinter as tk
from tkinter import ttk, messagebox, font, filedialog
from docx import Document
from docx.shared import Inches
from datetime import datetime
import pandas as pd
import os
from openpyxl import load_workbook

# === Building Selection Function ===
def select_building():
    # Ensure the main window is mapped so dialogs can show
    root.deiconify()
    building_window = tk.Toplevel()
    building_window.title("Select Building")
    building_window.configure(bg="black")
    # Center the window on the screen
    window_width = 400
    window_height = 350  # Increased height for Red Gym option
    screen_width = building_window.winfo_screenwidth()
    screen_height = building_window.winfo_screenheight()
    x = int((screen_width / 2) - (window_width / 2))
    y = int((screen_height / 2) - (window_height / 2))
    building_window.geometry(f"{window_width}x{window_height}+{x}+{y}")
    
    label = tk.Label(
        building_window, 
        text="SELECT BUILDING", 
        font=("Helvetica", 16, "bold"),
        fg="white", 
        bg="black"
    )
    label.pack(pady=20)
    
    selected_building = tk.StringVar()
    
    def confirm_selection():
        if not selected_building.get():
            messagebox.showwarning("Selection Required", "Please select a building before proceeding.")
            return
        
        # Set the global building variable
        global building
        building = selected_building.get()
        
        # Close the window
        building_window.destroy()
        
        # Now configure tabs based on selection
        configure_tabs_for_building()
    
    # Memorial Union Option
    mem_frame = tk.Frame(building_window, bg="black", bd=2, relief=tk.RIDGE, padx=10, pady=10)
    mem_frame.pack(fill=tk.X, padx=20, pady=5)
    
    mem_radio = tk.Radiobutton(
        mem_frame,
        text="Memorial Union",
        variable=selected_building,
        value="Memorial Union",
        bg="black",
        fg="white",
        selectcolor="black",
        font=("Helvetica", 12),
        activebackground="black",
        activeforeground="white"
    )
    mem_radio.pack(anchor=tk.W)
    
    # Union South Option
    south_frame = tk.Frame(building_window, bg="black", bd=2, relief=tk.RIDGE, padx=10, pady=10)
    south_frame.pack(fill=tk.X, padx=20, pady=5)
    
    south_radio = tk.Radiobutton(
        south_frame,
        text="Union South",
        variable=selected_building,
        value="Union South",
        bg="black",
        fg="white",
        selectcolor="black",
        font=("Helvetica", 12),
        activebackground="black",
        activeforeground="white"
    )
    south_radio.pack(anchor=tk.W)
    
    # Red Gym Option
    red_gym_frame = tk.Frame(building_window, bg="black", bd=2, relief=tk.RIDGE, padx=10, pady=10)
    red_gym_frame.pack(fill=tk.X, padx=20, pady=5)
    
    red_gym_radio = tk.Radiobutton(
        red_gym_frame,
        text="Red Gym",
        variable=selected_building,
        value="Red Gym",
        bg="black",
        fg="white",
        selectcolor="black",
        font=("Helvetica", 12),
        activebackground="black",
        activeforeground="white"
    )
    red_gym_radio.pack(anchor=tk.W)
    
    # Confirm Button
    confirm_btn = tk.Button(
        building_window,
        text="Confirm Selection",
        command=confirm_selection,
        bg="white",
        fg="black",
        font=("Helvetica", 11, "bold"),
        padx=10, 
        pady=5
    )
    confirm_btn.pack(pady=15)
    
    # Wait for this window to be destroyed before proceeding
    root.wait_window(building_window)

def configure_text_box(textbox, min_height=4):
    """Configure a text box to automatically resize based on content"""
    def update_height(event=None):
        # Get text content
        text = textbox.get("1.0", "end-1c")  # Exclude the automatic trailing newline
        
        # Get the width of the text box in pixels
        width = textbox.winfo_width()
        
        # If width is not properly initialized yet, use approximate calculation
        if width <= 1:
            if hasattr(textbox, 'text_font'):
                width = textbox.cget("width") * textbox.text_font.measure("0")
            else:
                # Default width approximation
                width = textbox.cget("width") * 7
        else:
            # Subtract some pixels for padding
            width = width - 10
        
        # Calculate the number of wrapped lines
        line_count = 0
        for line in text.split('\n'):
            if line.strip() == "":  # Empty line
                line_count += 1
                continue
                
            # Estimate line wrapping based on average character width
            if hasattr(textbox, 'text_font'):
                char_width = textbox.text_font.measure("0")  # Average character width
            else:
                char_width = 7  # Default approximation
            
            # Calculate how many lines this text would wrap to
            if char_width > 0:
                chars_per_line = max(1, width // char_width)
                wrapped_lines = (len(line) // chars_per_line) + 1
                line_count += wrapped_lines
            else:
                line_count += 1
        
        # Add 1 for the cursor line and any pending text
        line_count += 1
        
        # Set minimum height
        new_height = max(min_height, line_count)
        
        # Update height
        textbox.configure(height=new_height)
    
    # Create a font object to measure text
    try:
        font_name = textbox.cget("font")
        if isinstance(font_name, str):
            textbox.text_font = font.Font(font=font_name)
    except Exception:
        # If there's any error with font, we'll use approximations instead
        textbox.text_font = font.Font(family="Helvetica", size=11)
    
    # Bind events that should trigger height update
    textbox.bind("<KeyRelease>", update_height)
    textbox.bind("<FocusOut>", update_height)
    textbox.bind("<Configure>", update_height)  # Also update when the widget is resized
    
    # Initial height update (after a short delay to ensure the widget is rendered)
    textbox.after(100, update_height)
    
    return textbox

root = tk.Tk()
root.title("Night Report Generator")
root.configure(bg="black")
root.state('zoomed')  # Open in full screen
# root.geometry("750x700")  # Remove fixed geometry
root.withdraw()  # Hide the main window until building is selected

# Global building variable
building = ""

label_font = ("Helvetica", 11)
entry_bg = "black"
entry_fg = "white"
entry_font = ("Helvetica", 11)
entries = {}
building_traffic_boxes = []

# Add these global lists for all note sections
mechanical_boxes = []
mechanical_note_tags = []
production_boxes = []
production_note_tags = []
decibel_entries = []
patron_boxes = []
patron_note_tags = []
patron_emergency_flags = []
access_note_boxes = []
access_note_tags = []
cash_boxes = []
cash_note_tags = []
dining_boxes = []
dining_note_tags = []
hotel_boxes = []
misc_boxes = []
misc_note_tags = []

# Global variables for Security/CSC section
csc_entries = {}
csc_shifts = ["Morning", "Evening", "Special Event", "Chair Watch"]

# Add these global variables for Red Gym after the existing global lists
red_gym_building_tours_box = None
red_gym_deviations_entry = None
red_gym_deviation_boxes = []
red_gym_door_check_time = None
red_gym_door_check_day_type = None
red_gym_mail_boxes = []
red_gym_misc_boxes = []

# Memorial Union-specific
carding_boxes = []
terrace_boxes = []
terrace_note_tags = []
enforcement_boxes = []
enforcement_note_tags = []
enforcement_images = []  # Add this for storing image paths
alumni_boxes = []
alumni_note_tags = []
pier_boxes = []
pier_note_tags = []

# Tag options by tab (global)
MECHANICAL_TAG_OPTIONS = ["None", "FAMIS report", "Door Lock Failure", "Reset Elevator", "Custodial Support"]
PRODUCTION_TAG_OPTIONS = ["None", "Production Support", "AV Troubleshooting for Guest/Patron"]
PATRON_TAG_OPTIONS = ["None", "Lost & Found support", "First Aid Kit", "Incident report", "No dog policy", "Confirmed service animal", "Ban patron encounter", "Repeat Offender Encounter", "Severe Disruption / Verbal Escalation", "Patron Subject to 24-Hour Ban", "UWPD/911 Called", "Non-Patron Escort / Trespass", "Ensure outside alcohol policy", "Enforce smoking policy", "Removed Bikes & electric transport", "Wellness Check", "Disability Assistance/Patron Support", "CESO Support/Wedding walkthrough", "Patron Services/inquires/General Assistance"]
ACCESS_TAG_OPTIONS = ["None", "Loading dock access", "Employee Locker unlock", "Health Room unlock", "Door/cooler unlock"]
CASH_TAG_OPTIONS = ["None", "Cash Equipment Jam", "Cash Machine Problem"]
TERRACE_TAG_OPTIONS = ["None", "No dog policy", "Ensure outside alcohol policy & remove outside alcohol", "Enforce smoking policy", "Removed Bikes & electric transport", "Enforce no fishing policy", "Ban patron encounter", "Repeat Offender Encounter", "Confirmed service animal", "Wellness Check"]
DINING_TAG_OPTIONS = ["None", "Catering Order Pickup/inquiry"]
RED_GYM_MISC_TAG_OPTIONS = ["None", "Physical Plant"]

# Helper for adding tag dropdowns to a note (must be defined before use)
def add_tagging_to_note(tag_frame, tag_options, tag_vars, tag_dropdowns):
    def add_tag_dropdown():
        var = tk.StringVar(value="None")
        dropdown = ttk.Combobox(tag_frame, textvariable=var, values=tag_options, state="readonly", width=30)
        dropdown.pack(side="left", padx=(0, 5))
        tag_vars.append(var)
        tag_dropdowns.append(dropdown)
        def on_tag_change(event=None):
            if var.get() != "None" and not hasattr(dropdown, 'add_tag_btn'):
                add_btn = tk.Button(tag_frame, text="+ Add Tag", bg="white", fg="black", font=("Helvetica", 9, "bold"),
                                   command=lambda: [add_tag_dropdown(), add_btn.destroy()])
                add_btn.pack(side="left", padx=(0, 5))
                dropdown.add_tag_btn = add_btn
            elif var.get() == "None" and hasattr(dropdown, 'add_tag_btn'):
                dropdown.add_tag_btn.destroy()
                delattr(dropdown, 'add_tag_btn')
        dropdown.bind("<<ComboboxSelected>>", on_tag_change)
    add_tag_dropdown()

# === Tabs ===
notebook = ttk.Notebook(root)
notebook.pack(expand=1, fill="both")

style = ttk.Style()
style.theme_use('default')
style.configure('TNotebook', background='black', borderwidth=0)
style.configure('TNotebook.Tab', background='black', foreground='white', padding=10)
style.map('TNotebook.Tab', background=[('selected', '#333')], foreground=[('selected', 'white')])

def create_tab(title):
    frame = tk.Frame(notebook, bg="black")
    notebook.add(frame, text=title)
    return frame

def add_labeled_entry(parent, label_text, key, default=""):
    frame = tk.Frame(parent, bg="black")
    label = tk.Label(frame, text=label_text, fg="white", bg="black", font=label_font)
    entry = tk.Entry(frame, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, width=60)
    entry.insert(0, default)
    label.pack(anchor="w")
    entry.pack(fill="x")
    frame.pack(pady=5, padx=10, fill="x")
    entries[key] = entry

# Function to configure tabs based on selected building
def configure_tabs_for_building():
    global tabs
    
    # Update window title with selected building
    root.title(f"{building} - Night Report Generator")
    
    # Define all possible tabs
    all_tab_keys = [
        "Supervisor Info", "Building Traffic", "Mechanical", "Production", 
        "Patron Services", "Access", "Cash Office", "Carding Runs", 
        "Terrace Traffic", "Terrace Enforcement", "Alumni Park", "Goodspeed Pier",
        "Dining & Markets", "Hotel", "Misc", "Security"
    ]
    
    # Define the tabs to exclude for different buildings
    exclude_for_union_south = [
        "Carding Runs", "Terrace Traffic", "Terrace Enforcement", 
        "Alumni Park", "Goodspeed Pier", "Mail"
    ]
    
    exclude_for_memorial_union = [
        "Mail"
    ]
    
    red_gym_tabs_only = [
        "Supervisor Info", "Building Traffic", "Security", "Mail", "Misc"
    ]
    
    # Determine which tabs to create
    if building == "Red Gym":
        tab_keys_to_create = red_gym_tabs_only
    elif building == "Union South":
        tab_keys_to_create = [key for key in all_tab_keys if key not in exclude_for_union_south]
    else:  # Memorial Union
        tab_keys_to_create = [key for key in all_tab_keys if key not in exclude_for_memorial_union]
    
    # Create tabs
    tabs = {}
    for key in tab_keys_to_create:
        tab_title = key
        if key == "Access":
            tab_title = "Access/Lock/Unlock"
        elif key == "Dining & Markets":
            tab_title = "Dining Service & Markets"
        elif key == "Goodspeed Pier":
            tab_title = "Goodspeed Family Pier"
            
        tabs[key] = create_tab(tab_title)
    
    # Continue with the rest of the UI setup
    setup_ui_components()
    
    # Now show the main window
    root.deiconify()

def setup_ui_components():
    # === Supervisor Info tab ===
    today_str = datetime.now().strftime("%A, %B %d, %Y")
    add_labeled_entry(tabs["Supervisor Info"], "Date", "date", default=today_str)
    add_labeled_entry(tabs["Supervisor Info"], "Shift Hours", "shift_hours")
    add_labeled_entry(tabs["Supervisor Info"], "Building Manager(s)", "bms")
    
    # Only add additional fields for Memorial Union and Union South
    if building != "Red Gym":
        # Terrace Manager(s) only for Memorial Union
        if building == "Memorial Union":
            add_labeled_entry(tabs["Supervisor Info"], "Terrace Manager(s)", "terrace_managers")
        add_labeled_entry(tabs["Supervisor Info"], "Guest Service Specialist", "gss")
        add_labeled_entry(tabs["Supervisor Info"], "Operation Manager(s)", "operation_managers")
        add_labeled_entry(tabs["Supervisor Info"], "Custodial Supervisor(s)", "custodial")
        add_labeled_entry(tabs["Supervisor Info"], "Production Supervisor(s)", "production")
        add_labeled_entry(tabs["Supervisor Info"], "Retail & Dining Supervisor(s)", "retail")
        add_labeled_entry(tabs["Supervisor Info"], "Catering Supervisor(s)", "catering")
        add_labeled_entry(tabs["Supervisor Info"], "Event Manager(s)", "eventmanagers")
        add_labeled_entry(tabs["Supervisor Info"], "CAVR Desk Staff", "cavr")

    # === Building Traffic Tab ===
    # Create a container frame inside the tab to hold text boxes
    traffic_tab = tabs["Building Traffic"]
    traffic_notes_frame = tk.Frame(traffic_tab, bg="black")
    traffic_notes_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    def add_building_traffic_box(default_text=""):
        frame = tk.Frame(traffic_notes_frame, bg="black")
        label = tk.Label(frame, text=f"Building Traffic Note #{len(building_traffic_boxes)+1}:", fg="white", bg="black", font=label_font)
        textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        if default_text:
            textbox.insert("1.0", default_text)
        label.pack(anchor="w")
        textbox.pack(fill="both", expand=True, padx=5)
        frame.pack(pady=5, fill="x")
        # Configure text box to auto-resize
        configure_text_box(textbox)
        building_traffic_boxes.append(textbox)

    # Add first required box
    add_building_traffic_box()

    # Add the + Add Note button, stays at bottom
    add_box_btn = tk.Button(
        traffic_tab, text="+ Add Note", command=add_building_traffic_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_box_btn.pack(pady=10)

    # Red Gym specific tabs
    if building == "Red Gym":
        # === Red Gym Security Tab ===
        global red_gym_building_tours_box, red_gym_deviations_entry, red_gym_deviation_boxes
        global red_gym_door_check_time, red_gym_door_check_day_type
        
        security_tab = tabs["Security"]
        security_frame = tk.Frame(security_tab, bg="black")
        security_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Building Tours
        tours_label = tk.Label(security_frame, text="Building Tours:", fg="white", bg="black", font=label_font)
        tours_label.pack(anchor="w")
        red_gym_building_tours_box = tk.Text(security_frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        red_gym_building_tours_box.pack(fill="x", padx=5, pady=(0, 10))
        configure_text_box(red_gym_building_tours_box, min_height=3)
        
        # Deviations section
        deviations_label = tk.Label(security_frame, text="Deviations from standard building locking protocol:", fg="white", bg="black", font=label_font)
        deviations_label.pack(anchor="w", pady=(10, 0))
        
        deviations_frame = tk.Frame(security_frame, bg="black")
        deviations_frame.pack(fill="x", pady=5)
        
        deviations_text_label = tk.Label(deviations_frame, text="There were", fg="white", bg="black", font=label_font)
        deviations_text_label.pack(side="left")
        
        red_gym_deviations_entry = tk.Entry(deviations_frame, width=5, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        red_gym_deviations_entry.insert(0, "0")
        red_gym_deviations_entry.pack(side="left", padx=(5, 5))
        
        deviations_text_label2 = tk.Label(deviations_frame, text="deviations from the standard building locking protocol today.", fg="white", bg="black", font=label_font)
        deviations_text_label2.pack(side="left")
        
        # Container for deviation notes
        deviations_notes_frame = tk.Frame(security_frame, bg="black")
        deviations_notes_frame.pack(fill="x", pady=5)
        
        def update_deviation_notes():
            # Clear existing notes
            for widget in deviations_notes_frame.winfo_children():
                widget.destroy()
            red_gym_deviation_boxes.clear()
            
            try:
                num_deviations = int(red_gym_deviations_entry.get() or "0")
                if num_deviations > 0:
                    for i in range(num_deviations):
                        frame = tk.Frame(deviations_notes_frame, bg="black")
                        label = tk.Label(frame, text=f"Deviation {chr(97+i)}:", fg="white", bg="black", font=label_font)
                        textbox = tk.Text(frame, height=2, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
                        label.pack(anchor="w")
                        textbox.pack(fill="x", padx=5)
                        frame.pack(pady=2, fill="x")
                        configure_text_box(textbox, min_height=2)
                        red_gym_deviation_boxes.append(textbox)
            except ValueError:
                pass
        
        red_gym_deviations_entry.bind('<KeyRelease>', lambda e: update_deviation_notes())
        
        # Door check section
        door_check_frame = tk.Frame(security_frame, bg="black")
        door_check_frame.pack(fill="x", pady=(20, 5))
        
        door_check_label1 = tk.Label(door_check_frame, text="I checked to confirm that the Building Doors were locked and the door swipe scanner was red at", fg="white", bg="black", font=label_font)
        door_check_label1.pack(side="left")
        
        red_gym_door_check_time = tk.Entry(door_check_frame, width=10, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        red_gym_door_check_time.pack(side="left", padx=(5, 5))
        
        door_check_label2 = tk.Label(door_check_frame, text="on a", fg="white", bg="black", font=label_font)
        door_check_label2.pack(side="left")
        
        red_gym_door_check_day_type = ttk.Combobox(door_check_frame, values=["weekday", "weekend"], state="readonly", width=10)
        red_gym_door_check_day_type.pack(side="left", padx=(5, 0))
        
        door_check_label3 = tk.Label(door_check_frame, text=".", fg="white", bg="black", font=label_font)
        door_check_label3.pack(side="left")
        
        # === Red Gym Mail Tab ===
        global red_gym_mail_boxes
        mail_tab = tabs["Mail"]
        mail_frame = tk.Frame(mail_tab, bg="black")
        mail_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))
        
        def add_red_gym_mail_box(default_text=""):
            frame = tk.Frame(mail_frame, bg="black")
            label = tk.Label(frame, text=f"Mail Note #{len(red_gym_mail_boxes)+1}:", fg="white", bg="black", font=label_font)
            textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
            if default_text:
                textbox.insert("1.0", default_text)
            label.pack(anchor="w")
            textbox.pack(fill="both", expand=True, padx=5)
            frame.pack(pady=5, fill="x")
            configure_text_box(textbox, min_height=3)
            red_gym_mail_boxes.append(textbox)
        
        add_red_gym_mail_box()
        
        tk.Button(
            mail_tab, text="+ Add Note", command=add_red_gym_mail_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        ).pack(pady=10)
        
        # === Red Gym Misc Tab ===
        global red_gym_misc_boxes
        misc_tab = tabs["Misc"]
        misc_frame = tk.Frame(misc_tab, bg="black")
        misc_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))
        
        def add_red_gym_misc_box(default_text=""):
            frame = tk.Frame(misc_frame, bg="black")
            label = tk.Label(frame, text=f"Misc Note #{len(red_gym_misc_boxes)+1}:", fg="white", bg="black", font=label_font)
            textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
            if default_text:
                textbox.insert("1.0", default_text)
            label.pack(anchor="w")
            textbox.pack(fill="both", expand=True, padx=5)
            frame.pack(pady=5, fill="x")
            configure_text_box(textbox, min_height=3)
            red_gym_misc_boxes.append(textbox)
            
            # Tagging for Red Gym misc (single dropdown only)
            tag_vars = []
            tag_dropdowns = []
            tag_frame = tk.Frame(frame, bg="black")
            tag_frame.pack(anchor="w", pady=(2, 0))
            
            # Add single dropdown without the ability to add more
            var = tk.StringVar(value="None")
            dropdown = ttk.Combobox(tag_frame, textvariable=var, values=RED_GYM_MISC_TAG_OPTIONS, state="readonly", width=30)
            dropdown.pack(side="left", padx=(0, 5))
            tag_vars.append(var)
            tag_dropdowns.append(dropdown)
            
            misc_note_tags.append(tag_vars)
            frame.pack(pady=5, fill="x")
        
        add_red_gym_misc_box()
        
        tk.Button(
            misc_tab, text="+ Add Note", command=add_red_gym_misc_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        ).pack(pady=10)
        
        # === Generate Report Button ===
        submit_btn = tk.Button(
            root, text="Generate Report", command=generate_report,
            bg="white", fg="black", font=("Helvetica", 12, "bold"), padx=10, pady=6
        )
        submit_btn.pack(pady=10)
        
        return  # Exit early for Red Gym
    
    # Continue with existing code for Memorial Union and Union South
    # === Mechanical/Repairs/Custodial Tab ===
    mechanical_notes_frame = tk.Frame(tabs["Mechanical"], bg="black")
    mechanical_notes_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    def add_mechanical_box(default_text=""):
        frame = tk.Frame(mechanical_notes_frame, bg="black")
        label = tk.Label(frame, text=f"Mechanical Note #{len(mechanical_boxes)+1}:", fg="white", bg="black", font=label_font)
        textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        if default_text:
            textbox.insert("1.0", default_text)
        label.pack(anchor="w")
        textbox.pack(fill="both", expand=True, padx=5)
        configure_text_box(textbox)
        mechanical_boxes.append(textbox)

        # Tagging logic
        tag_vars = []  # List of tk.StringVar for this note
        tag_dropdowns = []  # List of dropdown widgets for this note

        tag_frame = tk.Frame(frame, bg="black")
        tag_frame.pack(anchor="w", pady=(2, 0))

        def add_tag_dropdown():
            var = tk.StringVar(value="None")
            dropdown = ttk.Combobox(tag_frame, textvariable=var, values=MECHANICAL_TAG_OPTIONS, state="readonly", width=22)
            dropdown.pack(side="left", padx=(0, 5))
            tag_vars.append(var)
            tag_dropdowns.append(dropdown)

            def on_tag_change(event=None):
                # Show +Add Tag button if a valid tag is selected and no button exists
                if var.get() != "None" and not hasattr(dropdown, 'add_tag_btn'):
                    add_btn = tk.Button(tag_frame, text="+ Add Tag", bg="white", fg="black", font=("Helvetica", 9, "bold"),
                                       command=lambda: [add_tag_dropdown(), add_btn.destroy()])
                    add_btn.pack(side="left", padx=(0, 5))
                    dropdown.add_tag_btn = add_btn
                elif var.get() == "None" and hasattr(dropdown, 'add_tag_btn'):
                    dropdown.add_tag_btn.destroy()
                    delattr(dropdown, 'add_tag_btn')
            dropdown.bind("<<ComboboxSelected>>", on_tag_change)

        add_tag_dropdown()  # Add the first dropdown
        mechanical_note_tags.append(tag_vars)
        frame.pack(pady=5, fill="x")

    # Add first required box
    add_mechanical_box()

    # Add Note button
    add_mechanical_btn = tk.Button(
        tabs["Mechanical"], text="+ Add Note", command=add_mechanical_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_mechanical_btn.pack(pady=10)

    # === Production Services Tab ===
    production_notes_frame = tk.Frame(tabs["Production"], bg="black")
    production_notes_frame.pack(fill="x", padx=10, pady=(10, 5))

    def add_production_note_box(default_text=""):
        frame = tk.Frame(production_notes_frame, bg="black")
        label = tk.Label(frame, text=f"Production Note #{len(production_boxes)+1}:", fg="white", bg="black", font=label_font)
        textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        if default_text:
            textbox.insert("1.0", default_text)
        label.pack(anchor="w")
        textbox.pack(fill="both", expand=True, padx=5)
        frame.pack(pady=5, fill="x")
        configure_text_box(textbox)
        production_boxes.append(textbox)
        # Tagging
        tag_vars = []
        tag_dropdowns = []
        tag_frame = tk.Frame(frame, bg="black")
        tag_frame.pack(anchor="w", pady=(2, 0))
        add_tagging_to_note(tag_frame, PRODUCTION_TAG_OPTIONS, tag_vars, tag_dropdowns)
        production_note_tags.append(tag_vars)
        frame.pack(pady=5, fill="x")

    # Add first required note
    add_production_note_box()

    # Button to add more production notes
    add_note_btn = tk.Button(
        tabs["Production"], text="+ Add Note", command=add_production_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_note_btn.pack(pady=(0, 10))

    # === Decibel Reading Table ===
    decibel_frame = tk.Frame(tabs["Production"], bg="black")
    decibel_frame.pack(padx=10, pady=(10, 5), fill="x")

    header = tk.Label(decibel_frame, text="Decibel Readings", fg="white", bg="black", font=("Helvetica", 12, "bold"))
    header.pack(anchor="w")

    # Container for decibel rows - this will help with proper placement
    decibel_rows_container = tk.Frame(decibel_frame, bg="black")
    decibel_rows_container.pack(fill="x", expand=True)

    def add_decibel_row():
        row_frame = tk.Frame(decibel_rows_container, bg="black")
        time_entry = tk.Entry(row_frame, width=15, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        reading_entry = tk.Entry(row_frame, width=15, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        location_entry = tk.Entry(row_frame, width=40, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)

        time_entry.insert(0, "Time")
        reading_entry.insert(0, "Reading (db)")
        location_entry.insert(0, "Location")

        # Add focus event handlers to select all text when clicked
        def on_focus_in(event):
            event.widget.select_range(0, tk.END)
            
        time_entry.bind("<FocusIn>", on_focus_in)
        reading_entry.bind("<FocusIn>", on_focus_in)
        location_entry.bind("<FocusIn>", on_focus_in)

        time_entry.pack(side="left", padx=5)
        reading_entry.pack(side="left", padx=5)
        location_entry.pack(side="left", padx=5)
        row_frame.pack(pady=3, anchor="w", fill="x")
        
        decibel_entries.append((time_entry, reading_entry, location_entry))

    # Add first decibel row
    add_decibel_row()

    # Button to add more decibel rows - kept outside the container
    add_decibel_btn = tk.Button(
        decibel_frame, text="+ Add Decibel Reading", command=add_decibel_row,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_decibel_btn.pack(pady=10)

    # === Patron Services Tab ===
    patron_notes_frame = tk.Frame(tabs["Patron Services"], bg="black")
    patron_notes_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    def add_patron_note_box(default_text=""):
        frame = tk.Frame(patron_notes_frame, bg="black")
        label = tk.Label(frame, text=f"Patron Note #{len(patron_boxes)+1}:", fg="white", bg="black", font=label_font)
        textbox = tk.Text(frame, height=6, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        if default_text:
            textbox.insert("1.0", default_text)
        label.pack(anchor="w")
        textbox.pack(fill="both", expand=True, padx=5)
        configure_text_box(textbox, min_height=6)
        patron_boxes.append(textbox)
        # Tagging
        tag_vars = []
        tag_dropdowns = []
        tag_frame = tk.Frame(frame, bg="black")
        tag_frame.pack(anchor="w", pady=(2, 0))
        add_tagging_to_note(tag_frame, PATRON_TAG_OPTIONS, tag_vars, tag_dropdowns)
        patron_note_tags.append(tag_vars)
        frame.pack(pady=5, fill="x")

    # Add first required box
    add_patron_note_box()

    # Add button to add more notes
    add_patron_btn = tk.Button(
        tabs["Patron Services"], text="+ Add Note", command=add_patron_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_patron_btn.pack(pady=10)

    # === Access/Lock/Unlock Tab ===
    global access_inputs, csc_shifts, csc_entries  # <-- This line ensures both are global
    access_tab = tabs["Access"]
    access_frame = tk.Frame(access_tab, bg="black")
    access_frame.pack(fill="both", expand=True, padx=10, pady=10)
    
    access_inputs = {}
    access_note_boxes = []

    def add_dropdown(label_text, key, options):
        frame = tk.Frame(access_frame, bg="black")
        label = tk.Label(frame, text=label_text, fg="white", bg="black", font=label_font)
        var = tk.StringVar(value=options[0])
        dropdown = ttk.Combobox(frame, textvariable=var, values=options, state="readonly", width=20)
        label.pack(anchor="w")
        dropdown.pack(fill="x")
        frame.pack(pady=5, fill="x")
        access_inputs[key] = var

    def add_entry(label_text, key):
        frame = tk.Frame(access_frame, bg="black")
        label = tk.Label(frame, text=label_text, fg="white", bg="black", font=label_font)
        entry = tk.Entry(frame, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, width=30)
        label.pack(anchor="w")
        entry.pack(fill="x")
        frame.pack(pady=5, fill="x")
        access_inputs[key] = entry

    add_dropdown("Loading Dock Arm Gate at Early Check:", "early_gate", ["Open", "Closed"])
    add_entry("Time of Early Check:", "early_time")

    add_dropdown("Loading Dock Arm Gate at Closing Check:", "close_gate", ["Open", "Closed"])
    add_entry("Time of Closing Check:", "close_time")

    add_dropdown("HID Scanners Status at Close:", "hid_status", ["Locked", "Unlocked"])
    
    # Set default based on building type
    if building == "Memorial Union":
        add_dropdown("Overhead Door Secured:", "door_status", ["Unsuccessfully", "Successfully"])
    else:  # Union South and other buildings
        add_dropdown("Overhead Door Secured:", "door_status", ["Successfully", "Unsuccessfully"])

    # === Optional Access Notes ===
    access_notes_label = tk.Label(access_frame, text="Additional Access Notes (Optional):", fg="white", bg="black", font=label_font)
    access_notes_label.pack(anchor="w", pady=(10, 0))

    access_notes_container = tk.Frame(access_frame, bg="black")
    access_notes_container.pack(fill="both", expand=True)

    def add_access_note():
        frame = tk.Frame(access_notes_container, bg="black")
        label = tk.Label(frame, text=f"Access Note #{len(access_note_boxes)+1}:", fg="white", bg="black", font=label_font)
        textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        label.pack(anchor="w")
        textbox.pack(fill="x", expand=True, padx=5)
        frame.pack(pady=5, fill="x")
        configure_text_box(textbox, min_height=3)
        access_note_boxes.append(textbox)
        # Tagging
        tag_vars = []
        tag_dropdowns = []
        tag_frame = tk.Frame(frame, bg="black")
        tag_frame.pack(anchor="w", pady=(2, 0))
        add_tagging_to_note(tag_frame, ACCESS_TAG_OPTIONS, tag_vars, tag_dropdowns)
        access_note_tags.append(tag_vars)
        frame.pack(pady=5, fill="x")

    # Add Note button
    add_note_btn = tk.Button(
        access_frame, text="+ Add Note", command=add_access_note,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_note_btn.pack(pady=10)

    # === Cash Office Tab ===
    cash_frame = tk.Frame(tabs["Cash Office"], bg="black")
    cash_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    def add_cash_note_box(default_text=""):
        frame = tk.Frame(cash_frame, bg="black")
        label = tk.Label(frame, text=f"Cash Office Note #{len(cash_boxes)+1}:", fg="white", bg="black", font=label_font)
        textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        if default_text:
            textbox.insert("1.0", default_text)
        label.pack(anchor="w")
        textbox.pack(fill="both", expand=True, padx=5)
        frame.pack(pady=5, fill="x")
        configure_text_box(textbox)
        cash_boxes.append(textbox)
        # Tagging
        tag_vars = []
        tag_dropdowns = []
        tag_frame = tk.Frame(frame, bg="black")
        tag_frame.pack(anchor="w", pady=(2, 0))
        add_tagging_to_note(tag_frame, CASH_TAG_OPTIONS, tag_vars, tag_dropdowns)
        cash_note_tags.append(tag_vars)
        frame.pack(pady=5, fill="x")

    add_cash_note_box()

    add_cash_btn = tk.Button(
        tabs["Cash Office"], text="+ Add Note", command=add_cash_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    )
    add_cash_btn.pack(pady=10)

    # Create the Memorial Union-specific tabs only if the selected building is Memorial Union
    if building == "Memorial Union":
        # === Carding Runs Tab ===
        carding_frame = tk.Frame(tabs["Carding Runs"], bg="black")
        carding_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        def add_carding_note_box(default_text=""):
            frame = tk.Frame(carding_frame, bg="black")
            label = tk.Label(frame, text=f"Carding Run Note #{len(carding_boxes)+1}:", fg="white", bg="black", font=label_font)
            textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
            if default_text:
                textbox.insert("1.0", default_text)
            label.pack(anchor="w")
            textbox.pack(fill="both", expand=True, padx=5)
            frame.pack(pady=5, fill="x")
            configure_text_box(textbox, min_height=3)
            carding_boxes.append(textbox)

        # Add first required note
        add_carding_note_box()

        # Add note button
        add_carding_btn = tk.Button(
            tabs["Carding Runs"], text="+ Add Note", command=add_carding_note_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        )
        add_carding_btn.pack(pady=10)

        # === Terrace Traffic Tab ===
        terrace_frame = tk.Frame(tabs["Terrace Traffic"], bg="black")
        terrace_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        def add_terrace_note_box(default_text=""):
            frame = tk.Frame(terrace_frame, bg="black")
            label = tk.Label(frame, text=f"Terrace Traffic Note #{len(terrace_boxes)+1}:", fg="white", bg="black", font=label_font)
            textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
            if default_text:
                textbox.insert("1.0", default_text)
            label.pack(anchor="w")
            textbox.pack(fill="both", expand=True, padx=5)
            frame.pack(pady=5, fill="x")
            configure_text_box(textbox)
            terrace_boxes.append(textbox)
            # Tagging
            tag_vars = []
            tag_dropdowns = []
            tag_frame = tk.Frame(frame, bg="black")
            tag_frame.pack(anchor="w", pady=(2, 0))
            add_tagging_to_note(tag_frame, TERRACE_TAG_OPTIONS, tag_vars, tag_dropdowns)
            terrace_note_tags.append(tag_vars)
            frame.pack(pady=5, fill="x")
        add_terrace_note_box()

        add_terrace_btn = tk.Button(
            tabs["Terrace Traffic"], text="+ Add Note", command=add_terrace_note_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        )
        add_terrace_btn.pack(pady=10)

        # === Terrace Enforcement Tab ===
        enforcement_frame = tk.Frame(tabs["Terrace Enforcement"], bg="black")
        enforcement_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        # Keep track of enforcement components for proper ordering
        enforcement_components = []

        def add_enforcement_image():
            image_frame = tk.Frame(enforcement_frame, bg="black")
            
            # Label for the image section
            image_label = tk.Label(image_frame, text=f"Enforcement Image #{len(enforcement_images)+1}:", fg="white", bg="black", font=label_font)
            image_label.pack(anchor="w")
            
            # Frame for image upload button and status
            upload_frame = tk.Frame(image_frame, bg="black")
            upload_frame.pack(fill="x", pady=2)
            
            # Variable to store image path
            image_path_var = tk.StringVar()
            enforcement_images.insert(0, image_path_var)  # Insert at beginning
            
            # Status label
            status_label = tk.Label(upload_frame, text="No image selected", fg="gray", bg="black", font=entry_font)
            status_label.pack(side="left", padx=(0, 10))
            
            def select_image():
                file_path = filedialog.askopenfilename(
                    title="Select Image",
                    filetypes=[
                        ("Image files", "*.jpg *.jpeg *.png *.gif *.bmp"),
                        ("All files", "*.*")
                    ]
                )
                if file_path:
                    image_path_var.set(file_path)
                    filename = file_path.split('/')[-1]  # Get just the filename
                    status_label.config(text=f"Selected: {filename}", fg="white")
            
            upload_btn = tk.Button(upload_frame, text="Select Image", command=select_image,
                                 bg="white", fg="black", font=("Helvetica", 9, "bold"))
            upload_btn.pack(side="left")
            
            # Text box for image description
            desc_label = tk.Label(image_frame, text="Description of enforcement action:", fg="white", bg="black", font=label_font)
            desc_label.pack(anchor="w", pady=(5, 0))
            
            textbox = tk.Text(image_frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
            textbox.pack(fill="x", padx=5)
            configure_text_box(textbox, min_height=3)
            enforcement_boxes.insert(0, textbox)  # Insert at beginning
            
            # Tagging
            tag_vars = []
            tag_dropdowns = []
            tag_frame = tk.Frame(image_frame, bg="black")
            tag_frame.pack(anchor="w", pady=(2, 0))
            add_tagging_to_note(tag_frame, TERRACE_TAG_OPTIONS, tag_vars, tag_dropdowns)
            enforcement_note_tags.insert(0, tag_vars)  # Insert at beginning
            
            # Pack the image frame and reorder all components
            enforcement_components.insert(0, image_frame)
            reorder_enforcement_components()

        def add_enforcement_note_box(default_text=""):
            frame = tk.Frame(enforcement_frame, bg="black")
            label = tk.Label(frame, text=f"Enforcement Note #{len(enforcement_boxes)+1}:", fg="white", bg="black", font=label_font)
            textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
            if default_text:
                textbox.insert("1.0", default_text)
            label.pack(anchor="w")
            textbox.pack(fill="both", expand=True, padx=5)
            configure_text_box(textbox)
            enforcement_boxes.append(textbox)
            # Tagging
            tag_vars = []
            tag_dropdowns = []
            tag_frame = tk.Frame(frame, bg="black")
            tag_frame.pack(anchor="w", pady=(2, 0))
            add_tagging_to_note(tag_frame, TERRACE_TAG_OPTIONS, tag_vars, tag_dropdowns)
            enforcement_note_tags.append(tag_vars)
            
            # Add to components and reorder
            enforcement_components.append(frame)
            reorder_enforcement_components()

        def reorder_enforcement_components():
            # Forget all components
            for component in enforcement_components:
                component.pack_forget()
            # Re-pack in correct order
            for component in enforcement_components:
                component.pack(fill="x", pady=5)

        # Button to add enforcement image with description
        add_image_btn = tk.Button(
            enforcement_frame, text="+ Add Enforcement Image", command=add_enforcement_image,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        )
        add_image_btn.pack(pady=5)

        add_enforcement_note_box()

        add_enforcement_btn = tk.Button(
            tabs["Terrace Enforcement"], text="+ Add Note", command=add_enforcement_note_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        )
        add_enforcement_btn.pack(pady=10)

        # === Alumni Park Tab ===
        alumni_frame = tk.Frame(tabs["Alumni Park"], bg="black")
        alumni_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        def add_alumni_note_box(default_text=""):
            frame = tk.Frame(alumni_frame, bg="black")
            label = tk.Label(frame, text=f"Alumni Park Note #{len(alumni_boxes)+1}:", fg="white", bg="black", font=label_font)
            textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
            if default_text:
                textbox.insert("1.0", default_text)
            label.pack(anchor="w")
            textbox.pack(fill="both", expand=True, padx=5)
            frame.pack(pady=5, fill="x")
            configure_text_box(textbox, min_height=3)
            alumni_boxes.append(textbox)
            # Tagging
            tag_vars = []
            tag_dropdowns = []
            tag_frame = tk.Frame(frame, bg="black")
            tag_frame.pack(anchor="w", pady=(2, 0))
            add_tagging_to_note(tag_frame, TERRACE_TAG_OPTIONS, tag_vars, tag_dropdowns)
            alumni_note_tags.append(tag_vars)
            frame.pack(pady=5, fill="x")
        add_alumni_note_box()

        tk.Button(
            tabs["Alumni Park"], text="+ Add Note", command=add_alumni_note_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        ).pack(pady=10)

        # === Goodspeed Pier Tab ===
        pier_frame = tk.Frame(tabs["Goodspeed Pier"], bg="black")
        pier_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        def add_pier_note_box(default_text=""):
            frame = tk.Frame(pier_frame, bg="black")
            label = tk.Label(frame, text=f"Pier Note #{len(pier_boxes)+1}:", fg="white", bg="black", font=label_font)
            textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
            if default_text:
                textbox.insert("1.0", default_text)
            label.pack(anchor="w")
            textbox.pack(fill="both", expand=True, padx=5)
            frame.pack(pady=5, fill="x")
            configure_text_box(textbox, min_height=3)
            pier_boxes.append(textbox)
            # Tagging
            tag_vars = []
            tag_dropdowns = []
            tag_frame = tk.Frame(frame, bg="black")
            tag_frame.pack(anchor="w", pady=(2, 0))
            add_tagging_to_note(tag_frame, TERRACE_TAG_OPTIONS, tag_vars, tag_dropdowns)
            pier_note_tags.append(tag_vars)
            frame.pack(pady=5, fill="x")
        add_pier_note_box()

        tk.Button(
            tabs["Goodspeed Pier"], text="+ Add Note", command=add_pier_note_box,
            bg="white", fg="black", font=("Helvetica", 10, "bold")
        ).pack(pady=10)

    # === Dining & Markets Tab ===
    dining_frame = tk.Frame(tabs["Dining & Markets"], bg="black")
    dining_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    def add_dining_note_box(default_text=""):
        frame = tk.Frame(dining_frame, bg="black")
        label = tk.Label(frame, text=f"Dining Note #{len(dining_boxes)+1}:", fg="white", bg="black", font=label_font)
        textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        if default_text:
            textbox.insert("1.0", default_text)
        label.pack(anchor="w")
        textbox.pack(fill="both", expand=True, padx=5)
        frame.pack(pady=5, fill="x")
        configure_text_box(textbox)
        dining_boxes.append(textbox)
        # Tagging
        tag_vars = []
        tag_dropdowns = []
        tag_frame = tk.Frame(frame, bg="black")
        tag_frame.pack(anchor="w", pady=(2, 0))
        add_tagging_to_note(tag_frame, DINING_TAG_OPTIONS, tag_vars, tag_dropdowns)
        dining_note_tags.append(tag_vars)
        frame.pack(pady=5, fill="x")
    add_dining_note_box()

    tk.Button(
        tabs["Dining & Markets"], text="+ Add Note", command=add_dining_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    ).pack(pady=10)

    # === Hotel Tab ===
    hotel_frame = tk.Frame(tabs["Hotel"], bg="black")
    hotel_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    def add_hotel_note_box(default_text=""):
        frame = tk.Frame(hotel_frame, bg="black")
        label = tk.Label(frame, text=f"Hotel Note #{len(hotel_boxes)+1}:", fg="white", bg="black", font=label_font)
        textbox = tk.Text(frame, height=3, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        if default_text:
            textbox.insert("1.0", default_text)
        label.pack(anchor="w")
        textbox.pack(fill="both", expand=True, padx=5)
        configure_text_box(textbox, min_height=3)
        hotel_boxes.append(textbox)
        frame.pack(pady=5, fill="x")
    add_hotel_note_box()
    tk.Button(
        tabs["Hotel"], text="+ Add Note", command=add_hotel_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    ).pack(pady=10)

    # === Miscellaneous Tab ===
    misc_frame = tk.Frame(tabs["Misc"], bg="black")
    misc_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    def add_misc_note_box(default_text=""):
        frame = tk.Frame(misc_frame, bg="black")
        label = tk.Label(frame, text=f"Misc Note #{len(misc_boxes)+1}:", fg="white", bg="black", font=label_font)
        textbox = tk.Text(frame, height=4, width=80, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font, wrap=tk.WORD)
        if default_text:
            textbox.insert("1.0", default_text)
        label.pack(anchor="w")
        textbox.pack(fill="both", expand=True, padx=5)
        frame.pack(pady=5, fill="x")
        configure_text_box(textbox)
        misc_boxes.append(textbox)
        frame.pack(pady=5, fill="x")
    add_misc_note_box()

    tk.Button(
        tabs["Misc"], text="+ Add Note", command=add_misc_note_box,
        bg="white", fg="black", font=("Helvetica", 10, "bold")
    ).pack(pady=10)

    # === Security Tab === (Changed from CSC Log)
    csc_tab = tabs["Security"]  # Changed from "CSC Log"
    csc_frame = tk.Frame(csc_tab, bg="black")
    csc_frame.pack(fill="both", expand=True, padx=10, pady=(10, 0))

    # Remove the local csc_entries and csc_shifts declarations since they're now global
    for shift in csc_shifts:
        section = tk.LabelFrame(csc_frame, text=shift, fg="white", bg="black", font=label_font, labelanchor="n", padx=10, pady=10)
        section.pack(fill="x", pady=5)

        # Requested
        req_label = tk.Label(section, text="Staff Requested:", fg="white", bg="black", font=label_font)
        req_entry = tk.Entry(section, width=10, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        req_label.grid(row=0, column=0, sticky="w")
        req_entry.grid(row=0, column=1, padx=5)

        # Present
        pres_label = tk.Label(section, text="Staff Present:", fg="white", bg="black", font=label_font)
        pres_entry = tk.Entry(section, width=10, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        pres_label.grid(row=0, column=2, sticky="w")
        pres_entry.grid(row=0, column=3, padx=5)

        # Names (optional)
        names_label = tk.Label(section, text="Names (comma-separated):", fg="white", bg="black", font=label_font)
        names_entry = tk.Entry(section, width=40, bg=entry_bg, fg=entry_fg, insertbackground="white", font=entry_font)
        names_label.grid(row=1, column=0, columnspan=2, sticky="w")
        names_entry.grid(row=1, column=2, columnspan=2, sticky="we", pady=5)

        csc_entries[shift] = {
            "requested": req_entry,
            "present": pres_entry,
            "names": names_entry
        }

    # === Generate Report Button ===
    submit_btn = tk.Button(
        root, text="Generate Report", command=generate_report,
        bg="white", fg="black", font=("Helvetica", 12, "bold"), padx=10, pady=6
    )
    submit_btn.pack(pady=10)

# === Generate Report Logic ===
def generate_report():
    try:
        doc = Document()
        # Include building name in the heading
        doc.add_heading(f'{building.upper()}\nBUILDING MANAGER\'S NIGHT REPORT', level=1)
        
        # Add bold paragraphs with regular user input
        def add_bold_para_with_input(bold_text, user_input):
            p = doc.add_paragraph()
            p.add_run(f"{bold_text}: ").bold = True
            p.add_run(user_input)
        
        add_bold_para_with_input("Date", entries["date"].get())
        add_bold_para_with_input("Shift Hours", entries["shift_hours"].get())
        add_bold_para_with_input("Building Manager(s)", entries["bms"].get())
        
        # Red Gym specific report format
        if building == "Red Gym":
            # Bold section headers
            p = doc.add_paragraph()
            p.add_run("\nNotes:").bold = True
            
            def add_bold_section_header(text):
                p = doc.add_paragraph()
                p.add_run(text).bold = True
                return p
                
            # Function to add indented paragraphs using Inches for better control
            def add_indented_paragraph(number, content):
                p = doc.add_paragraph("")
                p.paragraph_format.left_indent = Inches(0.25)
                p.add_run(f"{number}. ").bold = True
                p.add_run(content)
                return p
            
            def add_sub_indented_paragraph(letter, content):
                p = doc.add_paragraph("")
                p.paragraph_format.left_indent = Inches(0.75)
                p.add_run(f"{letter}. ").bold = True
                p.add_run(content)
                return p

            note_counter = 1

            # Building Traffic
            add_bold_section_header("Building Traffic")
            has_traffic_notes = False
            for box in building_traffic_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_traffic_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            if not has_traffic_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # Security
            add_bold_section_header("Security")
            
            # Building Tours
            tours_content = red_gym_building_tours_box.get("1.0", "end").strip()
            tours_text = f"Building Tours: {tours_content}" if tours_content else "Building Tours:"
            add_indented_paragraph(note_counter, tours_text)
            note_counter += 1
            
            # Deviations
            num_deviations = int(red_gym_deviations_entry.get() or "0")
            deviation_text = f"There were {num_deviations} deviations from the standard building locking protocol today."
            
            # Create paragraph for deviations with bold formatting
            p = doc.add_paragraph("")
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(f"{note_counter}. ").bold = True
            p.add_run("There were ")
            p.add_run(str(num_deviations)).bold = True
            p.add_run(" deviations from the standard building locking protocol today.")
            
            if num_deviations > 0:
                for i, box in enumerate(red_gym_deviation_boxes):
                    content = box.get("1.0", "end").strip()
                    letter = chr(97 + i)  # a, b, c, etc.
                    add_sub_indented_paragraph(letter, content)
            
            note_counter += 1
            
            # Door check
            door_time = red_gym_door_check_time.get().strip()   
            door_day_type = red_gym_door_check_day_type.get()
            
            # Create paragraph for door check with bold formatting
            p = doc.add_paragraph("")
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(f"{note_counter}. ").bold = True
            p.add_run("I checked to confirm that the Building Doors were locked and the door swipe scanner was red at ")
            p.add_run(door_time).bold = True
            p.add_run(" on a ")
            p.add_run(door_day_type).bold = True
            p.add_run(".")
            
            note_counter += 1

            # Mail
            add_bold_section_header("Mail")
            has_mail_notes = False
            for box in red_gym_mail_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_mail_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            if not has_mail_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # Miscellaneous
            add_bold_section_header("Miscellaneous")
            has_misc_notes = False
            for box in red_gym_misc_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_misc_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            if not has_misc_notes:
                add_indented_paragraph(note_counter, "None")
                note_counter += 1

        else:
            # Existing code for Memorial Union and Union South
            # Terrace Manager(s) only for Memorial Union
            if building == "Memorial Union":
                add_bold_para_with_input("Terrace Manager(s)", entries["terrace_managers"].get())
            add_bold_para_with_input("Event Manager(s)", entries["eventmanagers"].get())
            add_bold_para_with_input("Guest Service Specialist", entries["gss"].get())
            add_bold_para_with_input("Operation Manager(s)", entries["operation_managers"].get())
            add_bold_para_with_input("Custodial Supervisor(s)", entries["custodial"].get())
            add_bold_para_with_input("Production Supervisor(s)", entries["production"].get())
            add_bold_para_with_input("Retail & Dining Supervisor(s)", entries["retail"].get())
            add_bold_para_with_input("Catering Supervisor(s)", entries["catering"].get())
            add_bold_para_with_input("CAVR Desk Staff", entries["cavr"].get())

            # Bold section headers
            p = doc.add_paragraph()
            p.add_run("\nNotes:").bold = True
            
            def add_bold_section_header(text):
                p = doc.add_paragraph()
                p.add_run(text).bold = True
                return p
                
            # Function to add indented paragraphs using Inches for better control
            def add_indented_paragraph(number, content):
                # Create a paragraph with the content and number
                p = doc.add_paragraph("")
                p.paragraph_format.left_indent = Inches(0.25)  # Change from 0.5 to 0.25 inch indent
                
                # Add the number with bold formatting
                p.add_run(f"{number}. ").bold = True
                
                # Add the content directly
                p.add_run(content)
                
                return p

            # Start a global counter for note numbering across all sections
            note_counter = 1

            add_bold_section_header("Building Traffic")
            
            # Add building traffic notes with global counter
            has_traffic_notes = False
            for box in building_traffic_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_traffic_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_traffic_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            add_bold_section_header("Mechanical/Repairs/Custodial")
            
            # Add mechanical notes with continuing counter
            has_mechanical_notes = False
            for box in mechanical_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_mechanical_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_mechanical_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            add_bold_section_header("Production Services (Meetings, Events, Set-ups, AV)")
            
            # Add production notes with continuing counter
            has_production_notes = False
            for box in production_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_production_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_production_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1
                
            # === Decibel reading and Security table modifications ===
            
            if decibel_entries:
                add_bold_section_header("Decibel Readings")
                # Create a table for decibel readings
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Add header row - bold the headers
                header_cells = table.rows[0].cells
                header_cells[0].paragraphs[0].add_run("Time").bold = True
                header_cells[1].paragraphs[0].add_run("Reading (dB)").bold = True
                header_cells[2].paragraphs[0].add_run("Location").bold = True
                
                # Center-align header cells
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 1  # 1 = CENTER

                # Add data rows
                for time_entry, reading_entry, location_entry in decibel_entries:
                    time = time_entry.get().strip()
                    reading = reading_entry.get().strip()
                    location = location_entry.get().strip()
                    if time and reading and location and time != "Time" and reading != "Reading (db)" and location != "Location":
                        row_cells = table.add_row().cells
                        row_cells[0].text = time
                        row_cells[1].text = reading
                        row_cells[2].text = location
                        
                        # Center-align all cells in this row
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = 1  # 1 = CENTER

            add_bold_section_header("Patron Services (Membership, Patron Assistance, Problem Patrons)")
            
            # Add patron service notes with continuing counter
            has_patron_notes = False
            for box in patron_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_patron_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_patron_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            add_bold_section_header("Access/Lock/Unlock")

            # Special case for access notes - bold the user input instead
            access_notes = [
                {
                    "text": f"At the early check, the loading dock arm gate was {access_inputs['early_gate'].get().lower()} at {access_inputs['early_time'].get()}.",
                    "bold_parts": [access_inputs['early_gate'].get().lower(), access_inputs['early_time'].get()]
                },
                {
                    "text": f"At the closing check, the loading dock arm gate was {access_inputs['close_gate'].get().lower()} at {access_inputs['close_time'].get()}.",
                    "bold_parts": [access_inputs['close_gate'].get().lower(), access_inputs['close_time'].get()]
                },
                {
                    "text": f"At the closing check, the HID scanners were {access_inputs['hid_status'].get().lower()}.",
                    "bold_parts": [access_inputs['hid_status'].get().lower()]
                },
                {
                    "text": f"I {access_inputs['door_status'].get().lower()} secured the loading dock overhead door for the night.",
                    "bold_parts": [access_inputs['door_status'].get().lower()]
                }
            ]

            # Auto-generated sentences with bold user inputs and continuing counter with indentation
            for note_data in access_notes:
                p = doc.add_paragraph("")
                # Indent the entire paragraph using Inches
                p.paragraph_format.left_indent = Inches(0.25)  # Change from 0.5 to 0.25
                
                # Add the number with bold formatting
                p.add_run(f"{note_counter}. ").bold = True
                
                sentence = note_data["text"]
                bold_parts = note_data["bold_parts"]
                
                current_pos = 0
                for bold_part in bold_parts:
                    if bold_part in sentence[current_pos:]:
                        # Find position of bold part
                        part_pos = sentence.find(bold_part, current_pos)
                        
                        # Add text before the bold part
                        if part_pos > current_pos:
                            p.add_run(sentence[current_pos:part_pos])
                        
                        # Add the bold part
                        p.add_run(bold_part).bold = True
                        
                        # Update position
                        current_pos = part_pos + len(bold_part)
                
                # Add any remaining text after the last bold part
                if current_pos < len(sentence):
                    p.add_run(sentence[current_pos:])
                    
                note_counter += 1
                
            # User-entered access notes
            for box in access_note_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1

            add_bold_section_header("Cash Office")
            
            # Add cash office notes with continuing counter
            has_cash_notes = False
            for box in cash_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_cash_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_cash_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # Only include Memorial Union specific sections if the building is Memorial Union
            if building == "Memorial Union":
                add_bold_section_header("Carding Runs")
                
                # Add carding notes with continuing counter
                has_carding_notes = False
                for box in carding_boxes:
                    content = box.get("1.0", "end").strip()
                    if content:
                        has_carding_notes = True
                        add_indented_paragraph(note_counter, content)
                        note_counter += 1
                
                # Add an empty numbered note if no content
                if not has_carding_notes:
                    add_indented_paragraph(note_counter, "")
                    note_counter += 1

                add_bold_section_header("Terrace Traffic")
                
                # Add terrace traffic notes with continuing counter
                has_terrace_notes = False
                for box in terrace_boxes:
                    content = box.get("1.0", "end").strip()
                    if content:
                        has_terrace_notes = True
                        add_indented_paragraph(note_counter, content)
                        note_counter += 1
                
                # Add an empty numbered note if no content
                if not has_terrace_notes:
                    add_indented_paragraph(note_counter, "")
                    note_counter += 1

                add_bold_section_header("Terrace Enforcement")
                
                # Add enforcement notes with continuing counter
                has_enforcement_notes = False
                enforcement_index = 0  # Track index for both images and text notes
                
                for i, box in enumerate(enforcement_boxes):
                    content = box.get("1.0", "end").strip()
                    
                    # Check if this is an image entry (has corresponding image path)
                    if i < len(enforcement_images) and enforcement_images[i].get():
                        # This is an image with description
                        image_path = enforcement_images[i].get()
                        
                        # Add the image to the document
                        try:
                            # Insert image first with width fitting document and height of 400
                            doc.add_picture(image_path, width=Inches(6.5), height=Inches(250/72))  # 6.5 inches width (standard doc width), 400px height converted to inches
                            
                            # Add numbered description after the image
                            if content:
                                p = doc.add_paragraph("")
                                p.paragraph_format.left_indent = Inches(0.25)
                                p.add_run(f"{note_counter}. ").bold = True
                                p.add_run(content)
                            else:
                                # If no description, still add the number
                                p = doc.add_paragraph("")
                                p.paragraph_format.left_indent = Inches(0.25)
                                p.add_run(f"{note_counter}. ").bold = True
                            
                            has_enforcement_notes = True
                            note_counter += 1
                        except Exception as e:
                            # If image fails to load, just add the text
                            if content:
                                add_indented_paragraph(note_counter, content)
                                has_enforcement_notes = True
                                note_counter += 1
                    else:
                        # Regular text note
                        if content:
                            has_enforcement_notes = True
                            add_indented_paragraph(note_counter, content)
                            note_counter += 1
                

                # Add an empty numbered note if no content
                if not has_enforcement_notes:
                    add_indented_paragraph(note_counter, "")
                    note_counter += 1

           
            # === Dining Service & Markets ===
            add_bold_section_header("Dining Service & Markets")
            
            # Add dining notes with continuing counter
            has_dining_notes = False
            for box in dining_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_dining_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_dining_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # === Hotel ===
            add_bold_section_header("Hotel")
            
            # Add hotel notes with continuing counter
            has_hotel_notes = False
            for box in hotel_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_hotel_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_hotel_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # === Miscellaneous ===
            add_bold_section_header("Miscellaneous")
            
            # Add miscellaneous notes with continuing counter
            has_misc_notes = False
            for box in misc_boxes:
                content = box.get("1.0", "end").strip()
                if content:
                    has_misc_notes = True
                    add_indented_paragraph(note_counter, content)
                    note_counter += 1
            
            # Add an empty numbered note if no content
            if not has_misc_notes:
                add_indented_paragraph(note_counter, "")
                note_counter += 1

            # === Security Section === Only for Memorial Union and Union South
            add_bold_section_header("Security")  # Changed from "CSC Log"

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            # Bold the headers
            hdr_cells[0].paragraphs[0].add_run("Shift").bold = True
            hdr_cells[1].paragraphs[0].add_run("Staff Requested").bold = True
            hdr_cells[2].paragraphs[0].add_run("Staff Present").bold = True
            
            # Center-align header cells
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1  # 1 = CENTER

            for shift in csc_shifts:
                req_val = csc_entries[shift]["requested"].get().strip()
                pres_val = csc_entries[shift]["present"].get().strip()
                names_val = csc_entries[shift]["names"].get().strip()

                # Format names in parentheses only if present is a number
                pres_display = pres_val
                if pres_val.isdigit() and names_val:
                    pres_display = f"{pres_val} ({names_val})"

                row = table.add_row().cells
                
                # Bold the shift name - no tab character needed
                shift_run = row[0].paragraphs[0].add_run(shift)
                shift_run.bold = True
                
                # Set text directly
                row[1].text = req_val if req_val else "-"
                row[2].text = pres_display if pres_display else "-"
                
                # Center-align all cells in this row
                for cell in row:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 1  # 1 = CENTER

        # Excel Tally Update - Only for non-Red Gym buildings
        if building != "Red Gym":
            try:
                # Extract year and date for folder and filename
                user_date = entries["date"].get()
                parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
                current_year = parsed_date.strftime("%Y")
                current_month = parsed_date.strftime("%B")
                # Build master tag set from all tab tag lists
                master_tag_set = set()
                for taglist in [MECHANICAL_TAG_OPTIONS, PRODUCTION_TAG_OPTIONS, PATRON_TAG_OPTIONS, ACCESS_TAG_OPTIONS, CASH_TAG_OPTIONS, DINING_TAG_OPTIONS, TERRACE_TAG_OPTIONS]:
                    master_tag_set.update(tag for tag in taglist if tag != "None")
                # Add special-case tags
                master_tag_set.add("Hotel Request")
                master_tag_set.add("Carding Support/Lead Carding")
                master_tag_set.add("Decibel checked")
                master_tag_set.add("Patron Services/inquires/General Assistance")
                categories = sorted(master_tag_set)
                months = [
                    "January", "February", "March", "April", "May", "June",
                    "July", "August", "September", "October", "November", "December"
                ]
                # Create year and building folders if not exist
                year_dir = os.path.join(os.getcwd(), current_year)
                building_dir = os.path.join(year_dir, building)
                os.makedirs(building_dir, exist_ok=True)
                # Save tally as building_Tally_YYYY.xlsx in building folder
                building_short = "MU" if building == "Memorial Union" else "US"
                tally_filename = f"{building_short}_Tally_{current_year}.xlsx"
                tally_path = os.path.join(building_dir, tally_filename)
                # Load or create the Excel file for the current year/building
                if os.path.exists(tally_path):
                    df = pd.read_excel(tally_path, index_col=0)
                    for category in categories:
                        if category not in df.index:
                            df.loc[category] = [0] * len(df.columns)
                    # Remove any rows not in master list
                    for row in list(df.index):
                        if row not in categories:
                            df = df.drop(row)
                else:
                    df = pd.DataFrame(0, index=categories, columns=months)
                # Tally tags for all notes
                tag_counts = {cat: 0 for cat in categories}
                # Mechanical
                for tag_var_list, box in zip(mechanical_note_tags, mechanical_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Production
                for tag_var_list, box in zip(production_note_tags, production_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Patron
                for tag_var_list, box in zip(patron_note_tags, patron_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    if tags:
                        for tag in tags:
                            tag_counts[tag] += 1
                    else:
                        # If no tag but text, count as General Assistance
                        tag_counts["Patron Services/inquires/General Assistance"] += 1
                # Access
                for tag_var_list, box in zip(access_note_tags, access_note_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Cash
                for tag_var_list, box in zip(cash_note_tags, cash_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Dining
                for tag_var_list, box in zip(dining_note_tags, dining_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Memorial Union specific sections
                if building == "Memorial Union":
                    # Terrace
                    for tag_var_list, box in zip(terrace_note_tags, terrace_boxes):
                        content = box.get("1.0", "end").strip()
                        if not content:
                            continue
                        tags = set(var.get() for var in tag_var_list if var.get() != "None")
                        for tag in tags:
                            tag_counts[tag] += 1
                    # Enforcement
                    for tag_var_list, box in zip(enforcement_note_tags, enforcement_boxes):
                        content = box.get("1.0", "end").strip()
                        if not content:
                            continue
                        tags = set(var.get() for var in tag_var_list if var.get() != "None")
                        for tag in tags:
                            tag_counts[tag] += 1
                    # Alumni
                    for tag_var_list, box in zip(alumni_note_tags, alumni_boxes):
                        content = box.get("1.0", "end").strip()
                        if not content:
                            continue
                        tags = set(var.get() for var in tag_var_list if var.get() != "None")
                        for tag in tags:
                            tag_counts[tag] += 1
                    # Pier
                    for tag_var_list, box in zip(pier_note_tags, pier_boxes):
                        content = box.get("1.0", "end").strip()
                        if not content:
                            continue
                        tags = set(var.get() for var in tag_var_list if var.get() != "None")
                        for tag in tags:
                            tag_counts[tag] += 1
                # Hotel (special case)
                for box in hotel_boxes:
                    content = box.get("1.0", "end").strip()
                    if content:
                        tag_counts["Hotel Request"] +=   1
                # Decibel (special case)
                if any(time.get().strip() and time.get().strip() != "Time" for time, _, _ in decibel_entries):
                    tag_counts["Decibel checked"] += 1
                # Write tallies to Excel
                for category, count in tag_counts.items():
                    if category in df.index and current_month in df.columns:
                        df.loc[category, current_month] += count
                df = df.loc[categories]  # Ensure order
                df.to_excel(tally_path)
            except Exception as e:
                messagebox.showerror("Excel Error", f"Failed to update Excel tally: {e}")
        else:
            # Red Gym Excel Tally Update
            try:
                # Extract year and date for folder and filename
                user_date = entries["date"].get()
                parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
                current_year = parsed_date.strftime("%Y")
                current_month = parsed_date.strftime("%B")
                # Build Red Gym tag set
                red_gym_tag_set = set()
                for taglist in [RED_GYM_MISC_TAG_OPTIONS]:
                    red_gym_tag_set.update(tag for tag in taglist if tag != "None")
                categories = sorted(red_gym_tag_set)
                months = [
                    "January", "February", "March", "April", "May", "June",
                    "July", "August", "September", "October", "November", "December"
                ]
                # Create year and building folders if not exist
                year_dir = os.path.join(os.getcwd(), current_year)
                building_dir = os.path.join(year_dir, building)
                os.makedirs(building_dir, exist_ok=True)
                # Save tally as RG_Tally_YYYY.xlsx in Red Gym folder
                tally_filename = f"RG_Tally_{current_year}.xlsx"
                tally_path = os.path.join(building_dir, tally_filename)
                # Load or create the Excel file for the current year/building
                if os.path.exists(tally_path):
                    df = pd.read_excel(tally_path, index_col=0)
                    for category in categories:
                        if category not in df.index:
                            df.loc[category] = [0] * len(df.columns)
                    # Remove any rows not in master list
                    for row in list(df.index):
                        if row not in categories:
                            df = df.drop(row)
                else:
                    df = pd.DataFrame(0, index=categories, columns=months)
                # Tally tags for Red Gym misc notes
                tag_counts = {cat: 0 for cat in categories}
                # Red Gym Misc
                for tag_var_list, box in zip(misc_note_tags, red_gym_misc_boxes):
                    content = box.get("1.0", "end").strip()
                    if not content:
                        continue
                    tags = set(var.get() for var in tag_var_list if var.get() != "None")
                    for tag in tags:
                        tag_counts[tag] += 1
                # Write tallies to Excel
                for category, count in tag_counts.items():
                    if category in df.index and current_month in df.columns:
                        df.loc[category, current_month] += count
                df = df.loc[categories]  # Ensure order
                df.to_excel(tally_path)
            except Exception as e:
                messagebox.showerror("Excel Error", f"Failed to update Red Gym Excel tally: {e}")

        # Save report as MM-DD-YY.docx in the correct folder (for all building types)
        user_date = entries["date"].get()
        parsed_date = datetime.strptime(user_date, "%A, %B %d, %Y")
        current_year = parsed_date.strftime("%Y")
        # Create year and building folders if not exist
        year_dir = os.path.join(os.getcwd(), current_year)
        building_dir = os.path.join(year_dir, building)
        os.makedirs(building_dir, exist_ok=True)
        report_filename = f"{parsed_date.month}-{parsed_date.day}-{str(parsed_date.year)[2:]}.docx"
        report_path = os.path.join(building_dir, report_filename)
        doc.save(report_path)
        
        messagebox.showinfo("Success", f"Report saved as {report_path}")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# Wrap the initial UI setup in a function to be called after mainloop starts
def start_app():
    select_building()

# Schedule the app start after the mainloop starts
root.after(0, start_app)

root.mainloop()
